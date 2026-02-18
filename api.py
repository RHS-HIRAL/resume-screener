"""
Resume Screener API — FastAPI wrapper around the parsing logic.
Power Automate sends PDF/DOCX attachments here, gets JSON back.

Run:  uvicorn api:app --host 0.0.0.0 --port 8000
Deps: pip install fastapi uvicorn python-multipart openpyxl python-docx PyPDF2 python-dotenv google-generativeai groq pandas
"""

import io
import json
import re
import os
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional

import pandas as pd
import PyPDF2 as pdf
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse
import google.generativeai as genai
from groq import Groq

# ========== ENVIRONMENT ==========
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

GROQ_KEY_PRIMARY = os.getenv("GROQ_API_KEY")
GROQ_KEY_ALT = os.getenv("GROQ_API_KEY_ALT")

_GROQ_CLIENT_PRIMARY = Groq(api_key=GROQ_KEY_PRIMARY) if GROQ_KEY_PRIMARY else None
_GROQ_CLIENT_ALT = Groq(api_key=GROQ_KEY_ALT) if GROQ_KEY_ALT else None

# Output directory
OUTPUT_DIR = Path(__file__).parent / "resume_screener_output"
OUTPUT_DIR.mkdir(exist_ok=True)
RESULTS_EXCEL = OUTPUT_DIR / "resume_screening_results.xlsx"
LOG_FILE = OUTPUT_DIR / "llm_calls.jsonl"

EXCEL_COLUMNS = [
    "Name",
    "Resume Match Score",
    "Email",
    "Phone",
    "LinkedIn",
    "Current Job Title",
    "Current Organization",
    "Total Experience",
    "Total Jobs",
    "Average Tenure",
    "Location",
    "Role in JD",
    "Resume Summary",
    "Job History (org-title-jobDuration)",
    "Other Socials",
    "Processed At",
    "Job Description",
]

MATCH_SCORE_THRESHOLD = 60  # Default: only include candidates >= 60%

app = FastAPI(title="Resume Screener API", version="1.0")


# ========== LOGGING ==========
def log_failure(record: dict):
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(json.dumps(record, ensure_ascii=False) + "\n")


# ========== TEXT EXTRACTION ==========
def extract_text_from_pdf_bytes(content: bytes) -> str:
    reader = pdf.PdfReader(io.BytesIO(content))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_text_from_docx_bytes(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(filename: str, content: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf_bytes(content)
    elif name.endswith(".docx"):
        return extract_text_from_docx_bytes(content)
    return ""


# ========== LLM PROVIDERS ==========
def _call_gemini(prompt):
    model = genai.GenerativeModel("gemini-2.5-pro")
    return model.generate_content(prompt).text


def _call_groq(prompt, client, model_name):
    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
    )
    return response.choices[0].message.content


def _build_fallback_chain():
    chain = [("Gemini", None, None)]
    if _GROQ_CLIENT_PRIMARY:
        chain.append(("Groq-Primary/compound", _GROQ_CLIENT_PRIMARY, "groq/compound"))
    if _GROQ_CLIENT_ALT:
        chain.append(("Groq-Alt/compound", _GROQ_CLIENT_ALT, "groq/compound"))
    if _GROQ_CLIENT_PRIMARY:
        chain.append(
            ("Groq-Primary/llama", _GROQ_CLIENT_PRIMARY, "llama-3.3-70b-versatile")
        )
    if _GROQ_CLIENT_ALT:
        chain.append(("Groq-Alt/llama", _GROQ_CLIENT_ALT, "llama-3.3-70b-versatile"))
    return chain


FALLBACK_CHAIN = _build_fallback_chain()


def get_llm_response(prompt: str) -> str:
    ts = datetime.now().isoformat()
    for label, client, model_name in FALLBACK_CHAIN:
        try:
            if label == "Gemini":
                return _call_gemini(prompt)
            else:
                return _call_groq(prompt, client, model_name)
        except Exception as e:
            log_failure(
                {
                    "ts": ts,
                    "provider": label,
                    "error": str(e),
                    "prompt_chars": len(prompt),
                }
            )
            continue

    log_failure(
        {"ts": ts, "provider": "ALL_FAILED", "error": "Exhausted all providers"}
    )
    return "ERROR: All providers failed"


# ========== JSON CLEANING ==========
def _clean_llm_json(text):
    text = re.sub(r"^```(?:json)?\s*\n?", "", text.strip())
    text = re.sub(r"\n?```\s*$", "", text.strip())
    match = re.search(r"\{.*\}", text, re.DOTALL)
    return match.group(0) if match else None


# ========== BACKUP EXTRACTION ==========
_SKIP_WORDS = frozenset(
    [
        "resume",
        "cv",
        "email",
        "phone",
        "address",
        "linkedin",
        "objective",
        "summary",
        "experience",
    ]
)


def extract_name_backup(text):
    for line in text.split("\n")[:8]:
        line = line.strip()
        if not line or any(kw in line.lower() for kw in _SKIP_WORDS):
            continue
        if (
            2 <= len(line.split()) <= 4
            and re.match(r"^[A-Za-z\s\.]+$", line)
            and not line.isupper()
        ):
            return line
    return "N/A"


def extract_contact_backup(text):
    email_m = re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)
    email = email_m.group(0) if email_m else "N/A"
    phone = "N/A"
    for pat in [
        r"(?:\+91|91)[\s\-]?[6-9]\d{9}",
        r"[6-9]\d{9}",
        r"\d{3}[\s\-]?\d{3}[\s\-]?\d{4}",
    ]:
        m = re.search(pat, text)
        if m:
            phone = m.group(0)
            break
    li_m = re.search(r"linkedin\.com/in/([A-Za-z0-9\-]+)", text, re.IGNORECASE)
    linkedin = f"linkedin.com/in/{li_m.group(1)}" if li_m else "N/A"
    return email, phone, linkedin


def extract_experience_backup(text):
    for pat in [
        r"(?:total|overall)[\s\-:]*(?:experience|exp)[\s\-:]*(\d+(?:\.\d+)?)\s*(?:years?|yrs?)",
        r"(\d+(?:\.\d+)?)\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)",
    ]:
        m = re.search(pat, text, re.IGNORECASE)
        if m and 0 <= float(m.group(1)) <= 50:
            return f"{float(m.group(1))} years"
    return "N/A"


# ========== HELPERS ==========
def _count_jobs(jh):
    return len([e for e in jh.split(" | ") if e.strip()]) if jh and jh != "N/A" else 0


def _extract_years(s):
    if not s or s == "N/A":
        return 0.0
    m = re.search(r"(\d+(?:\.\d+)?)", s)
    return float(m.group(1)) if m else 0.0


def _avg_tenure(total_exp, total_jobs):
    try:
        years = _extract_years(total_exp)
        jobs = (
            int(total_jobs)
            if str(total_jobs).isdigit()
            else _count_jobs(str(total_jobs))
        )
        return f"{years / jobs:.1f} years" if years > 0 and jobs > 0 else "N/A"
    except Exception:
        return "N/A"


def _pick(ai_val, backup_val):
    return ai_val if ai_val and ai_val != "N/A" else backup_val


def _parse_score(score_str) -> float:
    """Extract numeric score from strings like '75%', '75', 'N/A', etc."""
    if not score_str or score_str in ("N/A", "Error"):
        return 0.0
    m = re.search(r"(\d+(?:\.\d+)?)", str(score_str))
    return float(m.group(1)) if m else 0.0


# ========== CORE PARSING ==========
def parse_single_resume(
    filename: str, resume_bytes: bytes, jd_text: str, jd_name: str, batch_ts: str
) -> dict:
    """Parse one resume and return a result dict"""
    try:
        resume_text = extract_text(filename, resume_bytes)
        if not resume_text or len(resume_text) < 50:
            return _error_record(filename, "Text extraction failed", jd_name, batch_ts)

        # AI extraction
        prompt = f"""You are an expert resume parser. Extract the following information from this resume with 100% accuracy.

CRITICAL INSTRUCTIONS:
1. Return ONLY valid JSON in the exact format shown below
2. If any field is not found, use "N/A"
3. For name: Look for the actual person's name (usually at the top)
4. For location: Look for city, state, or geographic location (NOT technical terms)
5. For current job: Find the most recent job title and company
6. For LinkedIn: Look for linkedin.com/in/ URLs or LinkedIn profile mentions
7. For job history: Format as "Company - Title - Duration" separated by " | "
8. For total experience: Extract as number only (e.g., "5" for 5 years)
9. Count all distinct jobs/positions in the work history

RESUME TEXT:
{resume_text[:5000]}

JOB DESCRIPTION:
{jd_text[:2000]}

Return JSON in this EXACT format:
{{
    "name": "Full Name Here",
    "role_in_jd": "Job title from JD",
    "current_job_title": "Current position title",
    "current_organization": "Current company name",
    "location": "City, State or geographic location",
    "phone": "Phone number",
    "email": "Email address",
    "linkedin": "LinkedIn profile URL or N/A",
    "other_socials": "Other social media or N/A",
    "total_experience": "X years",
    "job_history": "Company1 - Title1 - Duration1 | Company2 - Title2 - Duration2",
    "total_jobs": "Number of jobs",
    "match_score": "Calculate a strict 0-100% score based ONLY on the JD requirements. DEDUCT points if the candidate has the years of experience but lacks the specific 'Must Have' technical skills listed. A score of 100% means they are a perfect fit. A score of 60% or lower means they lack critical skills. Do not inflate the score for over-qualification in irrelevant areas.",
    "summary": "Brief professional summary in 2-3 sentences"
}}"""

        response = get_llm_response(prompt)
        if response.startswith("ERROR:"):
            return _error_record(filename, response, jd_name, batch_ts)

        ai = {}
        try:
            cleaned = _clean_llm_json(response)
            if cleaned:
                ai = json.loads(cleaned)
        except (json.JSONDecodeError, ValueError) as e:
            log_failure({"ts": batch_ts, "error": f"JSON parse: {e}", "file": filename})

        # Backup extraction
        b_name = extract_name_backup(resume_text)
        b_email, b_phone, b_linkedin = extract_contact_backup(resume_text)
        b_exp = extract_experience_backup(resume_text)

        job_history = ai.get("job_history", "N/A")
        total_jobs = ai.get("total_jobs", _count_jobs(job_history))
        total_exp = _pick(ai.get("total_experience"), b_exp)

        return {
            "Name": _pick(ai.get("name"), b_name),
            "Resume Match Score": ai.get("match_score", "N/A"),
            "Email": _pick(ai.get("email"), b_email),
            "Phone": _pick(ai.get("phone"), b_phone),
            "LinkedIn": _pick(ai.get("linkedin"), b_linkedin),
            "Current Job Title": ai.get("current_job_title", "N/A"),
            "Current Organization": ai.get("current_organization", "N/A"),
            "Total Experience": total_exp,
            "Total Jobs": str(total_jobs)
            if str(total_jobs).isdigit()
            else str(_count_jobs(job_history)),
            "Average Tenure": _avg_tenure(total_exp, total_jobs),
            "Location": ai.get("location", "N/A"),
            "Role in JD": ai.get("role_in_jd", "N/A"),
            "Resume Summary": ai.get("summary", "N/A"),
            "Job History (org-title-jobDuration)": job_history,
            "Other Socials": ai.get("other_socials", "N/A"),
            "Processed At": batch_ts,
            "Job Description": jd_name,
        }

    except Exception as e:
        log_failure({"ts": batch_ts, "error": f"parse_single_resume: {filename}: {e}"})
        return _error_record(filename, f"Processing error: {e}", jd_name, batch_ts)


def _error_record(filename, error_msg, jd_name, batch_ts):
    record = {col: "N/A" for col in EXCEL_COLUMNS}
    record.update(
        {
            "Name": filename,
            "Resume Match Score": "Error",
            "Resume Summary": error_msg,
            "Processed At": batch_ts,
            "Job Description": jd_name,
        }
    )
    return record


# ========== EXCEL OPERATIONS ==========
def save_to_excel(results: list) -> Path:
    """Save results to Excel, append to existing if present. Returns file path."""
    new_df = pd.DataFrame(results).reindex(columns=EXCEL_COLUMNS)

    if RESULTS_EXCEL.exists():
        try:
            existing_df = pd.read_excel(RESULTS_EXCEL, engine="openpyxl")
            combined_df = pd.concat([existing_df, new_df], ignore_index=True)
        except Exception:
            combined_df = new_df
    else:
        combined_df = new_df

    combined_df.to_excel(RESULTS_EXCEL, index=False, engine="openpyxl")
    return RESULTS_EXCEL


def build_filtered_excel(results: list, threshold: float) -> Path:
    """Build a separate Excel with only candidates above the match score threshold."""
    filtered = [
        r for r in results if _parse_score(r.get("Resume Match Score")) >= threshold
    ]

    if not filtered:
        filtered = [
            {
                "Name": "No candidates met the threshold",
                "Resume Match Score": f"< {threshold}%",
                **{
                    col: "N/A"
                    for col in EXCEL_COLUMNS
                    if col not in ("Name", "Resume Match Score")
                },
            }
        ]

    df = pd.DataFrame(filtered).reindex(columns=EXCEL_COLUMNS)
    # Sort by match score descending
    df["_sort_score"] = df["Resume Match Score"].apply(lambda x: _parse_score(x))
    df = df.sort_values("_sort_score", ascending=False).drop(columns=["_sort_score"])

    filtered_path = (
        OUTPUT_DIR / f"filtered_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    )
    df.to_excel(filtered_path, index=False, engine="openpyxl")
    return filtered_path


# =============================================
#  API ENDPOINTS
# =============================================


@app.post("/parse-resumes")
async def parse_resumes(
    resumes: list[UploadFile] = File(..., description="PDF/DOCX resume files"),
    jd_file: Optional[UploadFile] = File(
        None, description="JD as PDF/DOCX (provide this OR jd_text)"
    ),
    jd_text: Optional[str] = Form(
        None, description="JD as plain text (provide this OR jd_file)"
    ),
    jd_name: str = Form("Job Description", description="JD identifier/name"),
    score_threshold: float = Form(
        MATCH_SCORE_THRESHOLD, description="Min match score % for filtered results"
    ),
):
    """
    Parse multiple resumes against a JD.

    Returns:
    - all_results: Full JSON array of every parsed resume
    - filtered_results: Only candidates >= score_threshold
    - excel_path: Path to the full results Excel
    - filtered_excel_path: Path to the filtered Excel (for emailing)
    - summary: Quick stats
    """
    # Resolve JD text
    if jd_file:
        jd_bytes = await jd_file.read()
        jd_text_resolved = extract_text(jd_file.filename, jd_bytes)
        if not jd_text_resolved:
            raise HTTPException(400, "Failed to extract text from JD file")
    elif jd_text:
        jd_text_resolved = jd_text
    else:
        raise HTTPException(400, "Provide either jd_file or jd_text")

    batch_ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Parse each resume
    all_results = []
    for resume_file in resumes:
        content = await resume_file.read()
        result = parse_single_resume(
            resume_file.filename, content, jd_text_resolved, jd_name, batch_ts
        )
        all_results.append(result)

    # Save full results
    save_to_excel(all_results)

    # Build filtered Excel for TA team
    filtered_path = build_filtered_excel(all_results, score_threshold)

    # Filter JSON
    filtered_results = [
        r
        for r in all_results
        if _parse_score(r.get("Resume Match Score")) >= score_threshold
    ]
    filtered_results.sort(
        key=lambda r: _parse_score(r.get("Resume Match Score")), reverse=True
    )

    total = len(all_results)
    passed = len(filtered_results)
    errors = sum(1 for r in all_results if r["Resume Match Score"] == "Error")

    return {
        "all_results": all_results,
        "filtered_results": filtered_results,
        "excel_path": str(RESULTS_EXCEL),
        "filtered_excel_path": str(filtered_path),
        "summary": {
            "total_resumes": total,
            "passed_threshold": passed,
            "rejected": total - passed - errors,
            "errors": errors,
            "threshold": score_threshold,
        },
    }


@app.post("/parse-single")
async def parse_single(
    resume: UploadFile = File(..., description="Single PDF/DOCX resume"),
    jd_text: str = Form(..., description="JD as plain text"),
    jd_name: str = Form("Job Description"),
):
    """Parse a single resume — lightweight endpoint for per-attachment processing in Power Automate."""
    content = await resume.read()
    batch_ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    result = parse_single_resume(resume.filename, content, jd_text, jd_name, batch_ts)
    result["_score_numeric"] = _parse_score(result.get("Resume Match Score"))
    return result


@app.get("/download/{filename}")
async def download_file(filename: str):
    """Download a generated Excel file."""
    path = OUTPUT_DIR / filename
    if not path.exists():
        raise HTTPException(404, f"File not found: {filename}")
    return FileResponse(
        path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.get("/health")
async def health():
    return {"status": "ok", "providers": [label for label, _, _ in FALLBACK_CHAIN]}
