import streamlit as st
import google.generativeai as genai
import os
import PyPDF2 as pdf
from dotenv import load_dotenv
import json
import re
from docx import Document
import pandas as pd
from datetime import datetime
from pathlib import Path
from groq import Groq
from resume_screener_pipeline.old_pipeline import ResumeScreenerPipeline

# ========== ENVIRONMENT & API SETUP ==========
load_dotenv()

genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

GROQ_KEY_PRIMARY = os.getenv("GROQ_API_KEY")
GROQ_KEY_ALT = os.getenv("GROQ_API_KEY_ALT")


# ========= LOG DIRECTORY SETUP =========
APP_DIR = Path(__file__).parent
LOG_DIR = APP_DIR / "resume_screener_logs"
LOG_DIR.mkdir(exist_ok=True)

LLM_LOG_FILE = LOG_DIR / "llm_calls.jsonl"
RESULTS_EXCEL = LOG_DIR / "resume_screening_results.xlsx"

EXCEL_COLUMNS = [
    "Name",
    "Resume Match Score",
    "Email",
    "Phone",
    "LinkedIn",
    "Current Job Title",
    "Current Organization",
    "Total Experience",
    "Total Jobs",
    "Average Tenure",
    "Location",
    "Role in JD",
    "Resume Summary",
    "Job History (org-title-jobDuration)",
    "Other Socials",
    "Processed At",
    "Job Description",
]


# ========== LOGGING (errors/failures only) ==========
def log_llm_failure(record: dict):
    """Write only failures/fallbacks to the JSONL log"""
    with open(LLM_LOG_FILE, "a", encoding="utf-8") as f:
        f.write(json.dumps(record, ensure_ascii=False) + "\n")


def log_batch_separator(batch_ts, jd_name, resume_count):
    with open(LLM_LOG_FILE, "a", encoding="utf-8") as f:
        f.write(
            f"\n--- BATCH {batch_ts} | JD: {jd_name} | Resumes: {resume_count} ---\n"
        )


# ========== SESSION STATE ==========
def init_session_state():
    for key, default in {
        "results": None,
        "df": None,
        "processing_done": False,
        "llm_request_count": 0,
        "pending_conflicts": None,
        "non_conflict_results": None,
        "conflict_resolved": False,
        "pipeline_stage1_results": None,
    }.items():
        if key not in st.session_state:
            st.session_state[key] = default


# ========== TEXT EXTRACTION ==========
def extract_text_from_pdf(uploaded_file):
    reader = pdf.PdfReader(uploaded_file)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_text_from_docx(uploaded_file):
    doc = Document(uploaded_file)
    parts = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(uploaded_file):
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    elif name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    return ""


# ========== LLM PROVIDERS ==========
def _call_gemini(prompt):
    model = genai.GenerativeModel("gemini-2.5-pro")
    response = model.generate_content(prompt)
    return response.text


def _call_groq(prompt, api_key, model_name):
    client = Groq(api_key=api_key)
    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
    )
    usage = response.usage
    tokens = {
        "input": usage.prompt_tokens if usage else None,
        "output": usage.completion_tokens if usage else None,
        "total": usage.total_tokens if usage else None,
    }
    return response.choices[0].message.content, tokens


def _build_fallback_chain():
    """Gemini → Groq-Primary/compound → Groq-Alt/compound → Groq-Primary/llama → Groq-Alt/llama"""
    chain = [("Gemini", None, None)]  # (label, api_key, model)
    if GROQ_KEY_PRIMARY:
        chain.append(("Groq-Primary/compound", GROQ_KEY_PRIMARY, "groq/compound"))
    if GROQ_KEY_ALT:
        chain.append(("Groq-Alt/compound", GROQ_KEY_ALT, "groq/compound"))
    if GROQ_KEY_PRIMARY:
        chain.append(
            ("Groq-Primary/llama", GROQ_KEY_PRIMARY, "llama-3.3-70b-versatile")
        )
    if GROQ_KEY_ALT:
        chain.append(("Groq-Alt/llama", GROQ_KEY_ALT, "llama-3.3-70b-versatile"))
    return chain


FALLBACK_CHAIN = _build_fallback_chain()


def get_llm_response(prompt):
    st.session_state.llm_request_count += 1
    req_num = st.session_state.llm_request_count
    ts = datetime.now().isoformat()
    prompt_chars = len(prompt)

    for label, api_key, model_name in FALLBACK_CHAIN:
        try:
            if label == "Gemini":
                text = _call_gemini(prompt)
                return text
            else:
                text, _ = _call_groq(prompt, api_key, model_name)
                return text

        except Exception as e:
            # Estimate tokens for failed calls (rough: 1 token ≈ 4 chars)
            est_input = prompt_chars // 4
            token_info = {"input_est": est_input, "output": 0, "total_est": est_input}

            # Try to get real token counts from Groq errors if available
            if label != "Gemini":
                try:
                    # Some Groq errors carry usage in the response
                    if hasattr(e, "body") and isinstance(e.body, dict):
                        usage = e.body.get("usage", {})
                        if usage:
                            token_info = {
                                "input": usage.get("prompt_tokens", est_input),
                                "output": usage.get("completion_tokens", 0),
                                "total": usage.get("total_tokens", est_input),
                            }
                except Exception:
                    pass

            log_llm_failure(
                {
                    "req": req_num,
                    "ts": ts,
                    "provider": label,
                    "error": str(e),
                    "tokens": token_info,
                    "prompt_chars": prompt_chars,
                }
            )
            continue

    # All failed — log final failure
    log_llm_failure(
        {
            "req": req_num,
            "ts": ts,
            "provider": "ALL_FAILED",
            "error": "Exhausted all providers",
            "prompt_chars": prompt_chars,
        }
    )
    return "ERROR: All providers failed"


# ========== AI FIELD EXTRACTION ==========
def extract_all_fields_with_ai(resume_text, jd_text):
    prompt = f"""You are an expert resume parser. Extract the following information from this resume with 100% accuracy.

CRITICAL INSTRUCTIONS:
1. Return ONLY valid JSON in the exact format shown below
2. If any field is not found, use "N/A"
3. For name: Look for the actual person's name (usually at the top)
4. For location: Look for city, state, or geographic location (NOT technical terms)
5. For current job: Find the most recent job title and company
6. For LinkedIn: Look for linkedin.com/in/ URLs or LinkedIn profile mentions
7. For job history: Format as "Company - Title - Duration" separated by " | "
8. For total experience: Extract as number only (e.g., "5" for 5 years)
9. Count all distinct jobs/positions in the work history

RESUME TEXT:
{resume_text[:5000]}

JOB DESCRIPTION:
{jd_text[:2000]}

Return JSON in this EXACT format:
{{
    "name": "Full Name Here",
    "role_in_jd": "Job title from JD",
    "current_job_title": "Current position title",
    "current_organization": "Current company name",
    "location": "City, State or geographic location",
    "phone": "Phone number",
    "email": "Email address",
    "linkedin": "LinkedIn profile URL or N/A",
    "other_socials": "Other social media or N/A",
    "total_experience": "X years",
    "job_history": "Company1 - Title1 - Duration1 | Company2 - Title2 - Duration2",
    "total_jobs": "Number of jobs",
    "match_score": "Percentage of resume candidate match with JD (0-100%)",
    "summary": "Brief professional summary in 2-3 sentences"
}}"""

    response = get_llm_response(prompt)

    if response.startswith("ERROR:"):
        return {}

    try:
        json_match = re.search(r"\{.*\}", response, re.DOTALL)
        if json_match:
            return json.loads(json_match.group(0))
    except (json.JSONDecodeError, ValueError) as e:
        log_llm_failure(
            {
                "ts": datetime.now().isoformat(),
                "error": f"JSON parse: {e}",
                "response_preview": response[:300],
            }
        )
        st.error(f"JSON parsing error: {e}")

    return {}


# ========== BACKUP EXTRACTION ==========
_SKIP_WORDS = frozenset(
    [
        "resume",
        "cv",
        "email",
        "phone",
        "address",
        "linkedin",
        "objective",
        "summary",
        "experience",
    ]
)


def extract_name_backup(text):
    for line in text.split("\n")[:8]:
        line = line.strip()
        if not line or any(kw in line.lower() for kw in _SKIP_WORDS):
            continue
        if (
            2 <= len(line.split()) <= 4
            and re.match(r"^[A-Za-z\s\.]+$", line)
            and not line.isupper()
        ):
            return line
    return "N/A"


def extract_contact_backup(text):
    email_m = re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)
    email = email_m.group(0) if email_m else "N/A"

    phone = "N/A"
    for pat in [
        r"(?:\+91|91)[\s\-]?[6-9]\d{9}",
        r"[6-9]\d{9}",
        r"\d{3}[\s\-]?\d{3}[\s\-]?\d{4}",
    ]:
        m = re.search(pat, text)
        if m:
            phone = m.group(0)
            break

    li_m = re.search(r"linkedin\.com/in/([A-Za-z0-9\-]+)", text, re.IGNORECASE)
    linkedin = f"linkedin.com/in/{li_m.group(1)}" if li_m else "N/A"
    return email, phone, linkedin


def extract_experience_backup(text):
    for pat in [
        r"(?:total|overall)[\s\-:]*(?:experience|exp)[\s\-:]*(\d+(?:\.\d+)?)\s*(?:years?|yrs?)",
        r"(\d+(?:\.\d+)?)\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)",
    ]:
        m = re.search(pat, text, re.IGNORECASE)
        if m and 0 <= float(m.group(1)) <= 50:
            return f"{float(m.group(1))} years"
    return "N/A"


# ========== HELPERS ==========
def _count_jobs(job_history):
    if not job_history or job_history == "N/A":
        return 0
    return len([e for e in job_history.split(" | ") if e.strip()])


def _extract_years(exp_str):
    if not exp_str or exp_str == "N/A":
        return 0.0
    m = re.search(r"(\d+(?:\.\d+)?)", exp_str)
    return float(m.group(1)) if m else 0.0


def _avg_tenure(total_exp, total_jobs):
    try:
        years = _extract_years(total_exp)
        jobs = (
            int(total_jobs)
            if str(total_jobs).isdigit()
            else _count_jobs(str(total_jobs))
        )
        return f"{years / jobs:.1f} years" if years > 0 and jobs > 0 else "N/A"
    except Exception:
        return "N/A"


def _pick(ai_val, backup_val):
    return ai_val if ai_val and ai_val != "N/A" else backup_val


def _dedup_key(row):
    name = str(row.get("Name", "")).strip().lower()
    email = str(row.get("Email", "")).strip().lower()
    jd = str(row.get("Job Description", "")).strip().lower()
    return (name, email, jd)


# ========== RESUME PROCESSING ==========
def process_resume(resume_file, jd_text, jd_name, batch_ts):
    try:
        resume_text = extract_text(resume_file)
        if not resume_text or len(resume_text) < 50:
            return _error_record(
                resume_file.name, "Text extraction failed", jd_name, batch_ts
            )

        ai = extract_all_fields_with_ai(resume_text, jd_text)
        b_name = extract_name_backup(resume_text)
        b_email, b_phone, b_linkedin = extract_contact_backup(resume_text)
        b_exp = extract_experience_backup(resume_text)

        job_history = ai.get("job_history", "N/A")
        total_jobs = ai.get("total_jobs", _count_jobs(job_history))
        total_exp = _pick(ai.get("total_experience"), b_exp)

        return {
            "Name": _pick(ai.get("name"), b_name),
            "Resume Match Score": ai.get("match_score", "N/A"),
            "Email": _pick(ai.get("email"), b_email),
            "Phone": _pick(ai.get("phone"), b_phone),
            "LinkedIn": _pick(ai.get("linkedin"), b_linkedin),
            "Current Job Title": ai.get("current_job_title", "N/A"),
            "Current Organization": ai.get("current_organization", "N/A"),
            "Total Experience": total_exp,
            "Total Jobs": str(total_jobs)
            if str(total_jobs).isdigit()
            else str(_count_jobs(job_history)),
            "Average Tenure": _avg_tenure(total_exp, total_jobs),
            "Location": ai.get("location", "N/A"),
            "Role in JD": ai.get("role_in_jd", "N/A"),
            "Resume Summary": ai.get("summary", "N/A"),
            "Job History (org-title-jobDuration)": job_history,
            "Other Socials": ai.get("other_socials", "N/A"),
            "Processed At": batch_ts,
            "Job Description": jd_name,
        }

    except Exception as e:
        log_llm_failure(
            {
                "ts": datetime.now().isoformat(),
                "error": f"process_resume: {resume_file.name}: {e}",
            }
        )
        return _error_record(
            resume_file.name, f"Processing error: {e}", jd_name, batch_ts
        )


def _error_record(filename, error_msg, jd_name, batch_ts):
    record = {col: "N/A" for col in EXCEL_COLUMNS}
    record.update(
        {
            "Name": filename,
            "Resume Match Score": "Error",
            "Resume Summary": error_msg,
            "Processed At": batch_ts,
            "Job Description": jd_name,
        }
    )
    return record


# ========== DUPLICATE DETECTION ==========
def detect_duplicates(new_results):
    conflicts = []
    non_conflicts = []

    if not RESULTS_EXCEL.exists():
        return [], new_results

    try:
        existing_df = pd.read_excel(RESULTS_EXCEL, engine="openpyxl")
    except Exception:
        return [], new_results

    if existing_df.empty:
        return [], new_results

    existing_lookup = {}
    for idx, row in existing_df.iterrows():
        key = _dedup_key(row.to_dict())
        if key[0] == "n/a" and key[1] == "n/a":
            continue
        existing_lookup[key] = (row.to_dict(), idx)

    for new_row in new_results:
        key = _dedup_key(new_row)
        if (key[0] == "n/a" and key[1] == "n/a") or key not in existing_lookup:
            non_conflicts.append(new_row)
        else:
            old_row, old_idx = existing_lookup[key]
            conflicts.append(
                {
                    "key": key,
                    "old_row": old_row,
                    "new_row": new_row,
                    "old_excel_idx": old_idx,
                }
            )

    return conflicts, non_conflicts


# ========== EXCEL OPERATIONS ==========
def append_results_to_excel(results: list):
    if not results:
        return
    new_df = pd.DataFrame(results).reindex(columns=EXCEL_COLUMNS)

    if RESULTS_EXCEL.exists():
        try:
            existing_df = pd.read_excel(RESULTS_EXCEL, engine="openpyxl")
            combined_df = pd.concat([existing_df, new_df], ignore_index=True)
        except Exception:
            combined_df = new_df
    else:
        combined_df = new_df

    combined_df.to_excel(RESULTS_EXCEL, index=False, engine="openpyxl")
    return combined_df


def replace_rows_in_excel(replacements: list):
    if not replacements or not RESULTS_EXCEL.exists():
        return

    try:
        df = pd.read_excel(RESULTS_EXCEL, engine="openpyxl")
    except Exception:
        return

    for old_idx, new_row in replacements:
        if old_idx < len(df):
            for col in EXCEL_COLUMNS:
                if col in new_row:
                    df.at[old_idx, col] = new_row[col]

    df.to_excel(RESULTS_EXCEL, index=False, engine="openpyxl")


# ========== DISPLAY ==========
def display_results():
    results = st.session_state.results
    df = st.session_state.df
    total = len(results)

    st.success(f"✅ Processed {total} resumes — results saved to Excel.")

    st.subheader("📋 Results Preview")
    st.dataframe(df.head(5), width="stretch")

    successful = sum(1 for r in results if r["Resume Match Score"] != "Error")
    names_ok = sum(1 for r in results if r["Name"] != "N/A")
    tenure_ok = sum(1 for r in results if r["Average Tenure"] != "N/A")

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("📊 Total", total)
    c2.metric("✅ Successful", successful)
    c3.metric("👤 Names", names_ok)
    c4.metric("📈 Tenure", tenure_ok)

    tenure_values = []
    for r in results:
        if r["Average Tenure"] != "N/A":
            try:
                tenure_values.append(float(r["Average Tenure"].replace(" years", "")))
            except ValueError:
                pass

    if tenure_values:
        st.subheader("📊 Tenure Insights")
        c1, c2, c3 = st.columns(3)
        c1.metric("Average", f"{sum(tenure_values) / len(tenure_values):.1f} yrs")
        c2.metric("Highest", f"{max(tenure_values):.1f} yrs")
        c3.metric("Lowest", f"{min(tenure_values):.1f} yrs")

        for label, lo, hi in [
            ("0–2 yrs", 0, 2),
            ("2–5 yrs", 2, 5),
            ("5–10 yrs", 5, 10),
            ("10+ yrs", 10, 999),
        ]:
            count = sum(
                1 for t in tenure_values if (lo <= t <= hi if lo == 0 else lo < t <= hi)
            )
            st.write(
                f"• **{label}**: {count} ({count / len(tenure_values) * 100:.1f}%)"
            )

    with st.expander("📈 View All Results"):
        st.dataframe(df, width="stretch")

    st.subheader("🔍 Extraction Quality")
    for label, key in [
        ("Names", "Name"),
        ("Emails", "Email"),
        ("Phones", "Phone"),
        ("LinkedIn", "LinkedIn"),
        ("Tenure", "Average Tenure"),
        ("Job History", "Job History (org-title-jobDuration)"),
    ]:
        count = sum(1 for r in results if r[key] != "N/A")
        st.write(f"• **{label}**: {count}/{total} ({count / total * 100:.1f}%)")


def display_conflict_resolution():
    conflicts = st.session_state.pending_conflicts

    st.warning(
        f"⚠️ {len(conflicts)} duplicate(s) found — same candidate compared with the same JD before."
    )

    bulk_col1, bulk_col2, bulk_col3 = st.columns(3)
    with bulk_col1:
        if st.button("✅ Keep All Latest", type="primary"):
            _resolve_all_conflicts("new")
            return
    with bulk_col2:
        if st.button("⏪ Keep All Previous"):
            _resolve_all_conflicts("old")
            return
    with bulk_col3:
        if st.button("🔽 Choose Individually"):
            pass

    st.divider()

    for i, conflict in enumerate(conflicts):
        old = conflict["old_row"]
        new = conflict["new_row"]
        name = new.get("Name", "Unknown")
        jd = new.get("Job Description", "Unknown")

        st.subheader(f"Conflict {i + 1}: {name} vs {jd}")

        compare_fields = [
            "Resume Match Score",
            "Current Job Title",
            "Current Organization",
            "Total Experience",
            "Average Tenure",
            "Location",
            "Resume Summary",
        ]

        col_old, col_new = st.columns(2)
        with col_old:
            st.markdown(f"**Previous** (_{old.get('Processed At', '?')}_)")
            for field in compare_fields:
                st.write(f"**{field}:** {old.get(field, 'N/A')}")
        with col_new:
            st.markdown(f"**Latest** (_{new.get('Processed At', '?')}_)")
            for field in compare_fields:
                st.write(f"**{field}:** {new.get(field, 'N/A')}")

        choice = st.radio(
            f"Keep which for **{name}**?",
            options=["Latest", "Previous"],
            key=f"conflict_choice_{i}",
            horizontal=True,
        )
        conflict["user_choice"] = "new" if choice == "Latest" else "old"
        st.divider()

    if st.button("💾 Apply Selections & Save", type="primary"):
        _resolve_individual_conflicts()


def _resolve_all_conflicts(keep):
    conflicts = st.session_state.pending_conflicts
    non_conflicts = st.session_state.non_conflict_results

    if keep == "new":
        replace_rows_in_excel([(c["old_excel_idx"], c["new_row"]) for c in conflicts])

    if non_conflicts:
        append_results_to_excel(non_conflicts)

    _clear_conflict_state()
    st.rerun()


def _resolve_individual_conflicts():
    conflicts = st.session_state.pending_conflicts
    non_conflicts = st.session_state.non_conflict_results

    replacements = [
        (c["old_excel_idx"], c["new_row"])
        for c in conflicts
        if c.get("user_choice") == "new"
    ]
    if replacements:
        replace_rows_in_excel(replacements)

    if non_conflicts:
        append_results_to_excel(non_conflicts)

    _clear_conflict_state()
    st.rerun()


def _clear_conflict_state():
    st.session_state.pending_conflicts = None
    st.session_state.non_conflict_results = None
    st.session_state.conflict_resolved = True


# ========== MAIN UI ==========
def main():
    st.set_page_config(page_title="🎯 Resume Screener", layout="wide")
    init_session_state()

    st.title("🎯 Resume Screener")
    st.markdown("**AI-powered resume parsing with automatic Excel logging**")

    with st.sidebar:
        st.header("📊 Dashboard")
        st.metric("🤖 LLM Calls", st.session_state.llm_request_count)

        providers = ["Gemini"]
        if GROQ_KEY_PRIMARY:
            providers.extend(["Groq/compound", "Groq/llama"])
        if GROQ_KEY_ALT:
            providers.extend(["Groq-Alt/compound", "Groq-Alt/llama"])
        st.caption(f"**Fallback:** {' → '.join(providers)}")

        st.divider()
        st.subheader("📂 Files")
        st.code(str(LOG_DIR), language=None)
        st.caption("• `llm_calls.jsonl` — errors only")
        st.caption("• `resume_screening_results.xlsx`")

        if RESULTS_EXCEL.exists():
            try:
                row_count = len(pd.read_excel(RESULTS_EXCEL, engine="openpyxl"))
                st.metric("📊 Total Excel Rows", row_count)
            except Exception:
                pass

        st.divider()
        st.subheader("⚡ Pipeline Mode")
        use_two_stage = st.toggle(
            "Two-Stage Pipeline",
            value=False,
            help="Stage 1: Vector+BM25 filter → Stage 2: LLM score only top-K",
        )
        if use_two_stage:
            top_k = st.slider("Top-K for Stage 2", min_value=5, max_value=50, value=20)
            fusion_mode = st.radio(
                "Fusion Mode",
                ["Weighted Average", "RRF (Reciprocal Rank Fusion)"],
                index=0,
                help="Weighted: α-blended scores. RRF: rank-based, prevents score domination.",
            )
            fusion_mode_val = "rrf" if "RRF" in fusion_mode else "weighted"
            alpha = st.slider("Vector weight (α)", min_value=0.0, max_value=1.0, value=0.7, step=0.1,
                              help="α × vector + (1-α) × BM25 (only affects Weighted mode)",
                              disabled=(fusion_mode_val == "rrf"))
        else:
            top_k = 20
            alpha = 0.7
            fusion_mode_val = "weighted"

        st.divider()
        if st.button("🗑️ Clear Screen"):
            st.session_state.results = None
            st.session_state.df = None
            st.session_state.processing_done = False
            st.session_state.pending_conflicts = None
            st.session_state.non_conflict_results = None
            st.session_state.conflict_resolved = False
            st.session_state.pipeline_stage1_results = None
            st.rerun()

    if st.session_state.pending_conflicts:
        display_conflict_resolution()
        return

    col1, col2 = st.columns(2)
    with col1:
        jd_file = st.file_uploader("📋 Job Description", type=["pdf", "docx"])
    with col2:
        resume_files = st.file_uploader(
            "📄 Resumes", type=["pdf", "docx"], accept_multiple_files=True
        )

    if st.button("🚀 Start Processing", type="primary"):
        if not jd_file or not resume_files:
            st.error("⚠️ Upload both a JD and at least one resume")
        else:
            jd_text = extract_text(jd_file)
            if not jd_text:
                st.error("❌ Failed to extract JD text")
            elif use_two_stage:
                # ── Two-Stage Pipeline ────────────────────────────
                jd_name = jd_file.name
                batch_ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                log_batch_separator(batch_ts, jd_name, len(resume_files))

                with st.status("⚡ Two-Stage Pipeline...", expanded=True) as status:
                    # Extract all resume texts
                    st.write("📄 Extracting resume texts...")
                    resume_docs = []
                    for idx, rf in enumerate(resume_files, 1):
                        text = extract_text(rf)
                        if text and len(text) > 50:
                            resume_docs.append(
                                {"id": rf.name, "text": text, "metadata": {"filename": rf.name}}
                            )

                    st.write(f"🔍 Stage 1: Filtering {len(resume_docs)} resumes (top {top_k})...")
                    pipeline = ResumeScreenerPipeline(alpha=alpha, fusion_mode=fusion_mode_val)
                    pipeline_results = pipeline.run_with_files(
                        resume_docs=resume_docs,
                        jd_text=jd_text,
                        top_k=min(top_k, len(resume_docs)),
                    )
                    st.session_state.llm_request_count += pipeline.llm.request_count

                    status.update(label="✅ Pipeline complete!", state="complete", expanded=False)

                # Convert pipeline results to standard format
                results = []
                for pr in pipeline_results:
                    results.append({
                        "Name": pr.get("name", "N/A"),
                        "Resume Match Score": f"{pr.get('match_score', 'N/A')}%"
                            if isinstance(pr.get('match_score'), (int, float)) else "N/A",
                        "Email": pr.get("email", "N/A"),
                        "Phone": pr.get("phone", "N/A"),
                        "LinkedIn": pr.get("linkedin", "N/A"),
                        "Current Job Title": pr.get("current_job_title", "N/A"),
                        "Current Organization": pr.get("current_organization", "N/A"),
                        "Total Experience": pr.get("total_experience", "N/A"),
                        "Total Jobs": "N/A",
                        "Average Tenure": "N/A",
                        "Location": pr.get("location", "N/A"),
                        "Role in JD": jd_name,
                        "Resume Summary": pr.get("summary", "N/A"),
                        "Job History (org-title-jobDuration)": "N/A",
                        "Other Socials": "N/A",
                        "Processed At": batch_ts,
                        "Job Description": jd_name,
                    })

                # Show Stage 1 scores in a separate table
                st.subheader("⚡ Stage 1: Vector + BM25 Scores")
                stage1_df = pd.DataFrame(pipeline_results)[
                    ["filename", "name", "hybrid_score", "vector_score", "bm25_score", "match_score", "fit_report"]
                ].rename(columns={
                    "filename": "File", "name": "Name", "hybrid_score": "Hybrid Score",
                    "vector_score": "Vector Score", "bm25_score": "BM25 Score",
                    "match_score": "LLM Score", "fit_report": "Fit Report",
                })
                st.dataframe(stage1_df, use_container_width=True)

                st.info(
                    f"📊 Scanned **{len(resume_docs)}** resumes → "
                    f"shortlisted **{len(pipeline_results)}** → "
                    f"**{pipeline.llm.request_count}** LLM calls "
                    f"(saved {len(resume_docs) - len(pipeline_results)} API calls)"
                )

                st.session_state.results = results
                st.session_state.df = pd.DataFrame(results).reindex(columns=EXCEL_COLUMNS)
                st.session_state.processing_done = True

                conflicts, non_conflicts = detect_duplicates(results)
                if conflicts:
                    st.session_state.pending_conflicts = conflicts
                    st.session_state.non_conflict_results = non_conflicts
                    st.session_state.conflict_resolved = False
                    st.rerun()
                else:
                    append_results_to_excel(results)
            else:
                # ── Original single-resume processing ─────────────
                jd_name = jd_file.name
                batch_ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                log_batch_separator(batch_ts, jd_name, len(resume_files))

                results = []
                with st.status("🚀 Processing...", expanded=True) as status:
                    progress = st.progress(0)
                    file_display = st.empty()

                    for idx, rf in enumerate(resume_files, 1):
                        file_display.markdown(
                            f"**🔍 [{idx}/{len(resume_files)}]** `{rf.name}`"
                        )
                        results.append(process_resume(rf, jd_text, jd_name, batch_ts))
                        progress.progress(idx / len(resume_files))

                    file_display.empty()
                    status.update(
                        label="✅ Complete!", state="complete", expanded=False
                    )

                st.session_state.results = results
                st.session_state.df = pd.DataFrame(results).reindex(
                    columns=EXCEL_COLUMNS
                )
                st.session_state.processing_done = True

                conflicts, non_conflicts = detect_duplicates(results)

                if conflicts:
                    st.session_state.pending_conflicts = conflicts
                    st.session_state.non_conflict_results = non_conflicts
                    st.session_state.conflict_resolved = False
                    st.rerun()
                else:
                    append_results_to_excel(results)

    if st.session_state.processing_done and st.session_state.results:
        display_results()


if __name__ == "__main__":
    main()
