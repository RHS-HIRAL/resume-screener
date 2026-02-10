import streamlit as st
import google.generativeai as genai
import os
import PyPDF2 as pdf
from dotenv import load_dotenv
import json
import re
import tempfile
from docx import Document
import pandas as pd
from datetime import datetime
import logging
import io

# Load environment variables
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))


# ========== LOGGING SETUP ==========
def setup_logger():
    """Setup a logger that writes to both file and a StringIO buffer for UI display"""
    logger = logging.getLogger("ResumeScreener")
    logger.setLevel(logging.DEBUG)

    # Avoid duplicate handlers on Streamlit reruns
    if not logger.handlers:
        # File handler — persistent log file
        log_dir = tempfile.gettempdir()
        log_file = os.path.join(log_dir, "resume_screener.log")
        file_handler = logging.FileHandler(log_file, mode="a")
        file_handler.setLevel(logging.DEBUG)
        file_fmt = logging.Formatter(
            "%(asctime)s | %(levelname)-8s | %(message)s", datefmt="%Y-%m-%d %H:%M:%S"
        )
        file_handler.setFormatter(file_fmt)
        logger.addHandler(file_handler)

    return logger


logger = setup_logger()


def log_and_track(message, level="info"):
    """Log a message and also store it in session_state for UI display"""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    entry = f"[{timestamp}] [{level.upper()}] {message}"

    # Log to file
    getattr(logger, level, logger.info)(message)

    # Store in session_state for UI
    if "log_entries" not in st.session_state:
        st.session_state.log_entries = []
    st.session_state.log_entries.append(entry)


# ========== SESSION STATE INITIALIZATION ==========
def init_session_state():
    """Initialize all session state keys once"""
    defaults = {
        "results": None,  # Processed results (list of dicts)
        "df": None,  # DataFrame of results
        "processing_done": False,  # Flag: processing completed
        "history": [],  # History of all past runs
        "log_entries": [],  # Log entries for current session
        "llm_request_count": 0,  # Total LLM API calls made
        "llm_call_log": [],  # Detailed log of every LLM call
    }
    for key, default in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default


# ========== TEXT EXTRACTION ==========
def extract_text_from_pdf(uploaded_file):
    reader = pdf.PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        text += str(page.extract_text())
    return text


def extract_text_from_docx(uploaded_file):
    doc = Document(uploaded_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
    return text


def extract_text(uploaded_file):
    if uploaded_file.name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    elif uploaded_file.name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    return ""


# ========== AI-POWERED EXTRACTION ==========
def get_gemini_response(prompt):
    """Call Gemini API with full logging and request tracking"""
    st.session_state.llm_request_count += 1
    request_num = st.session_state.llm_request_count
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    log_and_track(f"LLM Request #{request_num} — Sending prompt ({len(prompt)} chars)")

    call_record = {
        "request_num": request_num,
        "timestamp": timestamp,
        "prompt_length": len(prompt),
        "prompt_preview": prompt[:300] + "..." if len(prompt) > 300 else prompt,
        "status": "pending",
        "response_length": 0,
        "response_preview": "",
        "error": None,
    }

    try:
        model = genai.GenerativeModel("gemini-2.5-flash")
        response = model.generate_content(prompt)
        response_text = response.text

        call_record["status"] = "success"
        call_record["response_length"] = len(response_text)
        call_record["response_preview"] = (
            response_text[:500] + "..." if len(response_text) > 500 else response_text
        )

        log_and_track(
            f"LLM Request #{request_num} — Success ({len(response_text)} chars response)"
        )

        st.session_state.llm_call_log.append(call_record)
        return response_text

    except Exception as e:
        call_record["status"] = "error"
        call_record["error"] = str(e)
        st.session_state.llm_call_log.append(call_record)

        log_and_track(f"LLM Request #{request_num} — ERROR: {str(e)}", level="error")
        return f"ERROR: {str(e)}"


def extract_all_fields_with_ai(resume_text, jd_text):
    """Extract all fields using a single, comprehensive AI call"""
    log_and_track(
        f"Starting AI field extraction (resume: {len(resume_text)} chars, JD: {len(jd_text)} chars)"
    )

    prompt = f"""
You are an expert resume parser. Extract the following information from this resume with 100% accuracy.

CRITICAL INSTRUCTIONS:
1. Return ONLY valid JSON in the exact format shown below
2. If any field is not found, use "N/A"
3. For name: Look for the actual person's name (usually at the top)
4. For location: Look for city, state, or geographic location (NOT technical terms)
5. For current job: Find the most recent job title and company
6. For LinkedIn: Look for linkedin.com/in/ URLs or LinkedIn profile mentions
7. For job history: Format as "Company - Title - Duration" separated by " | "
8. For total experience: Extract as number only (e.g., "5" for 5 years)
9. Count all distinct jobs/positions in the work history

RESUME TEXT:
{resume_text[:3000]}

JOB DESCRIPTION:
{jd_text[:1000]}

Return JSON in this EXACT format:
{{
    "name": "Full Name Here",
    "role_in_jd": "Job title from JD",
    "current_job_title": "Current position title",
    "current_organization": "Current company name",
    "location": "City, State or geographic location",
    "phone": "Phone number",
    "email": "Email address",
    "linkedin": "LinkedIn profile URL or N/A",
    "other_socials": "Other social media or N/A",
    "total_experience": "X years",
    "job_history": "Company1 - Title1 - Duration1 | Company2 - Title2 - Duration2",
    "total_jobs": "Number of jobs",
    "match_score": "XX%",
    "summary": "Brief professional summary in 2-3 sentences"
}}
"""
    response = get_gemini_response(prompt)

    try:
        json_match = re.search(r"\{.*\}", response, re.DOTALL)
        if json_match:
            data = json.loads(json_match.group(0))
            log_and_track(f"AI extraction successful — extracted {len(data)} fields")
            return data
    except Exception as e:
        log_and_track(f"JSON parsing error: {str(e)}", level="error")
        st.error(f"JSON parsing error: {str(e)}")

    return {}


# ========== BACKUP EXTRACTION FUNCTIONS ==========
def extract_name_backup(text):
    """Backup name extraction if AI fails"""
    lines = text.split("\n")[:8]
    for line in lines:
        line = line.strip()
        if any(
            keyword in line.lower()
            for keyword in [
                "resume",
                "cv",
                "email",
                "phone",
                "address",
                "linkedin",
                "objective",
                "summary",
                "experience",
            ]
        ):
            continue
        if line and 2 <= len(line.split()) <= 4:
            if re.match(r"^[A-Za-z\s\.]+$", line) and not line.isupper():
                return line
    return "N/A"


def extract_contact_backup(text):
    """Backup contact extraction"""
    email_match = re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)
    email = email_match.group(0) if email_match else "N/A"
    phone_patterns = [
        r"(?:\+91|91)[\s\-]?[6-9]\d{9}",
        r"[6-9]\d{9}",
        r"\d{3}[\s\-]?\d{3}[\s\-]?\d{4}",
    ]
    phone = "N/A"
    for pattern in phone_patterns:
        match = re.search(pattern, text)
        if match:
            phone = match.group(0)
            break
    linkedin_match = re.search(
        r"linkedin\.com/in/([A-Za-z0-9\-]+)", text, re.IGNORECASE
    )
    linkedin = f"linkedin.com/in/{linkedin_match.group(1)}" if linkedin_match else "N/A"
    return email, phone, linkedin


def extract_experience_backup(text):
    """Backup experience extraction"""
    patterns = [
        r"(?:total|overall)[\s\-:]*(?:experience|exp)[\s\-:]*(\d+(?:\.\d+)?)\s*(?:years?|yrs?)",
        r"(\d+(?:\.\d+)?)\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            exp = float(match.group(1))
            if 0 <= exp <= 50:
                return f"{exp} years"
    return "N/A"


def count_jobs_from_history(job_history):
    """Count number of jobs from job history string"""
    if not job_history or job_history == "N/A":
        return 0
    job_entries = job_history.split(" | ")
    return len([entry for entry in job_entries if entry.strip()])


def extract_numeric_experience(experience_str):
    """Extract numeric value from experience string"""
    if not experience_str or experience_str == "N/A":
        return 0
    match = re.search(r"(\d+(?:\.\d+)?)", experience_str)
    if match:
        return float(match.group(1))
    return 0


def calculate_average_tenure(total_experience, total_jobs):
    """Calculate average tenure per job"""
    try:
        exp_years = extract_numeric_experience(total_experience)
        jobs_count = (
            int(total_jobs)
            if str(total_jobs).isdigit()
            else count_jobs_from_history(str(total_jobs))
        )
        if exp_years > 0 and jobs_count > 0:
            avg_tenure = exp_years / jobs_count
            return f"{avg_tenure:.1f} years"
        else:
            return "N/A"
    except:
        return "N/A"


# ========== MAIN PROCESSING ==========
def process_resume_enhanced(resume_file, jd_text, index):
    """Enhanced resume processing with AI + backup extraction"""
    log_and_track(f"Processing resume #{index}: {resume_file.name}")

    try:
        resume_text = extract_text(resume_file)
        if not resume_text or len(resume_text) < 50:
            log_and_track(
                f"Resume #{index}: Text extraction failed (got {len(resume_text) if resume_text else 0} chars)",
                level="warning",
            )
            return create_error_record(index, "Text extraction failed")

        log_and_track(f"Resume #{index}: Extracted {len(resume_text)} chars of text")

        ai_data = extract_all_fields_with_ai(resume_text, jd_text)

        # Backup extractions
        backup_name = extract_name_backup(resume_text)
        backup_email, backup_phone, backup_linkedin = extract_contact_backup(
            resume_text
        )
        backup_experience = extract_experience_backup(resume_text)

        job_history = ai_data.get("job_history", "N/A")
        total_jobs = ai_data.get("total_jobs", count_jobs_from_history(job_history))
        total_experience = (
            ai_data.get("total_experience", backup_experience)
            if ai_data.get("total_experience")
            and ai_data.get("total_experience") != "N/A"
            else backup_experience
        )
        avg_tenure = calculate_average_tenure(total_experience, total_jobs)

        result = {
            "Sr.no": index,
            "Name": ai_data.get("name", backup_name)
            if ai_data.get("name") and ai_data.get("name") != "N/A"
            else backup_name,
            "Role in JD": ai_data.get("role_in_jd", "N/A"),
            "Current Job Title": ai_data.get("current_job_title", "N/A"),
            "Current Organization": ai_data.get("current_organization", "N/A"),
            "Location": ai_data.get("location", "N/A"),
            "Resume Match Score": ai_data.get("match_score", "N/A"),
            "Resume Summary": ai_data.get("summary", "N/A"),
            "Phone": ai_data.get("phone", backup_phone)
            if ai_data.get("phone") and ai_data.get("phone") != "N/A"
            else backup_phone,
            "Email": ai_data.get("email", backup_email)
            if ai_data.get("email") and ai_data.get("email") != "N/A"
            else backup_email,
            "LinkedIn": ai_data.get("linkedin", backup_linkedin)
            if ai_data.get("linkedin") and ai_data.get("linkedin") != "N/A"
            else backup_linkedin,
            "Other Socials": ai_data.get("other_socials", "N/A"),
            "Total Experience": total_experience,
            "Total Jobs": str(total_jobs)
            if str(total_jobs).isdigit()
            else str(count_jobs_from_history(job_history)),
            "Average Tenure": avg_tenure,
            "Job History (org-title-jobDuration)": job_history,
        }

        log_and_track(
            f"Resume #{index}: Completed — Name={result['Name']}, Score={result['Resume Match Score']}"
        )
        return result

    except Exception as e:
        log_and_track(f"Resume #{index}: Processing error — {str(e)}", level="error")
        return create_error_record(index, f"Processing error: {str(e)}")


def create_error_record(index, error_msg):
    """Create error record"""
    return {
        "Sr.no": index,
        "Name": "N/A",
        "Role in JD": "N/A",
        "Current Job Title": "N/A",
        "Current Organization": "N/A",
        "Location": "N/A",
        "Resume Match Score": "Error",
        "Resume Summary": error_msg,
        "Phone": "N/A",
        "Email": "N/A",
        "LinkedIn": "N/A",
        "Other Socials": "N/A",
        "Total Experience": "N/A",
        "Total Jobs": "N/A",
        "Average Tenure": "N/A",
        "Job History (org-title-jobDuration)": "N/A",
    }


# ========== RESULTS DISPLAY (separated from processing) ==========
def display_results():
    """Display results from session_state — survives reruns"""
    results = st.session_state.results
    df = st.session_state.df

    st.success(f"✅ Successfully processed {len(results)} resumes!")

    # Show sample results for verification
    st.subheader("📋 Sample Results (First 3 rows)")
    st.dataframe(df.head(3), width="stretch")

    # Statistics
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("📊 Total Resumes", len(results))
    with col2:
        successful = sum(1 for r in results if r["Resume Match Score"] != "Error")
        st.metric("✅ Successful", successful)
    with col3:
        names_extracted = sum(1 for r in results if r["Name"] != "N/A")
        st.metric("👤 Names Extracted", names_extracted)
    with col4:
        tenure_calculated = sum(1 for r in results if r["Average Tenure"] != "N/A")
        st.metric("📈 Tenure Calculated", tenure_calculated)

    # Show average tenure insights
    st.subheader("📊 Average Tenure Insights")
    valid_tenures = [
        r["Average Tenure"] for r in results if r["Average Tenure"] != "N/A"
    ]
    tenure_values = []
    for tenure in valid_tenures:
        try:
            value = float(tenure.replace(" years", ""))
            tenure_values.append(value)
        except:
            pass

    if tenure_values:
        avg_tenure_overall = sum(tenure_values) / len(tenure_values)
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("📊 Average Tenure Overall", f"{avg_tenure_overall:.1f} years")
        with col2:
            st.metric("📈 Highest Tenure", f"{max(tenure_values):.1f} years")
        with col3:
            st.metric("📉 Lowest Tenure", f"{min(tenure_values):.1f} years")

        st.subheader("📊 Tenure Distribution")
        tenure_ranges = {
            "0-2 years": sum(1 for t in tenure_values if 0 <= t <= 2),
            "2-5 years": sum(1 for t in tenure_values if 2 < t <= 5),
            "5-10 years": sum(1 for t in tenure_values if 5 < t <= 10),
            "10+ years": sum(1 for t in tenure_values if t > 10),
        }
        for range_name, count in tenure_ranges.items():
            percentage = (count / len(tenure_values)) * 100
            st.write(f"• **{range_name}**: {count} candidates ({percentage:.1f}%)")

    # Full results
    with st.expander("📈 View All Results"):
        st.dataframe(df, width="stretch")

    # ---------- DOWNLOAD (no rerun / no clearing) ----------
    # Pre-generate the Excel bytes so the download_button uses cached data
    # and does NOT trigger reprocessing
    excel_buffer = io.BytesIO()
    df.to_excel(excel_buffer, index=False, engine="openpyxl")
    excel_bytes = excel_buffer.getvalue()

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    st.download_button(
        label="📥 Download Results (Excel)",
        data=excel_bytes,
        file_name=f"resume_results_{timestamp}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    # CSV download option too
    csv_bytes = df.to_csv(index=False).encode("utf-8")
    st.download_button(
        label="📥 Download Results (CSV)",
        data=csv_bytes,
        file_name=f"resume_results_{timestamp}.csv",
        mime="text/csv",
    )

    # Quality check
    st.subheader("🔍 Extraction Quality Check")
    quality_metrics = {
        "Names Extracted": f"{names_extracted}/{len(results)} ({names_extracted / len(results) * 100:.1f}%)",
        "Emails Extracted": f"{sum(1 for r in results if r['Email'] != 'N/A')}/{len(results)}",
        "Phones Extracted": f"{sum(1 for r in results if r['Phone'] != 'N/A')}/{len(results)}",
        "LinkedIn Extracted": f"{sum(1 for r in results if r['LinkedIn'] != 'N/A')}/{len(results)}",
        "Average Tenure Calculated": f"{tenure_calculated}/{len(results)} ({tenure_calculated / len(results) * 100:.1f}%)",
        "Job History Extracted": f"{sum(1 for r in results if r['Job History (org-title-jobDuration)'] != 'N/A')}/{len(results)}",
    }
    for metric, value in quality_metrics.items():
        st.write(f"• **{metric}**: {value}")

    # LLM usage stats
    st.subheader("🤖 LLM Usage Stats")
    st.metric("Total LLM API Calls", st.session_state.llm_request_count)


def display_history():
    """Display history of past processing runs"""
    if not st.session_state.history:
        st.info("No processing history yet. Run your first batch above!")
        return

    for i, run in enumerate(reversed(st.session_state.history), 1):
        with st.expander(
            f"Run {len(st.session_state.history) - i + 1} — {run['timestamp']} — {run['total_resumes']} resumes"
        ):
            st.write(f"**Resumes processed:** {run['total_resumes']}")
            st.write(f"**Successful:** {run['successful']}")
            st.write(f"**LLM calls made:** {run['llm_calls']}")
            st.dataframe(pd.DataFrame(run["results"]).head(5), width="stretch")


def display_logs():
    """Display session logs and LLM call details"""
    tab1, tab2 = st.tabs(["📝 Activity Log", "🤖 LLM Call Details"])

    with tab1:
        if st.session_state.log_entries:
            # Show most recent first
            for entry in reversed(st.session_state.log_entries[-100:]):
                st.text(entry)
        else:
            st.info("No log entries yet.")

    with tab2:
        if st.session_state.llm_call_log:
            for call in reversed(st.session_state.llm_call_log):
                status_icon = "✅" if call["status"] == "success" else "❌"
                with st.expander(
                    f"{status_icon} Request #{call['request_num']} — {call['timestamp']}"
                ):
                    st.write(f"**Status:** {call['status']}")
                    st.write(f"**Prompt length:** {call['prompt_length']} chars")
                    st.write(f"**Response length:** {call['response_length']} chars")
                    if call["error"]:
                        st.error(f"Error: {call['error']}")
                    st.text_area(
                        "Prompt preview",
                        call["prompt_preview"],
                        height=100,
                        key=f"prompt_{call['request_num']}",
                    )
                    st.text_area(
                        "Response preview",
                        call["response_preview"],
                        height=100,
                        key=f"resp_{call['request_num']}",
                    )
        else:
            st.info("No LLM calls made yet.")


# ========== STREAMLIT UI ==========
def main():
    st.set_page_config(page_title="🎯 Resume Screener", layout="wide")
    init_session_state()

    st.title("🎯 Resume Screener")
    st.markdown(
        "**Fine-tuned AI extraction for precise results with Average Tenure Calculation**"
    )

    # Sidebar — LLM stats + history + logs
    with st.sidebar:
        st.header("📊 Session Dashboard")
        st.metric("🤖 Total LLM Calls", st.session_state.llm_request_count)
        st.metric("📁 Runs Completed", len(st.session_state.history))

        if st.button("🗑️ Clear Current Results"):
            st.session_state.results = None
            st.session_state.df = None
            st.session_state.processing_done = False
            st.rerun()

    # File upload
    col1, col2 = st.columns(2)
    with col1:
        jd_file = st.file_uploader("📋 Upload Job Description", type=["pdf", "docx"])
    with col2:
        resume_files = st.file_uploader(
            "📄 Upload Resume Files", type=["pdf", "docx"], accept_multiple_files=True
        )

    # Processing trigger
    if st.button("🚀 Start Accurate Processing", type="primary"):
        if not jd_file or not resume_files:
            st.error("⚠️ Please upload both JD and resume files")
        else:
            jd_text = extract_text(jd_file)
            if not jd_text:
                st.error("❌ Failed to extract JD text")
            else:
                log_and_track(
                    f"=== NEW RUN: {len(resume_files)} resumes, JD: {jd_file.name} ==="
                )
                llm_calls_before = st.session_state.llm_request_count
                results = []

                with st.status(
                    "🚀 Orchestrating AI Extraction...", expanded=True
                ) as status:
                    progress_bar = st.progress(0)
                    current_file_display = st.empty()

                    for idx, resume_file in enumerate(resume_files, 1):
                        current_file_display.markdown(
                            f"**🔍 Analyzing File {idx} of {len(resume_files)}:** `{resume_file.name}`"
                        )
                        result = process_resume_enhanced(resume_file, jd_text, idx)
                        results.append(result)
                        progress_bar.progress(idx / len(resume_files))

                    current_file_display.empty()
                    status.update(
                        label="✅ AI Analysis Complete!",
                        state="complete",
                        expanded=False,
                    )

                # Persist results in session_state
                st.session_state.results = results
                st.session_state.df = pd.DataFrame(results)
                st.session_state.processing_done = True

                # Save to history
                llm_calls_this_run = (
                    st.session_state.llm_request_count - llm_calls_before
                )
                history_entry = {
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "total_resumes": len(results),
                    "successful": sum(
                        1 for r in results if r["Resume Match Score"] != "Error"
                    ),
                    "llm_calls": llm_calls_this_run,
                    "results": results,
                }
                st.session_state.history.append(history_entry)
                log_and_track(
                    f"=== RUN COMPLETE: {len(results)} resumes, {llm_calls_this_run} LLM calls ==="
                )

    # ---- Always display results if they exist (survives reruns) ----
    if st.session_state.processing_done and st.session_state.results:
        display_results()

    # ---- Tabs for history and logs ----
    st.divider()
    tab_history, tab_logs = st.tabs(["📜 Processing History", "📋 Logs & LLM Calls"])

    with tab_history:
        display_history()

    with tab_logs:
        display_logs()


if __name__ == "__main__":
    main()
