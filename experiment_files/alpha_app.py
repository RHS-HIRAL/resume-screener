"""
Resume Screener — Streamlit UI (Batch Mode)
Select a folder, auto-match JD, and score ALL resumes in one go.
Supports force-rescore with conflict resolution (Previous / New / Average).
"""

import io
import json
import os
import re
import time
from datetime import datetime
from pathlib import Path
from urllib.parse import quote

import requests
import streamlit as st
from dotenv import load_dotenv

load_dotenv()

# ── SharePoint folder paths ──────────────────────────────────────────────────
SP_ORIGINAL_RESUMES_FOLDER = os.getenv("SHAREPOINT_BASE_FOLDER", "Resumes")
SP_TEXT_RESUMES_FOLDER = os.getenv(
    "SHAREPOINT_TEXT_RESUMES_FOLDER", "Text Files/NewResumes"
)
SP_TEXT_JD_FOLDER = os.getenv("SHAREPOINT_TEXT_JD_FOLDER", "Text Files/JobDescriptions")

# ── Optional SharePoint / MSAL integration ────────────────────────────────────
try:
    import msal

    HAS_MSAL = True
except ImportError:
    HAS_MSAL = False

# ── Optional text-extraction libraries ───────────────────────────────────────
try:
    import PyPDF2 as pdf_lib

    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document as DocxDocument

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ═══════════════════════════════════════════════════════════════════════════
# PAGE CONFIG & STYLES
# ═══════════════════════════════════════════════════════════════════════════

st.set_page_config(
    page_title="Resume Screener",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display:ital@0;1&family=DM+Sans:wght@300;400;500;600&family=JetBrains+Mono:wght@400;600&display=swap');

html, body, [class*="css"] { font-family: 'DM Sans', sans-serif; }

.stApp {
    background: #0d0f14;
    background-image:
        radial-gradient(ellipse at 20% 10%, rgba(99,102,241,0.12) 0%, transparent 60%),
        radial-gradient(ellipse at 80% 80%, rgba(16,185,129,0.08) 0%, transparent 50%);
}

[data-testid="stSidebar"] { background: #13151c !important; border-right: 1px solid rgba(255,255,255,0.06) !important; }
[data-testid="stSidebar"] * { color: #c8ccd8 !important; }

.rs-card {
    background: #161a24; border: 1px solid rgba(255,255,255,0.07); border-radius: 14px;
    padding: 1.6rem 1.8rem; margin-bottom: 1.2rem; box-shadow: 0 4px 30px rgba(0,0,0,0.3); transition: border-color 0.2s;
}
.rs-card:hover { border-color: rgba(99,102,241,0.3); }

.rs-title { font-family: 'DM Serif Display', serif; font-size: 2.6rem; color: #f0f2fa; letter-spacing: -0.02em; line-height: 1.1; margin: 0 0 0.3rem 0; }
.rs-subtitle { font-size: 0.95rem; color: #6b7280; letter-spacing: 0.02em; text-transform: uppercase; font-weight: 500; margin: 0 0 2rem 0; }
.rs-section-header { font-family: 'DM Serif Display', serif; font-size: 1.35rem; color: #e8eaf2; margin: 0 0 0.9rem 0; display: flex; align-items: center; gap: 0.5rem; }

.score-number { font-family: 'DM Serif Display', serif; font-size: 4.5rem; line-height: 1; font-weight: 400; background: linear-gradient(135deg, #818cf8, #34d399); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; }
.score-label { font-size: 0.75rem; color: #6b7280; text-transform: uppercase; letter-spacing: 0.1em; font-weight: 600; margin-top: 0.3rem; }

.badge { display: inline-block; padding: 0.22rem 0.75rem; border-radius: 999px; font-size: 0.72rem; font-weight: 700; letter-spacing: 0.06em; text-transform: uppercase; }
.badge-match       { background: rgba(16,185,129,0.15); color: #34d399; border: 1px solid rgba(16,185,129,0.3); }
.badge-partial     { background: rgba(251,191,36,0.13); color: #fbbf24; border: 1px solid rgba(251,191,36,0.3); }
.badge-no-match    { background: rgba(239,68,68,0.12);  color: #f87171; border: 1px solid rgba(239,68,68,0.3); }
.badge-skipped     { background: rgba(107,114,128,0.15); color: #9ca3af; border: 1px solid rgba(107,114,128,0.3); }

.param-card { background: #1a1f2e; border: 1px solid rgba(255,255,255,0.06); border-radius: 10px; padding: 1rem 1.2rem; margin-bottom: 0.7rem; }
.param-title { font-size: 0.78rem; color: #8b90a0; text-transform: uppercase; letter-spacing: 0.08em; font-weight: 600; margin-bottom: 0.4rem; }
.param-summary { font-size: 0.9rem; color: #c8ccd8; line-height: 1.5; }

.profile-field { display: flex; justify-content: space-between; align-items: baseline; padding: 0.55rem 0; border-bottom: 1px solid rgba(255,255,255,0.05); }
.profile-field:last-child { border-bottom: none; }
.pf-label { font-size: 0.75rem; color: #6b7280; text-transform: uppercase; letter-spacing: 0.08em; font-weight: 600; }
.pf-value { font-size: 0.9rem; color: #dde1ee; text-align: right; max-width: 65%; word-break: break-word; }

.info-pill { display: inline-flex; align-items: center; gap: 0.4rem; background: rgba(99,102,241,0.1); border: 1px solid rgba(99,102,241,0.2); border-radius: 999px; padding: 0.3rem 0.9rem; font-size: 0.8rem; color: #818cf8; font-weight: 500; margin: 0.2rem 0.2rem 0.2rem 0; }

.json-viewer { background: #0f1117; border: 1px solid rgba(255,255,255,0.08); border-radius: 10px; padding: 1.2rem; font-family: 'JetBrains Mono', monospace; font-size: 0.78rem; color: #a8b2c8; overflow-x: auto; max-height: 400px; overflow-y: auto; white-space: pre; }

/* Batch results table */
.batch-row { display: flex; align-items: center; gap: 1rem; padding: 0.7rem 1rem; border-bottom: 1px solid rgba(255,255,255,0.05); }
.batch-row:hover { background: rgba(99,102,241,0.05); }
.batch-name { flex: 1; font-size: 0.88rem; color: #dde1ee; }
.batch-score { font-family: 'JetBrains Mono', monospace; font-size: 1rem; font-weight: 700; min-width: 3rem; text-align: center; }
.batch-status { font-size: 0.75rem; min-width: 6rem; text-align: center; }

.stButton>button { background: linear-gradient(135deg, #6366f1, #4f46e5) !important; color: white !important; border: none !important; border-radius: 8px !important; padding: 0.6rem 1.8rem !important; font-family: 'DM Sans', sans-serif !important; font-weight: 600 !important; font-size: 0.9rem !important; letter-spacing: 0.02em !important; transition: all 0.2s !important; box-shadow: 0 4px 15px rgba(99,102,241,0.25) !important; }
.stButton>button:hover { transform: translateY(-1px) !important; box-shadow: 0 6px 20px rgba(99,102,241,0.4) !important; }
.stTextInput>div>div>input, .stTextArea>div>div>textarea, .stSelectbox>div>div>div { background: #1a1f2e !important; color: #dde1ee !important; border: 1px solid rgba(255,255,255,0.1) !important; border-radius: 8px !important; font-family: 'DM Sans', sans-serif !important; }
label { color: #8b90a0 !important; font-size: 0.8rem !important; font-weight: 600 !important; text-transform: uppercase !important; letter-spacing: 0.08em !important; }
.stMarkdown p { color: #c8ccd8; }
.stProgress > div > div { background: linear-gradient(90deg, #6366f1, #34d399) !important; }
hr { border-color: rgba(255,255,255,0.07) !important; }
.rs-divider { height: 1px; background: linear-gradient(90deg, transparent, rgba(99,102,241,0.4), transparent); margin: 1.5rem 0; }
.stAlert { border-radius: 10px !important; }
</style>
""",
    unsafe_allow_html=True,
)


# ═══════════════════════════════════════════════════════════════════════════
# HELPERS — Text Extraction & API
# ═══════════════════════════════════════════════════════════════════════════


def extract_text_from_pdf(file_bytes: bytes) -> str:
    if not HAS_PYPDF:
        return ""
    try:
        reader = pdf_lib.PdfReader(io.BytesIO(file_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        st.error(f"PDF extraction failed: {e}")
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    if not HAS_DOCX:
        return ""
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    except Exception as e:
        st.error(f"DOCX extraction failed: {e}")
        return ""


def call_api(api_url: str, resume_text: str, jd_text: str) -> dict | None:
    """Call the FastAPI backend to analyze a single resume against a JD."""
    endpoint = api_url.rstrip("/") + "/api/v1/analyze-resume"
    payload = {"resume_text": resume_text, "jd_text": jd_text}
    try:
        resp = requests.post(endpoint, json=payload, timeout=180)
        resp.raise_for_status()
        return resp.json()
    except Exception as e:
        return {"error": str(e)}


def save_result(result: dict, role: str, full_name: str) -> Path:
    safe_role = re.sub(r"[^\w\-]", "_", role.strip()) or "General"
    safe_name = re.sub(r"[^\w\-]", "_", full_name.strip()) or "Unknown"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir = Path("extracted_json_output") / safe_role
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{safe_name}_{ts}.json"
    out_path.write_text(
        json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    return out_path


# ═══════════════════════════════════════════════════════════════════════════
# SHAREPOINT — Client
# ═══════════════════════════════════════════════════════════════════════════


class SharePointClient:
    """
    Handles all SharePoint interactions: listing files, downloading text,
    checking/pushing MatchScore on original resume files.
    """

    GRAPH_BASE = "https://graph.microsoft.com/v1.0"
    SCOPES = ["https://graph.microsoft.com/.default"]

    def __init__(
        self,
        tenant_id: str,
        client_id: str,
        client_secret: str,
        site_domain: str,
        site_path: str,
        drive_name: str,
    ):
        self._msal_app = msal.ConfidentialClientApplication(
            client_id=client_id,
            client_credential=client_secret,
            authority=f"https://login.microsoftonline.com/{tenant_id}",
        )
        self.site_domain = site_domain
        self.site_path = site_path.strip("/")
        self.drive_name = drive_name
        self._site_id: str | None = None
        self._drive_id: str | None = None

    # ── Auth ───────────────────────────────────────────────────────────────

    def _headers(self) -> dict:
        result = self._msal_app.acquire_token_silent(self.SCOPES, account=None)
        if not result:
            result = self._msal_app.acquire_token_for_client(scopes=self.SCOPES)
        return {
            "Authorization": f"Bearer {result['access_token']}",
            "Content-Type": "application/json",
        }

    def _get_drive_id(self) -> str:
        if self._drive_id:
            return self._drive_id
        url_site = f"{self.GRAPH_BASE}/sites/{self.site_domain}:/{self.site_path}"
        self._site_id = requests.get(
            url_site, headers=self._headers(), timeout=30
        ).json()["id"]

        url_drives = f"{self.GRAPH_BASE}/sites/{self._site_id}/drives"
        drives = (
            requests.get(url_drives, headers=self._headers(), timeout=30)
            .json()
            .get("value", [])
        )

        for d in drives:
            if d["name"].lower() == self.drive_name.lower():
                self._drive_id = d["id"]
                return self._drive_id
        self._drive_id = drives[0]["id"]
        return self._drive_id

    # ── Folder / File Listing ─────────────────────────────────────────────

    def _list_folder_children(self, folder_path: str) -> list[dict]:
        drive_id = self._get_drive_id()
        encoded = quote(folder_path.strip("/"), safe="/")
        url = f"{self.GRAPH_BASE}/drives/{drive_id}/root:/{encoded}:/children?$select=id,name,file,folder&$top=999"
        items = []
        while url:
            resp = requests.get(url, headers=self._headers(), timeout=30)
            if not resp.ok:
                break
            data = resp.json()
            items.extend(data.get("value", []))
            url = data.get("@odata.nextLink")
        return items

    def list_resumes_grouped(self) -> dict[str, list[dict]]:
        """Return {subfolder_name: [list of .txt resume files]} from the text resumes folder."""
        subfolders = [
            i
            for i in self._list_folder_children(SP_TEXT_RESUMES_FOLDER)
            if "folder" in i
        ]
        groups = {}
        for sf in subfolders:
            files = [
                {"id": f["id"], "name": f["name"]}
                for f in self._list_folder_children(
                    f"{SP_TEXT_RESUMES_FOLDER}/{sf['name']}"
                )
                if "file" in f and f["name"].lower().endswith(".txt")
            ]
            if files:
                groups[sf["name"]] = files
        return groups

    def list_jd_files(self) -> list[dict]:
        return [
            {"id": f["id"], "name": f["name"]}
            for f in self._list_folder_children(SP_TEXT_JD_FOLDER)
            if "file" in f and f["name"].lower().endswith(".txt")
        ]

    def download_text_content(self, item_id: str) -> str:
        drive_id = self._get_drive_id()
        url = f"{self.GRAPH_BASE}/drives/{drive_id}/items/{item_id}/content"
        resp = requests.get(
            url, headers=self._headers(), timeout=60, allow_redirects=True
        )
        resp.raise_for_status()
        return resp.content.decode("utf-8", errors="replace")

    # ── MatchScore Operations ─────────────────────────────────────────────

    def find_exact_resume(self, subfolder: str, stem: str) -> dict | None:
        """Find the original resume (PDF or DOCX) in the correct subfolder by stem."""
        drive_id = self._get_drive_id()
        folder_path = f"{SP_ORIGINAL_RESUMES_FOLDER}/{subfolder}".strip("/")
        encoded_folder = quote(folder_path, safe="/")

        url = f"{self.GRAPH_BASE}/drives/{drive_id}/root:/{encoded_folder}:/children?$select=id,name,file&$expand=listItem"
        resp = requests.get(url, headers=self._headers(), timeout=30)

        if not resp.ok:
            return None

        for item in resp.json().get("value", []):
            if "file" in item:
                item_stem = Path(item["name"]).stem
                if item_stem.lower() == stem.lower():
                    return item
        return None

    def fetch_match_score(
        self, subfolder: str, stem: str
    ) -> tuple[str | None, int | None]:
        """Return (item_id, existing_score) for the original resume. Score is None if not set."""
        item = self.find_exact_resume(subfolder, stem)
        if not item:
            return None, None

        item_id = item["id"]
        fields = item.get("listItem", {}).get("fields", {})
        raw = fields.get("MatchScore")

        if raw is None or str(raw).strip() in ("", "0", "None"):
            return item_id, None

        try:
            return item_id, int(round(float(raw)))
        except (TypeError, ValueError):
            return item_id, None

    def push_match_score(
        self, subfolder: str, stem: str, score: int, confirmed_item_id: str = ""
    ) -> tuple[str, str]:
        """Write the MatchScore column on the original resume file in SharePoint."""
        drive_id = self._get_drive_id()

        if confirmed_item_id:
            item_id = confirmed_item_id
            filename_display = stem
        else:
            item = self.find_exact_resume(subfolder, stem)
            if not item:
                return (
                    "ERROR",
                    f"Original file for '{stem}' not found in '{SP_ORIGINAL_RESUMES_FOLDER}/{subfolder}'.",
                )
            item_id = item["id"]
            filename_display = item["name"]

        url = f"{self.GRAPH_BASE}/drives/{drive_id}/items/{item_id}/listItem/fields"
        resp = requests.patch(
            url, headers=self._headers(), json={"MatchScore": score}, timeout=30
        )

        if resp.status_code == 200:
            return (
                "OK",
                f"MatchScore = {score} written to `{filename_display}`.",
            )
        return "ERROR", f"HTTP {resp.status_code}: {resp.text[:200]}"


# ═══════════════════════════════════════════════════════════════════════════
# HELPERS — Utilities
# ═══════════════════════════════════════════════════════════════════════════


def _make_sp_client() -> SharePointClient:
    cfg = st.session_state.get("sp_config", {})
    return SharePointClient(
        tenant_id=cfg["tenant_id"],
        client_id=cfg["client_id"],
        client_secret=cfg["client_secret"],
        site_domain=cfg["site_domain"],
        site_path=cfg["site_path"],
        drive_name=cfg["drive_name"],
    )


def _extract_job_id_from_subfolder(subfolder: str) -> str:
    m = re.match(r"^(\w+?)_", subfolder)
    return m.group(1) if m else ""


def _find_jd_for_job_id(job_id: str, jd_files: list[dict]) -> dict | None:
    if not job_id:
        return None
    pattern = re.compile(rf"(?:^|[_\-]){re.escape(job_id)}(?:[_\-\.])", re.IGNORECASE)
    for f in jd_files:
        if pattern.search(f["name"]):
            return f
    return None


def _badge_html(status: str) -> str:
    s = status.lower().replace(" ", "-")
    if "match" == s:
        return '<span class="badge badge-match">Match</span>'
    elif "partial" in s:
        return '<span class="badge badge-partial">Partial</span>'
    elif "no" in s:
        return '<span class="badge badge-no-match">No Match</span>'
    elif "skip" in s:
        return '<span class="badge badge-skipped">Skipped</span>'
    else:
        return f'<span class="badge badge-partial">{status}</span>'


def _score_color(score: int) -> str:
    if score >= 70:
        return "#34d399"
    elif score >= 45:
        return "#fbbf24"
    else:
        return "#f87171"


# ═══════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════════════════

with st.sidebar:
    st.markdown(
        """
    <div style="padding: 0.5rem 0 1.5rem 0;">
        <div style="font-family:'DM Serif Display',serif; font-size:1.4rem; color:#e8eaf2; line-height:1.1;">Resume<br>Screener</div>
        <div style="font-size:0.7rem; color:#4b5563; text-transform:uppercase; letter-spacing:0.12em; font-weight:600; margin-top:0.3rem;">Batch Scoring Mode</div>
    </div>
    """,
        unsafe_allow_html=True,
    )

    st.markdown("---")
    st.markdown(
        '<span style="font-size:0.75rem;color:#6b7280;text-transform:uppercase;letter-spacing:0.1em;font-weight:600;">⚙ Configuration</span>',
        unsafe_allow_html=True,
    )

    api_url = st.text_input(
        "FastAPI Server URL",
        value=st.session_state.get("api_url", "http://localhost:8000"),
        key="api_url_input",
    )
    st.session_state["api_url"] = api_url

    st.markdown("---")

    _sp_tenant = os.getenv("AZURE_TENANT_ID", "")
    _sp_client = os.getenv("AZURE_CLIENT_ID", "")
    _sp_secret = os.getenv("AZURE_CLIENT_SECRET", "")
    _sp_domain = os.getenv("SHAREPOINT_SITE_DOMAIN", "")
    _sp_sitepath = os.getenv("SHAREPOINT_SITE_PATH", "")
    _sp_drive = os.getenv("SHAREPOINT_DRIVE_NAME", "Documents")

    _sp_ready = all(
        [_sp_tenant, _sp_client, _sp_secret, _sp_domain, _sp_sitepath, _sp_drive]
    )

    if _sp_ready and HAS_MSAL:
        st.markdown(
            f'<div style="font-size:0.82rem;color:#34d399;">✅ Connected to <b>{_sp_domain}</b></div>',
            unsafe_allow_html=True,
        )
    elif not HAS_MSAL:
        st.warning("⚠️ `msal` not installed")
    else:
        st.warning("⚠️ Missing AZURE_* / SHAREPOINT_* env vars")

    st.session_state["sp_config"] = {
        "enabled": _sp_ready and HAS_MSAL,
        "tenant_id": _sp_tenant,
        "client_id": _sp_client,
        "client_secret": _sp_secret,
        "site_domain": _sp_domain,
        "site_path": _sp_sitepath,
        "drive_name": _sp_drive,
    }

    st.markdown("---")
    st.markdown(
        '<span style="font-size:0.75rem;color:#6b7280;text-transform:uppercase;letter-spacing:0.1em;font-weight:600;">📊 Batch Stats</span>',
        unsafe_allow_html=True,
    )
    _batch_results = st.session_state.get("batch_results", [])
    if _batch_results:
        scored = [r for r in _batch_results if r.get("status") == "scored"]
        skipped = [r for r in _batch_results if r.get("status") == "skipped"]
        failed = [r for r in _batch_results if r.get("status") == "error"]
        st.markdown(
            f"Scored: **{len(scored)}** &nbsp; Skipped: **{len(skipped)}** &nbsp; Failed: **{len(failed)}**"
        )
        if scored:
            avg = sum(r["score"] for r in scored) / len(scored)
            st.markdown(f"Average Score: **{avg:.0f}**")
    else:
        st.markdown("No batch run yet.")


# ═══════════════════════════════════════════════════════════════════════════
# MAIN PAGE
# ═══════════════════════════════════════════════════════════════════════════

st.markdown(
    """<div style="padding: 2rem 0 0.5rem 0;">
        <div class="rs-title">Resume Screener</div>
        <div class="rs-subtitle">Batch Folder Scoring — Select folder, match JD, score all resumes</div>
    </div>""",
    unsafe_allow_html=True,
)

sp_cfg = st.session_state.get("sp_config", {})
sp_enabled = sp_cfg.get("enabled", False)

if not sp_enabled:
    st.error(
        "SharePoint integration is not configured. Please set the AZURE_* and SHAREPOINT_* environment variables."
    )
    st.stop()

# ── Load SharePoint file lists ────────────────────────────────────────────

if "sp_resume_groups" not in st.session_state:
    with st.spinner("📂 Loading files from SharePoint…"):
        try:
            _b = _make_sp_client()
            st.session_state["sp_resume_groups"] = _b.list_resumes_grouped()
            st.session_state["sp_jd_files"] = _b.list_jd_files()
        except Exception as _load_err:
            st.error(f"❌ Could not load SharePoint files: {_load_err}")
            st.stop()

_ref_col, _ = st.columns([1, 5])
with _ref_col:
    if st.button("🔄 Refresh File List", key="sp_refresh"):
        for _k in list(st.session_state.keys()):
            if _k in ("sp_resume_groups", "sp_jd_files") or _k.startswith(
                "sp_content_"
            ):
                del st.session_state[_k]
        st.rerun()


# ═══════════════════════════════════════════════════════════════════════════
# STEP 1: SELECT FOLDER + JD
# ═══════════════════════════════════════════════════════════════════════════

st.markdown('<div class="rs-divider"></div>', unsafe_allow_html=True)
st.markdown(
    '<div class="rs-section-header">① Select Job Role Folder & Job Description</div>',
    unsafe_allow_html=True,
)

_resume_groups = st.session_state.get("sp_resume_groups", {})
_jd_files = st.session_state.get("sp_jd_files", [])

col_folder, col_jd = st.columns(2)

with col_folder:
    st.markdown(
        '<div class="rs-card"><div class="param-title">Job Role Folder</div>',
        unsafe_allow_html=True,
    )
    folder_options = ["— Select folder —"] + sorted(_resume_groups.keys())
    chosen_folder = st.selectbox(
        "Folder",
        options=folder_options,
        key="batch_folder_select",
        label_visibility="collapsed",
    )
    if chosen_folder and chosen_folder != "— Select folder —":
        n = len(_resume_groups.get(chosen_folder, []))
        st.markdown(
            f'<span class="info-pill">📄 {n} resume(s) found</span>',
            unsafe_allow_html=True,
        )
    st.markdown("</div>", unsafe_allow_html=True)

with col_jd:
    st.markdown(
        '<div class="rs-card"><div class="param-title">Job Description</div>',
        unsafe_allow_html=True,
    )

    # Auto-detect JD based on folder name
    auto_jd = None
    if chosen_folder and chosen_folder != "— Select folder —":
        job_id = _extract_job_id_from_subfolder(chosen_folder)
        auto_jd = _find_jd_for_job_id(job_id, _jd_files)

    jd_options = ["— Select JD —"] + [f["name"] for f in _jd_files]
    default_jd_idx = 0
    if auto_jd:
        try:
            default_jd_idx = jd_options.index(auto_jd["name"])
        except ValueError:
            default_jd_idx = 0

    chosen_jd = st.selectbox(
        "JD File",
        options=jd_options,
        index=default_jd_idx,
        key="batch_jd_select",
        label_visibility="collapsed",
    )

    if auto_jd and chosen_jd == auto_jd["name"]:
        st.markdown(
            '<span class="info-pill">🔗 Auto-matched from folder name</span>',
            unsafe_allow_html=True,
        )
    st.markdown("</div>", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════
# STEP 2: BATCH OPTIONS
# ═══════════════════════════════════════════════════════════════════════════

st.markdown('<div class="rs-divider"></div>', unsafe_allow_html=True)
st.markdown(
    '<div class="rs-section-header">② Scoring Options</div>',
    unsafe_allow_html=True,
)

opt_col1, opt_col2 = st.columns(2)

with opt_col1:
    force_rescore = st.checkbox(
        "🔁 Force rescore ALL resumes (even those already scored)",
        value=False,
        key="batch_force_rescore",
    )

with opt_col2:
    conflict_policy = "use_new"  # default
    if force_rescore:
        conflict_policy = st.radio(
            "When a previous score exists, use:",
            options=["use_new", "use_average", "keep_previous"],
            format_func=lambda x: {
                "use_new": "🆕 New Score (overwrite)",
                "use_average": "📊 Average of Old + New",
                "keep_previous": "📌 Keep Previous (still re-analyze but don't update SP)",
            }[x],
            key="batch_conflict_policy",
            horizontal=False,
        )


# ═══════════════════════════════════════════════════════════════════════════
# STEP 3: RUN BATCH
# ═══════════════════════════════════════════════════════════════════════════

st.markdown('<div class="rs-divider"></div>', unsafe_allow_html=True)

_folder_valid = chosen_folder and chosen_folder != "— Select folder —"
_jd_valid = chosen_jd and chosen_jd != "— Select JD —"
ready = _folder_valid and _jd_valid

col_run, col_clear = st.columns([3, 1])
with col_run:
    run_clicked = st.button(
        "🚀  Start Batch Scoring",
        disabled=not ready,
        use_container_width=True,
        type="primary",
    )
with col_clear:
    if st.button("🗑  Clear Results", use_container_width=True):
        st.session_state.pop("batch_results", None)
        st.session_state.pop("batch_details", None)
        st.rerun()


# ═══════════════════════════════════════════════════════════════════════════
# BATCH EXECUTION LOGIC
# ═══════════════════════════════════════════════════════════════════════════

if run_clicked and ready:
    api_server = st.session_state.get("api_url", "http://localhost:8000").strip()
    sp_client = _make_sp_client()

    # 1. Download the JD text
    jd_item = next(f for f in _jd_files if f["name"] == chosen_jd)
    with st.spinner(f"📋 Downloading JD: `{chosen_jd}`…"):
        jd_text = sp_client.download_text_content(jd_item["id"])

    if not jd_text.strip():
        st.error("❌ JD text is empty. Cannot proceed.")
        st.stop()

    # 2. Get list of resumes in the folder
    resume_files = _resume_groups.get(chosen_folder, [])

    st.markdown(
        f"""<div class="rs-card">
        <div class="rs-section-header">⏳ Processing {len(resume_files)} resume(s) in <code>{chosen_folder}</code></div>
        </div>""",
        unsafe_allow_html=True,
    )

    progress_bar = st.progress(0, text="Starting batch…")
    status_container = st.container()
    batch_results = []
    batch_details = {}

    for idx, resume_file in enumerate(resume_files):
        fname = resume_file["name"]
        stem = Path(fname).stem
        pct = (idx + 1) / len(resume_files)
        progress_bar.progress(
            pct, text=f"Processing {idx + 1}/{len(resume_files)}: {fname}"
        )

        result_entry = {
            "filename": fname,
            "stem": stem,
            "status": "pending",
            "score": None,
            "old_score": None,
            "final_score": None,
            "message": "",
        }

        try:
            # A. Check existing MatchScore
            item_id, old_score = sp_client.fetch_match_score(chosen_folder, stem)
            result_entry["old_score"] = old_score

            if old_score is not None and not force_rescore:
                # Already scored, skip
                result_entry["status"] = "skipped"
                result_entry["score"] = old_score
                result_entry["final_score"] = old_score
                result_entry["message"] = f"Already scored ({old_score}). Skipped."
                batch_results.append(result_entry)
                with status_container:
                    st.markdown(
                        f"⏭️ **{fname}** — Existing score: **{old_score}** (skipped)"
                    )
                continue

            # B. Download resume text
            r_cache_key = f"sp_content_{resume_file['id']}"
            if r_cache_key in st.session_state:
                resume_text = st.session_state[r_cache_key]
            else:
                resume_text = sp_client.download_text_content(resume_file["id"])
                st.session_state[r_cache_key] = resume_text

            if not resume_text.strip():
                result_entry["status"] = "error"
                result_entry["message"] = "Resume text is empty."
                batch_results.append(result_entry)
                with status_container:
                    st.markdown(f"⚠️ **{fname}** — Empty text, skipped.")
                continue

            # C. Call the API (key rotation happens server-side)
            api_result = call_api(api_server, resume_text, jd_text)

            if api_result and "error" in api_result:
                result_entry["status"] = "error"
                result_entry["message"] = api_result["error"]
                batch_results.append(result_entry)
                with status_container:
                    st.markdown(
                        f"❌ **{fname}** — API error: {api_result['error'][:100]}"
                    )
                continue

            if not api_result:
                result_entry["status"] = "error"
                result_entry["message"] = "Empty API response."
                batch_results.append(result_entry)
                continue

            # D. Extract the new score
            new_score = round(
                api_result.get("function_1_resume_jd_matching", {}).get(
                    "overall_match_score", 0
                )
            )
            result_entry["score"] = new_score
            batch_details[stem] = api_result

            # E. Determine final score based on conflict policy
            if old_score is not None and force_rescore:
                if conflict_policy == "use_new":
                    final_score = new_score
                elif conflict_policy == "use_average":
                    final_score = round((old_score + new_score) / 2)
                else:  # keep_previous
                    final_score = old_score
            else:
                final_score = new_score

            result_entry["final_score"] = final_score

            # F. Push to SharePoint
            push_status, push_msg = sp_client.push_match_score(
                chosen_folder,
                stem,
                final_score,
                confirmed_item_id=item_id or "",
            )

            if push_status == "OK":
                result_entry["status"] = "scored"
                result_entry["message"] = push_msg
                status_icon = "✅"
            else:
                result_entry["status"] = "error"
                result_entry["message"] = push_msg
                status_icon = "⚠️"

            with status_container:
                old_tag = f" (was {old_score})" if old_score is not None else ""
                st.markdown(
                    f"{status_icon} **{fname}** — Score: **{final_score}**{old_tag}"
                )

            # Save JSON locally too
            try:
                candidate_name = (
                    api_result.get("function_2_resume_data_extraction", {})
                    .get("personal_information", {})
                    .get("full_name", stem)
                )
                save_result(api_result, chosen_folder, candidate_name)
            except Exception:
                pass

        except Exception as e:
            result_entry["status"] = "error"
            result_entry["message"] = str(e)
            with status_container:
                st.markdown(f"❌ **{fname}** — Error: {str(e)[:120]}")

        batch_results.append(result_entry)

        # Small delay to avoid hammering the API
        time.sleep(0.5)

    progress_bar.progress(1.0, text="✅ Batch complete!")
    st.session_state["batch_results"] = batch_results
    st.session_state["batch_details"] = batch_details


# ═══════════════════════════════════════════════════════════════════════════
# BATCH RESULTS DISPLAY
# ═══════════════════════════════════════════════════════════════════════════

if "batch_results" in st.session_state and st.session_state["batch_results"]:
    st.markdown('<div class="rs-divider"></div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="rs-section-header">📊 Batch Results</div>',
        unsafe_allow_html=True,
    )

    results = st.session_state["batch_results"]
    details = st.session_state.get("batch_details", {})

    # Summary stats
    scored = [r for r in results if r["status"] == "scored"]
    skipped = [r for r in results if r["status"] == "skipped"]
    errors = [r for r in results if r["status"] == "error"]

    s1, s2, s3, s4 = st.columns(4)
    with s1:
        st.markdown(
            f"""<div class="rs-card" style="text-align:center;">
            <div class="score-number" style="font-size:2.5rem;">{len(results)}</div>
            <div class="score-label">Total</div></div>""",
            unsafe_allow_html=True,
        )
    with s2:
        st.markdown(
            f"""<div class="rs-card" style="text-align:center;">
            <div style="font-size:2.5rem;font-weight:700;color:#34d399;">{len(scored)}</div>
            <div class="score-label">Scored</div></div>""",
            unsafe_allow_html=True,
        )
    with s3:
        st.markdown(
            f"""<div class="rs-card" style="text-align:center;">
            <div style="font-size:2.5rem;font-weight:700;color:#fbbf24;">{len(skipped)}</div>
            <div class="score-label">Skipped</div></div>""",
            unsafe_allow_html=True,
        )
    with s4:
        st.markdown(
            f"""<div class="rs-card" style="text-align:center;">
            <div style="font-size:2.5rem;font-weight:700;color:#f87171;">{len(errors)}</div>
            <div class="score-label">Failed</div></div>""",
            unsafe_allow_html=True,
        )

    # Sort by final score descending (scored first, then skipped, then errors)
    def _sort_key(r):
        order = {"scored": 0, "skipped": 1, "error": 2, "pending": 3}
        return (order.get(r["status"], 9), -(r.get("final_score") or 0))

    sorted_results = sorted(results, key=_sort_key)

    # Results table
    st.markdown(
        '<div class="rs-card" style="padding:0.8rem 1rem;">',
        unsafe_allow_html=True,
    )

    # Header row
    st.markdown(
        """<div style="display:flex;align-items:center;gap:1rem;padding:0.6rem 1rem;border-bottom:2px solid rgba(255,255,255,0.1);margin-bottom:0.3rem;">
            <div style="flex:1;font-size:0.72rem;color:#6b7280;text-transform:uppercase;font-weight:700;letter-spacing:0.1em;">Candidate</div>
            <div style="min-width:5rem;text-align:center;font-size:0.72rem;color:#6b7280;text-transform:uppercase;font-weight:700;letter-spacing:0.1em;">Old</div>
            <div style="min-width:5rem;text-align:center;font-size:0.72rem;color:#6b7280;text-transform:uppercase;font-weight:700;letter-spacing:0.1em;">New</div>
            <div style="min-width:5rem;text-align:center;font-size:0.72rem;color:#6b7280;text-transform:uppercase;font-weight:700;letter-spacing:0.1em;">Final</div>
            <div style="min-width:6rem;text-align:center;font-size:0.72rem;color:#6b7280;text-transform:uppercase;font-weight:700;letter-spacing:0.1em;">Status</div>
        </div>""",
        unsafe_allow_html=True,
    )

    for r in sorted_results:
        old_display = str(r["old_score"]) if r["old_score"] is not None else "—"
        new_display = str(r["score"]) if r["score"] is not None else "—"
        final = r.get("final_score")
        final_display = str(final) if final is not None else "—"
        final_color = _score_color(final) if final else "#6b7280"

        status_map = {
            "scored": '<span class="badge badge-match">Scored</span>',
            "skipped": '<span class="badge badge-skipped">Skipped</span>',
            "error": '<span class="badge badge-no-match">Error</span>',
        }
        status_badge = status_map.get(
            r["status"], '<span class="badge badge-partial">?</span>'
        )

        st.markdown(
            f"""<div class="batch-row">
                <div class="batch-name">{r["stem"]}</div>
                <div style="min-width:5rem;text-align:center;font-family:'JetBrains Mono',monospace;color:#6b7280;">{old_display}</div>
                <div style="min-width:5rem;text-align:center;font-family:'JetBrains Mono',monospace;color:#818cf8;">{new_display}</div>
                <div style="min-width:5rem;text-align:center;font-family:'JetBrains Mono',monospace;font-weight:700;color:{final_color};">{final_display}</div>
                <div class="batch-status">{status_badge}</div>
            </div>""",
            unsafe_allow_html=True,
        )

    st.markdown("</div>", unsafe_allow_html=True)

    # ── Expandable detail per resume ──────────────────────────────────────
    st.markdown('<div class="rs-divider"></div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="rs-section-header">🔍 Detailed Results (per resume)</div>',
        unsafe_allow_html=True,
    )

    for r in sorted_results:
        stem = r["stem"]
        detail = details.get(stem)
        if not detail:
            continue

        match_data = detail.get("function_1_resume_jd_matching", {})
        extract_data = detail.get("function_2_resume_data_extraction", {})
        personal = extract_data.get("personal_information", {})
        candidate_name = personal.get("full_name", stem)
        score_val = match_data.get("overall_match_score", 0)

        with st.expander(
            f"{'✅' if r['status'] == 'scored' else '⚠️'} {candidate_name} — Score: {score_val}",
            expanded=False,
        ):
            # Match parameters
            params = [
                "experience",
                "education",
                "location",
                "project_history_relevance",
                "tools_used",
                "certifications",
            ]

            pcols = st.columns(3)
            for i, param in enumerate(params):
                p = match_data.get(param, {})
                with pcols[i % 3]:
                    status_badge = _badge_html(p.get("status", "—"))
                    st.markdown(
                        f"""<div class="param-card">
                            <div class="param-title">{param.replace("_", " ")}</div>
                            {status_badge}
                            <div class="param-summary" style="margin-top:0.4rem;">{p.get("summary", "—")}</div>
                        </div>""",
                        unsafe_allow_html=True,
                    )

            # Candidate info
            st.markdown(
                '<div class="param-title" style="margin-top:0.8rem;">Candidate Profile</div>',
                unsafe_allow_html=True,
            )
            info_items = {
                "Name": personal.get("full_name", "—"),
                "Email": personal.get("email", "—"),
                "Phone": personal.get("phone", "—"),
                "Location": personal.get("location", "—"),
                "Current Role": extract_data.get("current_employment", {}).get(
                    "current_job_title", "—"
                ),
                "Organization": extract_data.get("current_employment", {}).get(
                    "current_organization", "—"
                ),
                "Experience": f"{extract_data.get('career_metrics', {}).get('total_experience_in_years', '—')} yrs",
            }
            for label, value in info_items.items():
                st.markdown(
                    f"""<div class="profile-field">
                        <span class="pf-label">{label}</span>
                        <span class="pf-value">{value}</span>
                    </div>""",
                    unsafe_allow_html=True,
                )

            # Raw JSON
            with st.expander("📄 Raw JSON", expanded=False):
                st.markdown(
                    f'<div class="json-viewer">{json.dumps(detail, indent=2, ensure_ascii=False)}</div>',
                    unsafe_allow_html=True,
                )


# ═══════════════════════════════════════════════════════════════════════════
# FORCE RESCORE — Interactive Per-Resume Score Resolution
# (Alternative mode: manual review after batch with "keep_previous" policy)
# ═══════════════════════════════════════════════════════════════════════════

_batch_results = st.session_state.get("batch_results", [])
_batch_details = st.session_state.get("batch_details", {})

# Check if there are any scored items with old_score (meaning conflicts existed)
_conflict_items = [
    r
    for r in _batch_results
    if r.get("old_score") is not None
    and r.get("score") is not None
    and r["status"] == "scored"
    and r["old_score"] != r["score"]
]

if _conflict_items and st.session_state.get("batch_force_rescore", False):
    st.markdown('<div class="rs-divider"></div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="rs-section-header">⚖️ Score Conflicts — Manual Override</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        "These resumes had existing scores that differed from the new analysis. "
        "The batch policy was applied automatically, but you can manually override individual scores below."
    )

    for r in _conflict_items:
        stem = r["stem"]
        old = r["old_score"]
        new = r["score"]
        avg = round((old + new) / 2)
        current_final = r.get("final_score", new)

        st.markdown(
            f"""<div class="rs-card">
            <div style="display:flex;justify-content:space-between;align-items:center;">
                <div>
                    <div style="font-size:0.95rem;color:#dde1ee;font-weight:600;">{stem}</div>
                    <div style="font-size:0.78rem;color:#6b7280;margin-top:0.2rem;">
                        Current final: <span style="color:{_score_color(current_final)};font-weight:700;">{current_final}</span>
                    </div>
                </div>
                <div style="display:flex;gap:1.2rem;">
                    <div style="text-align:center;"><div style="font-size:1.3rem;font-weight:700;color:#6b7280;">{old}</div><div style="font-size:0.65rem;color:#6b7280;text-transform:uppercase;">Previous</div></div>
                    <div style="text-align:center;"><div style="font-size:1.3rem;font-weight:700;color:#818cf8;">{new}</div><div style="font-size:0.65rem;color:#818cf8;text-transform:uppercase;">New</div></div>
                    <div style="text-align:center;"><div style="font-size:1.3rem;font-weight:700;color:#34d399;">{avg}</div><div style="font-size:0.65rem;color:#34d399;text-transform:uppercase;">Average</div></div>
                </div>
            </div>
            </div>""",
            unsafe_allow_html=True,
        )

        oc1, oc2, oc3 = st.columns(3)
        with oc1:
            if st.button(
                f"Use Previous ({old})",
                key=f"override_prev_{stem}",
                use_container_width=True,
            ):
                try:
                    sp = _make_sp_client()
                    s, m = sp.push_match_score(
                        chosen_folder
                        if _folder_valid
                        else st.session_state.get("batch_folder_select", ""),
                        stem,
                        old,
                    )
                    if s == "OK":
                        st.success(f"✅ {stem} → {old}")
                    else:
                        st.error(m)
                except Exception as e:
                    st.error(str(e))

        with oc2:
            if st.button(
                f"Use New ({new})",
                key=f"override_new_{stem}",
                use_container_width=True,
            ):
                try:
                    sp = _make_sp_client()
                    s, m = sp.push_match_score(
                        chosen_folder
                        if _folder_valid
                        else st.session_state.get("batch_folder_select", ""),
                        stem,
                        new,
                    )
                    if s == "OK":
                        st.success(f"✅ {stem} → {new}")
                    else:
                        st.error(m)
                except Exception as e:
                    st.error(str(e))

        with oc3:
            if st.button(
                f"Use Average ({avg})",
                key=f"override_avg_{stem}",
                use_container_width=True,
            ):
                try:
                    sp = _make_sp_client()
                    s, m = sp.push_match_score(
                        chosen_folder
                        if _folder_valid
                        else st.session_state.get("batch_folder_select", ""),
                        stem,
                        avg,
                    )
                    if s == "OK":
                        st.success(f"✅ {stem} → {avg}")
                    else:
                        st.error(m)
                except Exception as e:
                    st.error(str(e))
