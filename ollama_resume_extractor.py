"""
ollama_resume_extractor.py
──────────────────────────
Extract structured candidate data from resume files (PDF / DOCX) using a
local Ollama LLM.  Two extraction modes are supported:

  1. **structured** (default) – uses Ollama's `format` parameter to enforce a
     JSON schema at the grammar level.  The model *must* return data that
     conforms to the schema.
  2. **function_call** – uses Ollama's tool-calling / function-calling API.
     A tool named `save_candidate` is defined with the same schema; the model
     "calls" that tool and the script captures the arguments.

Usage examples:
    # Single resume (structured mode, default)
    python ollama_resume_extractor.py --resume resume.pdf

    # Single resume + Job Description for match scoring
    python ollama_resume_extractor.py --resume resume.pdf --jd jd.pdf

    # Batch – process every PDF/DOCX in a folder
    python ollama_resume_extractor.py --dir ./tmp_resumes --jd jd.pdf

    # Use function-calling mode instead of structured output
    python ollama_resume_extractor.py --resume resume.pdf --mode function_call

    # Choose a different Ollama model
    python ollama_resume_extractor.py --resume resume.pdf --model mistral

    # Save output to a JSON file
    python ollama_resume_extractor.py --resume resume.pdf -o output.json
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

import PyPDF2 as pdf
import requests
from docx import Document


# ═══════════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════════
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
DEFAULT_MODEL = "llama3"

# ═══════════════════════════════════════════════════════════════════════════════
# CANDIDATE SCHEMA  (shared by both modes)
# ═══════════════════════════════════════════════════════════════════════════════
CANDIDATE_SCHEMA = {
    "type": "object",
    "properties": {
        "full_name": {
            "type": "string", 
            "description": "Candidate's full name"
        },
        "current_job_title": {
            "type": "string", 
            "description": "Most recent professional title or N/A"
        },
        "explicit_skillset": {
            "type": "array",
            "items": {"type": "string"},
            "description": "List of skills explicitly listed in a Skills section"
        },
        "experience_skillset": {
            "type": "array",
            "items": {"type": "string"},
            "description": "List of skills inferred from job descriptions"
        },
        "job_history": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "company": {"type": "string"},
                    "title": {"type": "string"},
                    "duration": {"type": "string"},
                    "description": {"type": "string"}
                },
                "required": ["company", "title", "duration"]
            },
            "description": "List of past employment details"
        },
        "education": {
            "type": "array", 
            "items": {"type": "string"},
            "description": "List of degrees and institutions"
        },
        "certifications": {
            "type": "array", 
            "items": {"type": "string"},
            "description": "List of certifications"
        },
        # Optional fields (populated only if JD is provided)
        "match_score": {
            "type": "integer", 
            "description": "0-100 score matching JD"
        },
        "role_fit_summary": {
            "type": "string", 
            "description": "Brief explanation of fit"
        }
    },
    "required": [
        "full_name", 
        "current_job_title", 
        "explicit_skillset", 
        "experience_skillset", 
        "job_history", 
        "education", 
        "certifications"
    ],
}


# ═══════════════════════════════════════════════════════════════════════════════
# TOOL DEFINITION  (for function-calling mode)
# ═══════════════════════════════════════════════════════════════════════════════
SAVE_CANDIDATE_TOOL = {
    "type": "function",
    "function": {
        "name": "save_candidate",
        "description": (
            "Save the extracted candidate profile.  Call this function ONCE "
            "with every field populated from the resume."
        ),
        "parameters": CANDIDATE_SCHEMA,
    },
}


# ═══════════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION  (mirrors experiment_app.py)
# ═══════════════════════════════════════════════════════════════════════════════
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    with open(file_path, "rb") as f:
        reader = pdf.PdfReader(f)
        return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    doc = Document(file_path)
    parts = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(file_path: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}  (only .pdf and .docx)")


# ═══════════════════════════════════════════════════════════════════════════════
# PROMPT TEMPLATE
# ═══════════════════════════════════════════════════════════════════════════════
def build_prompt(resume_text: str, jd_text: str | None = None) -> str:
    """Build a concise extraction prompt focusing on specific candidate details."""
    
    jd_block = ""
    if jd_text:
        jd_block = f"""
JOB DESCRIPTION:
{jd_text[:2000]}

ADDITIONAL TASKS:
- match_score: (0-100) Relevance to JD
- role_fit_summary: Brief explanation of fit
"""

    return f"""You are a precise resume parser. Extract the following fields from the text below.

REQUIRED JSON FIELDS:
1. full_name: Candidate's full name
2. current_job_title: Most recent professional title
3. explicit_skillset: List of skills found in a dedicated 'Skills' section
4. experience_skillset: List of skills inferred or extracted specifically from job descriptions/history
5. job_history: List of objects containing {{ "company", "title", "duration", "description" }}
6. education: List of degrees, majors, and institutions
7. certifications: List of professional certifications

{jd_block}

RESUME TEXT:
{resume_text[:6000]}

CRITICAL: Return ONLY valid JSON. Use "N/A" if a field is not found.
"""


# ═══════════════════════════════════════════════════════════════════════════════
# OLLAMA CALLS
# ═══════════════════════════════════════════════════════════════════════════════
def _ollama_chat(messages: list, model: str, **extra) -> dict:
    """Low-level wrapper around POST /api/chat."""
    payload = {"model": model, "messages": messages, "stream": False, **extra}
    resp = requests.post(f"{OLLAMA_BASE_URL}/api/chat", json=payload, timeout=300)
    resp.raise_for_status()
    return resp.json()


# ── Mode 1: Structured Output ────────────────────────────────────────────────
def extract_structured(resume_text: str, jd_text: str | None, model: str) -> dict:
    """
    Uses Ollama's `format` parameter to enforce the JSON schema.
    The model is *forced* to return conforming JSON — no regex parsing needed.
    """
    prompt = build_prompt(resume_text, jd_text)
    messages = [
        {
            "role": "system",
            "content": (
                "You are a resume-parsing assistant.  Return the candidate "
                "profile as a JSON object matching the provided schema.  "
                "Do NOT include any text outside the JSON."
            ),
        },
        {"role": "user", "content": prompt},
    ]

    result = _ollama_chat(messages, model, format=CANDIDATE_SCHEMA)
    content = result.get("message", {}).get("content", "")

    try:
        return json.loads(content)
    except json.JSONDecodeError:
        # Fallback: try to extract JSON from the response via regex
        m = re.search(r"\{.*\}", content, re.DOTALL)
        if m:
            return json.loads(m.group(0))
        raise ValueError(f"Model did not return valid JSON:\n{content[:500]}")


# ── Mode 2: Function / Tool Calling ──────────────────────────────────────────
def extract_function_call(resume_text: str, jd_text: str | None, model: str) -> dict:
    """
    Uses Ollama's tool-calling API.  We define a `save_candidate` tool and
    instruct the model to call it with the extracted data.
    """
    prompt = build_prompt(resume_text, jd_text)
    prompt += (
        "\n\nAfter extracting the data, call the `save_candidate` function "
        "with ALL extracted fields."
    )

    messages = [
        {
            "role": "system",
            "content": (
                "You are a resume-parsing assistant.  Extract candidate data "
                "from the resume and call the save_candidate tool with the results."
            ),
        },
        {"role": "user", "content": prompt},
    ]

    result = _ollama_chat(messages, model, tools=[SAVE_CANDIDATE_TOOL])

    # The model's response should contain a tool_calls list
    msg = result.get("message", {})
    tool_calls = msg.get("tool_calls", [])

    if tool_calls:
        # Pick the first (and usually only) tool call
        call = tool_calls[0]
        fn = call.get("function", {})
        args = fn.get("arguments", {})
        if isinstance(args, str):
            args = json.loads(args)
        return args

    # Fallback: model may have returned JSON in content instead of a tool call
    content = msg.get("content", "")
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", content, re.DOTALL)
        if m:
            return json.loads(m.group(0))
        raise ValueError(
            "Model did not invoke the save_candidate tool and "
            f"returned no parseable JSON:\n{content[:500]}"
        )


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN DISPATCHER
# ═══════════════════════════════════════════════════════════════════════════════
EXTRACTORS = {
    "structured": extract_structured,
    "function_call": extract_function_call,
}


def process_single(
    resume_path: str,
    jd_text: str | None,
    model: str,
    mode: str,
) -> dict:
    """Process one resume file and return the candidate dict."""
    extractor = EXTRACTORS[mode]
    resume_text = extract_text(resume_path)
    if not resume_text or len(resume_text) < 50:
        return {"error": f"Could not extract meaningful text from {resume_path}"}

    try:
        data = extractor(resume_text, jd_text, model)
        data["_source_file"] = Path(resume_path).name
        data["_processed_at"] = datetime.now().isoformat()
        return data
    except Exception as e:
        return {
            "error": str(e),
            "_source_file": Path(resume_path).name,
            "_processed_at": datetime.now().isoformat(),
        }


def collect_files(directory: str) -> list[str]:
    """Gather all PDF/DOCX files from a directory."""
    exts = {".pdf", ".docx"}
    return sorted(
        str(p) for p in Path(directory).iterdir()
        if p.is_file() and p.suffix.lower() in exts
    )


# ═══════════════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════════════
def main():
    parser = argparse.ArgumentParser(
        description="Extract candidate data from resumes using a local Ollama LLM.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--resume", type=str, help="Path to a single resume (PDF/DOCX)")
    group.add_argument("--dir", type=str, help="Path to a directory of resumes")

    parser.add_argument("--jd", type=str, default=None,
                        help="Optional job description file for match scoring")
    parser.add_argument("--model", type=str, default=DEFAULT_MODEL,
                        help=f"Ollama model name (default: {DEFAULT_MODEL})")
    parser.add_argument("--mode", choices=["structured", "function_call"],
                        default="structured",
                        help="Extraction mode: 'structured' (schema-enforced) or "
                             "'function_call' (tool-calling API)")
    parser.add_argument("-o", "--output", type=str, default=None,
                        help="Save JSON output to this file")

    args = parser.parse_args()

    # ── Load optional JD ──────────────────────────────────────────────────
    jd_text = None
    if args.jd:
        jd_text = extract_text(args.jd)
        if not jd_text:
            print(f"⚠️  Warning: could not extract text from JD '{args.jd}'",
                  file=sys.stderr)

    # ── Gather resume file(s) ─────────────────────────────────────────────
    if args.resume:
        files = [args.resume]
    else:
        files = collect_files(args.dir)
        if not files:
            print(f"❌ No PDF/DOCX files found in '{args.dir}'", file=sys.stderr)
            sys.exit(1)

    print(f"🚀 Processing {len(files)} resume(s)  |  model={args.model}  "
          f"|  mode={args.mode}", file=sys.stderr)

    # ── Process ───────────────────────────────────────────────────────────
    results = []
    for i, fpath in enumerate(files, 1):
        print(f"  [{i}/{len(files)}] {Path(fpath).name} …", end=" ",
              file=sys.stderr, flush=True)
        result = process_single(fpath, jd_text, args.model, args.mode)
        results.append(result)
        status = "✅" if "error" not in result else "❌"
        print(status, file=sys.stderr)

    # ── Output ────────────────────────────────────────────────────────────
    output = results if len(results) > 1 else results[0]
    pretty = json.dumps(output, indent=2, ensure_ascii=False)

    # Always print to stdout
    print(pretty)

    # Optionally save to file
    if args.output:
        Path(args.output).write_text(pretty, encoding="utf-8")
        print(f"\n💾 Saved to {args.output}", file=sys.stderr)


if __name__ == "__main__":
    main()
