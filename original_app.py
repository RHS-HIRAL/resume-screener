import streamlit as st
import google.generativeai as genai
import os
import PyPDF2 as pdf
from dotenv import load_dotenv
import json
import re
import tempfile
from docx import Document
import pandas as pd
from datetime import datetime

# Load environment variables
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))


# ========== TEXT EXTRACTION ==========
def extract_text_from_pdf(uploaded_file):
    reader = pdf.PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        text += str(page.extract_text())
    return text


def extract_text_from_docx(uploaded_file):
    doc = Document(uploaded_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
    return text


def extract_text(uploaded_file):
    if uploaded_file.name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    elif uploaded_file.name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    return ""


# ========== AI-POWERED EXTRACTION ==========
def get_gemini_response(prompt):
    try:
        model = genai.GenerativeModel("gemini-2.5-flash")
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"ERROR: {str(e)}"


def extract_all_fields_with_ai(resume_text, jd_text):
    """Extract all fields using a single, comprehensive AI call"""
    prompt = f"""
You are an expert resume parser. Extract the following information from this resume with 100% accuracy.

CRITICAL INSTRUCTIONS:
1. Return ONLY valid JSON in the exact format shown below
2. If any field is not found, use "N/A"
3. For name: Look for the actual person's name (usually at the top)
4. For location: Look for city, state, or geographic location (NOT technical terms)
5. For current job: Find the most recent job title and company
6. For LinkedIn: Look for linkedin.com/in/ URLs or LinkedIn profile mentions
7. For job history: Format as "Company - Title - Duration" separated by " | "
8. For total experience: Extract as number only (e.g., "5" for 5 years)
9. Count all distinct jobs/positions in the work history

RESUME TEXT:
{resume_text[:3000]}

JOB DESCRIPTION:
{jd_text[:1000]}

Return JSON in this EXACT format:
{{
    "name": "Full Name Here",
    "role_in_jd": "Job title from JD",
    "current_job_title": "Current position title",
    "current_organization": "Current company name",
    "location": "City, State or geographic location",
    "phone": "Phone number",
    "email": "Email address",
    "linkedin": "LinkedIn profile URL or N/A",
    "other_socials": "Other social media or N/A",
    "total_experience": "X years",
    "job_history": "Company1 - Title1 - Duration1 | Company2 - Title2 - Duration2",
    "total_jobs": "Number of jobs",
    "match_score": "XX%",
    "summary": "Brief professional summary in 2-3 sentences"
}}
"""
    response = get_gemini_response(prompt)
    print("GEMINI RAW RESPONSE:", response)
    try:
        # Clean the response and extract JSON
        json_match = re.search(r"\{.*\}", response, re.DOTALL)
        if json_match:
            data = json.loads(json_match.group(0))
            return data
    except Exception as e:
        st.error(f"JSON parsing error: {str(e)}")
    return {}


# ========== BACKUP EXTRACTION FUNCTIONS ==========
def extract_name_backup(text):
    """Backup name extraction if AI fails"""
    lines = text.split("\n")[:8]
    for line in lines:
        line = line.strip()
        if any(
            keyword in line.lower()
            for keyword in [
                "resume",
                "cv",
                "email",
                "phone",
                "address",
                "linkedin",
                "objective",
                "summary",
                "experience",
            ]
        ):
            continue
        if line and 2 <= len(line.split()) <= 4:
            if re.match(r"^[A-Za-z\s\.]+$", line) and not line.isupper():
                return line
    return "N/A"


def extract_contact_backup(text):
    """Backup contact extraction"""
    email_match = re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)
    email = email_match.group(0) if email_match else "N/A"
    phone_patterns = [
        r"(?:\+91|91)[\s\-]?[6-9]\d{9}",
        r"[6-9]\d{9}",
        r"\d{3}[\s\-]?\d{3}[\s\-]?\d{4}",
    ]
    phone = "N/A"
    for pattern in phone_patterns:
        match = re.search(pattern, text)
        if match:
            phone = match.group(0)
            break
    linkedin_match = re.search(
        r"linkedin\.com/in/([A-Za-z0-9\-]+)", text, re.IGNORECASE
    )
    linkedin = f"linkedin.com/in/{linkedin_match.group(1)}" if linkedin_match else "N/A"
    return email, phone, linkedin


def extract_experience_backup(text):
    """Backup experience extraction"""
    patterns = [
        r"(?:total|overall)[\s\-:]*(?:experience|exp)[\s\-:]*(\d+(?:\.\d+)?)\s*(?:years?|yrs?)",
        r"(\d+(?:\.\d+)?)\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            exp = float(match.group(1))
            if 0 <= exp <= 50:
                return f"{exp} years"
    return "N/A"


def count_jobs_from_history(job_history):
    """Count number of jobs from job history string"""
    if not job_history or job_history == "N/A":
        return 0
    job_entries = job_history.split(" | ")
    return len([entry for entry in job_entries if entry.strip()])


def extract_numeric_experience(experience_str):
    """Extract numeric value from experience string"""
    if not experience_str or experience_str == "N/A":
        return 0
    match = re.search(r"(\d+(?:\.\d+)?)", experience_str)
    if match:
        return float(match.group(1))
    return 0


def calculate_average_tenure(total_experience, total_jobs):
    """Calculate average tenure per job"""
    try:
        exp_years = extract_numeric_experience(total_experience)
        jobs_count = (
            int(total_jobs)
            if str(total_jobs).isdigit()
            else count_jobs_from_history(str(total_jobs))
        )
        if exp_years > 0 and jobs_count > 0:
            avg_tenure = exp_years / jobs_count
            return f"{avg_tenure:.1f} years"
        else:
            return "N/A"
    except:
        return "N/A"


# ========== MAIN PROCESSING ==========
def process_resume_enhanced(resume_file, jd_text, index):
    """Enhanced resume processing with AI + backup extraction"""
    try:
        resume_text = extract_text(resume_file)
        if not resume_text or len(resume_text) < 50:
            return create_error_record(index, "Text extraction failed")
        ai_data = extract_all_fields_with_ai(resume_text, jd_text)
        backup_name = extract_name_backup(resume_text)
        backup_email, backup_phone, backup_linkedin = extract_contact_backup(
            resume_text
        )
        backup_experience = extract_experience_backup(resume_text)
        job_history = ai_data.get("job_history", "N/A")
        total_jobs = ai_data.get("total_jobs", count_jobs_from_history(job_history))
        total_experience = (
            ai_data.get("total_experience", backup_experience)
            if ai_data.get("total_experience")
            and ai_data.get("total_experience") != "N/A"
            else backup_experience
        )
        avg_tenure = calculate_average_tenure(total_experience, total_jobs)
        result = {
            "Sr.no": index,
            "Name": ai_data.get("name", backup_name)
            if ai_data.get("name") and ai_data.get("name") != "N/A"
            else backup_name,
            "Role in JD": ai_data.get("role_in_jd", "N/A"),
            "Current Job Title": ai_data.get("current_job_title", "N/A"),
            "Current Organization": ai_data.get("current_organization", "N/A"),
            "Location": ai_data.get("location", "N/A"),
            "Resume Match Score": ai_data.get("match_score", "N/A"),
            "Resume Summary": ai_data.get("summary", "N/A"),
            "Phone": ai_data.get("phone", backup_phone)
            if ai_data.get("phone") and ai_data.get("phone") != "N/A"
            else backup_phone,
            "Email": ai_data.get("email", backup_email)
            if ai_data.get("email") and ai_data.get("email") != "N/A"
            else backup_email,
            "LinkedIn": ai_data.get("linkedin", backup_linkedin)
            if ai_data.get("linkedin") and ai_data.get("linkedin") != "N/A"
            else backup_linkedin,
            "Other Socials": ai_data.get("other_socials", "N/A"),
            "Total Experience": total_experience,
            "Total Jobs": str(total_jobs)
            if str(total_jobs).isdigit()
            else str(count_jobs_from_history(job_history)),
            "Average Tenure": avg_tenure,
            "Job History (org-title-jobDuration)": job_history,
        }
        return result
    except Exception as e:
        return create_error_record(index, f"Processing error: {str(e)}")


def create_error_record(index, error_msg):
    """Create error record"""
    return {
        "Sr.no": index,
        "Name": "N/A",
        "Role in JD": "N/A",
        "Current Job Title": "N/A",
        "Current Organization": "N/A",
        "Location": "N/A",
        "Resume Match Score": "Error",
        "Resume Summary": error_msg,
        "Phone": "N/A",
        "Email": "N/A",
        "LinkedIn": "N/A",
        "Other Socials": "N/A",
        "Total Experience": "N/A",
        "Total Jobs": "N/A",
        "Average Tenure": "N/A",
        "Job History (org-title-jobDuration)": "N/A",
    }


# ========== STREAMLIT UI ==========
def main():
    st.set_page_config(page_title="🎯 Resume Screener", layout="wide")

    st.title("🎯 Resume Screener")
    st.markdown(
        "**Fine-tuned AI extraction for precise results with Average Tenure Calculation**"
    )

    # File upload
    col1, col2 = st.columns(2)
    with col1:
        jd_file = st.file_uploader("📋 Upload Job Description", type=["pdf", "docx"])
    with col2:
        resume_files = st.file_uploader(
            "📄 Upload Resume Files", type=["pdf", "docx"], accept_multiple_files=True
        )

    if st.button("🚀 Start Accurate Processing", type="primary"):
        if not jd_file or not resume_files:
            st.error("⚠️ Please upload both JD and resume files")
            return

        jd_text = extract_text(jd_file)
        if not jd_text:
            st.error("❌ Failed to extract JD text")
            return

        results = []
        progress_bar = st.progress(0)
        for idx, resume_file in enumerate(resume_files, 1):
            st.info(f"🔍 Processing: {resume_file.name} ({idx}/{len(resume_files)})")
            result = process_resume_enhanced(resume_file, jd_text, idx)
            results.append(result)
            progress_bar.progress(idx / len(resume_files))

        # Display results
        progress_bar.empty()
        df = pd.DataFrame(results)
        st.success(f"✅ Successfully processed {len(results)} resumes!")

        # Show sample results for verification
        st.subheader("📋 Sample Results (First 3 rows)")
        st.dataframe(df.head(3), use_container_width=True)

        # Statistics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("📊 Total Resumes", len(results))
        with col2:
            successful = sum(1 for r in results if r["Resume Match Score"] != "Error")
            st.metric("✅ Successful", successful)
        with col3:
            names_extracted = sum(1 for r in results if r["Name"] != "N/A")
            st.metric("👤 Names Extracted", names_extracted)
        with col4:
            tenure_calculated = sum(1 for r in results if r["Average Tenure"] != "N/A")
            st.metric("📈 Tenure Calculated", tenure_calculated)

        # Show average tenure insights
        st.subheader("📊 Average Tenure Insights")
        valid_tenures = [
            r["Average Tenure"] for r in results if r["Average Tenure"] != "N/A"
        ]
        tenure_values = []
        for tenure in valid_tenures:
            try:
                value = float(tenure.replace(" years", ""))
                tenure_values.append(value)
            except:
                pass

        if tenure_values:
            avg_tenure_overall = sum(tenure_values) / len(tenure_values)
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric(
                    "📊 Average Tenure Overall", f"{avg_tenure_overall:.1f} years"
                )
            with col2:
                st.metric("📈 Highest Tenure", f"{max(tenure_values):.1f} years")
            with col3:
                st.metric("📉 Lowest Tenure", f"{min(tenure_values):.1f} years")

            st.subheader("📊 Tenure Distribution")
            tenure_ranges = {
                "0-2 years": sum(1 for t in tenure_values if 0 <= t <= 2),
                "2-5 years": sum(1 for t in tenure_values if 2 < t <= 5),
                "5-10 years": sum(1 for t in tenure_values if 5 < t <= 10),
                "10+ years": sum(1 for t in tenure_values if t > 10),
            }
            for range_name, count in tenure_ranges.items():
                percentage = (count / len(tenure_values)) * 100
                st.write(f"• **{range_name}**: {count} candidates ({percentage:.1f}%)")

        # Full results
        with st.expander("📈 View All Results"):
            st.dataframe(df, use_container_width=True)

        # Download
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        excel_path = os.path.join(
            tempfile.gettempdir(), f"accurate_resume_results_{timestamp}.xlsx"
        )
        df.to_excel(excel_path, index=False)
        with open(excel_path, "rb") as f:
            st.download_button(
                label="📥 Download Accurate Results",
                data=f,
                file_name=f"accurate_resume_results_{timestamp}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        st.subheader("🔍 Extraction Quality Check")
        quality_metrics = {
            "Names Extracted": f"{names_extracted}/{len(results)} ({names_extracted / len(results) * 100:.1f}%)",
            "Emails Extracted": f"{sum(1 for r in results if r['Email'] != 'N/A')}/{len(results)}",
            "Phones Extracted": f"{sum(1 for r in results if r['Phone'] != 'N/A')}/{len(results)}",
            "LinkedIn Extracted": f"{sum(1 for r in results if r['LinkedIn'] != 'N/A')}/{len(results)}",
            "Average Tenure Calculated": f"{tenure_calculated}/{len(results)} ({tenure_calculated / len(results) * 100:.1f}%)",
            "Job History Extracted": f"{sum(1 for r in results if r['Job History (org-title-jobDuration)'] != 'N/A')}/{len(results)}",
        }
        for metric, value in quality_metrics.items():
            st.write(f"• **{metric}**: {value}")


if __name__ == "__main__":
    main()
