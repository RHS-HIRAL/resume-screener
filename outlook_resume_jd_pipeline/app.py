"""
resume_pipeline.py — Unified Resume Pipeline
═══════════════════════════════════════════════
Combines two workflows into a single, self-contained script:

  Pipeline 1 — Email Fetch & Upload
    Monitors an Outlook mailbox for new application emails, parses candidate
    info, downloads resume attachments (PDF/DOCX), and uploads them to
    SharePoint under Resumes/<JobID>_<JobRole>/<Candidate>_<JobID>.ext

  Pipeline 2 — Text Extraction & Upload
    Iterates through all resume files already on SharePoint, extracts raw
    text (PyPDF2 / python-docx / OCR fallback), and uploads .txt versions
    to Text Files/NewResumes/<Role>/<Filename>.txt

Usage:
    python resume_pipeline.py                  # Run both sequentially
    python resume_pipeline.py --all            # Same as above
    python resume_pipeline.py --fetch-emails   # Pipeline 1 only
    python resume_pipeline.py --extract-text   # Pipeline 2 only
"""

import argparse
import hashlib
import io
import logging
import os
import re
import shutil
import sys
import time
from dataclasses import dataclass, field
from datetime import datetime, timedelta, timezone
from pathlib import Path
from urllib.parse import quote

import msal
import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv

load_dotenv()

# ── Optional OCR dependencies (graceful fallback) ────────────────────────────
try:
    import fitz as pymupdf  # PyMuPDF

    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import pytesseract
    from PIL import Image

    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False

# ── Optional text extraction libraries ───────────────────────────────────────
try:
    import PyPDF2

    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from docx import Document as DocxDocument

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ═══════════════════════════════════════════════════════════════════════════════
#  CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════════


class Config:
    """All values loaded from environment variables or .env file."""

    # ─── Entra ID (Azure AD) Credentials ────────────────────────────────────
    TENANT_ID = os.getenv("AZURE_TENANT_ID", "")
    CLIENT_ID = os.getenv("AZURE_CLIENT_ID", "")
    CLIENT_SECRET = os.getenv("AZURE_CLIENT_SECRET", "")

    # ─── Microsoft Graph API ─────────────────────────────────────────────────
    GRAPH_BASE_URL = "https://graph.microsoft.com/v1.0"
    AUTHORITY = f"https://login.microsoftonline.com/{TENANT_ID}"
    SCOPES = ["https://graph.microsoft.com/.default"]

    # ─── Mailbox to Monitor ──────────────────────────────────────────────────
    MAILBOX_USER = os.getenv("MAILBOX_USER", "recruitment@yourcompany.com")

    # ─── Email Filtering ─────────────────────────────────────────────────────
    SUBJECT_KEYWORDS = ["new application received"]
    LOOKBACK_HOURS = 100

    # ─── SharePoint Target ───────────────────────────────────────────────────
    SHAREPOINT_SITE_DOMAIN = os.getenv(
        "SHAREPOINT_SITE_DOMAIN", "yourcompany.sharepoint.com"
    )
    SHAREPOINT_SITE_PATH = os.getenv("SHAREPOINT_SITE_PATH", "/sites/Recruitment")
    SHAREPOINT_DRIVE_NAME = os.getenv("SHAREPOINT_DRIVE_NAME", "Documents")
    SHAREPOINT_BASE_FOLDER = os.getenv("SHAREPOINT_BASE_FOLDER", "Resumes")

    # ─── Text Extraction Target ──────────────────────────────────────────────
    TEXT_RESUMES_FOLDER = os.getenv(
        "SHAREPOINT_TEXT_RESUMES_FOLDER", "Text Files/NewResumes"
    )

    # ─── File Naming ─────────────────────────────────────────────────────────
    FILE_NAME_TEMPLATE = "{name}_{job_id}{ext}"
    SUBFOLDER_TEMPLATE = "{job_id}_{job_role}"

    # ─── Notifications & Logging ─────────────────────────────────────────────
    TEAMS_WEBHOOK_URL = os.getenv("TEAMS_WEBHOOK_URL", "")
    LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO")
    LOG_FILE = os.getenv("LOG_FILE", "logs/pipeline.log")
    TEMP_DIR = os.getenv("TEMP_DIR", "./tmp_resumes")

    # ─── OCR Config ──────────────────────────────────────────────────────────
    TESSERACT_CMD = os.getenv(
        "TESSERACT_CMD", r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    )
    OCR_DPI = int(os.getenv("OCR_DPI", "300"))


# ═══════════════════════════════════════════════════════════════════════════════
#  LOGGING
# ═══════════════════════════════════════════════════════════════════════════════


def setup_logging():
    fmt = "%(asctime)s │ %(levelname)-7s │ %(name)-22s │ %(message)s"
    os.makedirs("logs", exist_ok=True)
    logging.basicConfig(
        level=getattr(logging, Config.LOG_LEVEL, logging.INFO),
        format=fmt,
        handlers=[
            logging.StreamHandler(sys.stdout),
            logging.FileHandler(Config.LOG_FILE, encoding="utf-8"),
        ],
    )


logger = logging.getLogger("pipeline")


# ═══════════════════════════════════════════════════════════════════════════════
#  AUTH — Unified Microsoft Graph Authentication
# ═══════════════════════════════════════════════════════════════════════════════


class GraphAuthProvider:
    """Handles MSAL token acquisition for Microsoft Graph API."""

    def __init__(self):
        self._app = msal.ConfidentialClientApplication(
            client_id=Config.CLIENT_ID,
            client_credential=Config.CLIENT_SECRET,
            authority=Config.AUTHORITY,
        )

    def get_access_token(self) -> str:
        result = self._app.acquire_token_silent(Config.SCOPES, account=None)
        if not result:
            logger.info("Acquiring new token via client credentials.")
            result = self._app.acquire_token_for_client(scopes=Config.SCOPES)

        if "access_token" in result:
            return result["access_token"]

        error = result.get("error_description", result.get("error", "Unknown error"))
        raise RuntimeError(f"Could not acquire token: {error}")

    def get_headers(self) -> dict:
        return {
            "Authorization": f"Bearer {self.get_access_token()}",
            "Content-Type": "application/json",
        }


# ═══════════════════════════════════════════════════════════════════════════════
#  DATA MODELS
# ═══════════════════════════════════════════════════════════════════════════════

SUBJECT_PATTERN = re.compile(
    r"for the position:\s*(.+?)\s*\[(\w+)\]\s*$", re.IGNORECASE
)
JOB_ID_BRACKET = re.compile(r"\[(\w+)\]\s*$")
FIELD_PATTERNS = {
    "job_opening": re.compile(r"Job\s*Opening\s*:\s*(.+)", re.IGNORECASE),
    "name": re.compile(r"Name\s*:\s*(.+)", re.IGNORECASE),
    "email": re.compile(r"Email\s*:\s*(\S+@\S+\.\S+)", re.IGNORECASE),
    "phone": re.compile(r"Phone\s*:\s*([\d\s\+\-().]+)", re.IGNORECASE),
    "resume_url": re.compile(r"Resume\s*:\s*(https?://\S+)", re.IGNORECASE),
}


@dataclass
class CandidateInfo:
    name: str = ""
    email: str = ""
    phone: str = ""
    job_role: str = ""
    job_id: str = ""
    resume_url: str = ""
    attachments: list[dict] = field(default_factory=list)
    source_email_id: str = ""
    source_subject: str = ""
    received_datetime: str = ""
    received_date: str = ""

    @property
    def safe_name(self) -> str:
        cleaned = re.sub(r"[^\w\s\-]", "", self.name).strip()
        return "_".join(w.capitalize() for w in cleaned.split()) or "Unknown"

    @property
    def safe_job_id(self) -> str:
        return re.sub(r"[^\w\-]", "", self.job_id).strip() or "NO-ID"

    @property
    def safe_job_role(self) -> str:
        cleaned = re.sub(r"[^\w\s\-]", "", self.job_role).strip()
        return "_".join(cleaned.split())[:80] or "General"


# ═══════════════════════════════════════════════════════════════════════════════
#  EMAIL FETCHER
# ═══════════════════════════════════════════════════════════════════════════════


class EmailFetcher:
    """Fetches and parses candidate application emails from an Outlook mailbox."""

    def __init__(self, auth_headers: dict):
        self.headers = auth_headers
        self.base = Config.GRAPH_BASE_URL

    def fetch_recent_emails(self) -> list[CandidateInfo]:
        raw_emails = self._get_emails_since(hours=Config.LOOKBACK_HOURS)
        logger.info(
            "Fetched %d emails from the last %d hours.",
            len(raw_emails),
            Config.LOOKBACK_HOURS,
        )

        candidates: list[CandidateInfo] = []
        for msg in raw_emails:
            if not self._is_relevant(msg):
                continue
            candidates.append(self._parse_email(msg))
        return candidates

    def _get_emails_since(self, hours: int) -> list[dict]:
        since = (datetime.now(timezone.utc) - timedelta(hours=hours)).strftime(
            "%Y-%m-%dT%H:%M:%SZ"
        )
        url = (
            f"{self.base}/users/{Config.MAILBOX_USER}/messages"
            f"?$filter=receivedDateTime ge {since}"
            f"&$orderby=receivedDateTime desc"
            f"&$top=200"
            f"&$select=id,subject,from,receivedDateTime,body,hasAttachments"
        )
        all_messages = []
        while url:
            resp = requests.get(url, headers=self.headers, timeout=30)
            resp.raise_for_status()
            data = resp.json()
            all_messages.extend(data.get("value", []))
            url = data.get("@odata.nextLink")
        return all_messages

    def get_attachment_content(self, email_id: str, attachment_id: str) -> bytes:
        url = (
            f"{self.base}/users/{Config.MAILBOX_USER}"
            f"/messages/{email_id}/attachments/{attachment_id}/$value"
        )
        resp = requests.get(url, headers=self.headers, timeout=60)
        resp.raise_for_status()
        return resp.content

    def get_attachments_metadata(self, email_id: str) -> list[dict]:
        url = (
            f"{self.base}/users/{Config.MAILBOX_USER}"
            f"/messages/{email_id}/attachments?$select=id,name,contentType"
        )
        resp = requests.get(url, headers=self.headers, timeout=30)
        resp.raise_for_status()
        return resp.json().get("value", [])

    def _is_relevant(self, msg: dict) -> bool:
        if not Config.SUBJECT_KEYWORDS:
            return True
        subject = (msg.get("subject") or "").lower()
        return any(kw in subject for kw in Config.SUBJECT_KEYWORDS)

    def _parse_email(self, msg: dict) -> CandidateInfo:
        subject = msg.get("subject", "")
        body_html = msg.get("body", {}).get("content", "")
        body_text = BeautifulSoup(body_html, "html.parser").get_text(separator="\n")
        received_dt = msg.get("receivedDateTime", "")

        candidate = CandidateInfo(
            source_email_id=msg["id"],
            source_subject=subject,
            received_datetime=received_dt,
            received_date=received_dt[:10],
        )

        subj_match = SUBJECT_PATTERN.search(subject)
        if subj_match:
            candidate.job_role = subj_match.group(1).strip()
            candidate.job_id = subj_match.group(2).strip()

        for line in body_text.splitlines():
            line = line.strip()
            if not line:
                continue
            for field_key, pattern in FIELD_PATTERNS.items():
                m = pattern.search(line)
                if not m:
                    continue
                value = m.group(1).strip()

                if field_key == "name":
                    candidate.name = value.title()
                elif field_key == "email":
                    candidate.email = value.lower()
                elif field_key == "phone":
                    candidate.phone = re.sub(r"[^\d+\-() ]", "", value).strip()
                elif field_key == "resume_url":
                    candidate.resume_url = value
                elif field_key == "job_opening":
                    if not candidate.job_id:
                        bracket = JOB_ID_BRACKET.search(value)
                        if bracket:
                            candidate.job_id = bracket.group(1)
                            candidate.job_role = value[: bracket.start()].strip()
                        else:
                            candidate.job_role = candidate.job_role or value

        if not candidate.resume_url and msg.get("hasAttachments"):
            attachments = self.get_attachments_metadata(msg["id"])
            valid_ext = (".pdf", ".docx", ".doc")
            valid_mime = ["pdf", "wordprocessingml", "msword"]

            candidate.attachments = [
                {
                    "id": a["id"],
                    "name": a["name"],
                    "content_type": a.get("contentType", ""),
                }
                for a in attachments
                if a.get("name", "").lower().endswith(valid_ext)
                or any(ct in a.get("contentType", "").lower() for ct in valid_mime)
            ]

        if not candidate.name:
            sender = msg.get("from", {}).get("emailAddress", {}).get("name", "")
            candidate.name = sender.title() if sender else "Unknown"
        if not candidate.email:
            candidate.email = (
                msg.get("from", {}).get("emailAddress", {}).get("address", "")
            )

        return candidate


# ═══════════════════════════════════════════════════════════════════════════════
#  SHAREPOINT MANAGER — Unified Upload, Download, Listing, Metadata
# ═══════════════════════════════════════════════════════════════════════════════

FIELD_MAP = {
    "CandidateName": "CandidateName",
    "CandidateEmail": "CandidateEmail",
    "CandidatePhone": "CandidatePhone",
    "JobID": "JobID",
    "JobRole": "JobRole",
    "SourceEmailID": "SourceEmailID",
}


class SharePointManager:
    """
    Single class for all SharePoint operations across both pipelines:
    - Folder management and file uploads (Pipeline 1)
    - File listing, downloading, and text upload (Pipeline 2)
    """

    def __init__(self, auth_headers: dict):
        self.headers = auth_headers
        self.base = Config.GRAPH_BASE_URL
        self._site_id: str | None = None
        self._drive_id: str | None = None
        self._ensured_folders: set[str] = set()

    # ── Site & Drive Resolution ───────────────────────────────────────────────

    def _get_site_id(self) -> str:
        if self._site_id:
            return self._site_id
        domain = Config.SHAREPOINT_SITE_DOMAIN
        path = Config.SHAREPOINT_SITE_PATH.strip("/")
        url = f"{self.base}/sites/{domain}:/{path}"
        resp = requests.get(url, headers=self.headers, timeout=30)
        resp.raise_for_status()
        self._site_id = resp.json()["id"]
        return self._site_id

    def _get_drive_id(self) -> str:
        if self._drive_id:
            return self._drive_id
        url = f"{self.base}/sites/{self._get_site_id()}/drives"
        resp = requests.get(url, headers=self.headers, timeout=30)
        resp.raise_for_status()
        for d in resp.json().get("value", []):
            if d["name"].lower() == Config.SHAREPOINT_DRIVE_NAME.lower():
                self._drive_id = d["id"]
                return self._drive_id
        self._drive_id = resp.json()["value"][0]["id"]
        return self._drive_id

    # ── Folder Management ─────────────────────────────────────────────────────

    def _ensure_folder(self, drive_id: str, folder_path: str) -> None:
        if folder_path in self._ensured_folders:
            return
        current = ""
        for part in folder_path.strip("/").split("/"):
            current = f"{current}/{part}" if current else part
            if current in self._ensured_folders:
                continue

            check_url = f"{self.base}/drives/{drive_id}/root:/{quote(current)}"
            if (
                requests.get(check_url, headers=self.headers, timeout=15).status_code
                == 404
            ):
                create_url = f"{self.base}/drives/{drive_id}/root/children"
                if "/" in current:
                    parent_encoded = quote("/".join(current.split("/")[:-1]))
                    create_url = f"{self.base}/drives/{drive_id}/root:/{parent_encoded}:/children"

                requests.post(
                    create_url,
                    headers=self.headers,
                    json={
                        "name": part,
                        "folder": {},
                        "@microsoft.graph.conflictBehavior": "fail",
                    },
                    timeout=15,
                )
            self._ensured_folders.add(current)

    def ensure_base_folder(self) -> None:
        self._ensure_folder(
            self._get_drive_id(), Config.SHAREPOINT_BASE_FOLDER.strip("/")
        )

    # ── File Existence Check ──────────────────────────────────────────────────

    def file_exists(self, remote_path: str) -> bool:
        """Check if a file exists on SharePoint at the given path."""
        drive_id = self._get_drive_id()
        encoded = quote(remote_path.strip("/"))
        url = f"{self.base}/drives/{drive_id}/root:/{encoded}"
        try:
            resp = requests.get(url, headers=self.headers, timeout=10)
            return resp.status_code == 200
        except Exception:
            return False

    # ── File Metadata ─────────────────────────────────────────────────────────

    def get_file_metadata(self, folder: str, filename: str) -> dict | None:
        """Returns listItem fields if the file exists, else None."""
        drive_id = self._get_drive_id()
        encoded_path = quote(f"{folder}/{filename}")
        url = f"{self.base}/drives/{drive_id}/root:/{encoded_path}?$expand=listItem"
        try:
            resp = requests.get(url, headers=self.headers, timeout=15)
            if resp.status_code == 200:
                return resp.json().get("listItem", {}).get("fields", {})
            return None
        except Exception:
            return None

    def _set_metadata(self, drive_id: str, item_id: str, metadata: dict) -> None:
        if item_id == "resumable_upload_complete":
            return
        url = f"{self.base}/drives/{drive_id}/items/{item_id}/listItem/fields"
        fields = {FIELD_MAP[k]: v for k, v in metadata.items() if k in FIELD_MAP and v}
        if fields:
            requests.patch(url, headers=self.headers, json=fields, timeout=30)

    # ── Folder / File Listing ─────────────────────────────────────────────────

    def list_children(self, folder_path: str) -> list[dict]:
        """List all children (files & folders) in a SharePoint folder."""
        drive_id = self._get_drive_id()
        encoded = quote(folder_path.strip("/"), safe="/")
        url = (
            f"{self.base}/drives/{drive_id}/root:/{encoded}:/children"
            f"?$select=id,name,file,folder&$top=999"
        )
        items = []
        while url:
            resp = requests.get(url, headers=self.headers, timeout=30)
            if not resp.ok:
                break
            data = resp.json()
            items.extend(data.get("value", []))
            url = data.get("@odata.nextLink")
        return items

    def list_subfolders(self, folder_path: str) -> list[dict]:
        """Return only subfolder items."""
        return [
            {"name": i["name"], "id": i.get("id", "")}
            for i in self.list_children(folder_path)
            if "folder" in i
        ]

    def list_files(self, folder_path: str, extensions: tuple[str, ...] = ()) -> list[dict]:
        """
        Return file items in a folder. Optionally filter by extension.
        Each item has keys: id, name, download_url.
        """
        children = self.list_children(folder_path)
        results = []
        for item in children:
            if "file" not in item:
                continue
            name = item.get("name", "")
            if extensions and not name.lower().endswith(extensions):
                continue
            results.append(
                {
                    "id": item.get("id", ""),
                    "name": name,
                    "download_url": item.get("@microsoft.graph.downloadUrl", ""),
                }
            )
        return sorted(results, key=lambda f: f["name"])

    # ── File Download ─────────────────────────────────────────────────────────

    def download_file(self, item_id: str, dest_path: Path) -> bool:
        """Download a file by item ID to a local path."""
        drive_id = self._get_drive_id()
        url = f"{self.base}/drives/{drive_id}/items/{item_id}/content"
        try:
            resp = requests.get(
                url, headers=self.headers, timeout=60, allow_redirects=True
            )
            resp.raise_for_status()
            dest_path.write_bytes(resp.content)
            return True
        except Exception as e:
            logger.error("Download failed for item %s: %s", item_id, e)
            return False

    def download_file_by_url(self, download_url: str, dest_path: Path) -> bool:
        """Download a file using its @microsoft.graph.downloadUrl."""
        try:
            with requests.get(download_url, stream=True, timeout=60) as r:
                r.raise_for_status()
                with open(dest_path, "wb") as f:
                    for chunk in r.iter_content(chunk_size=8192):
                        f.write(chunk)
            return True
        except Exception as e:
            logger.error("Download failed: %s", e)
            return False

    # ── File Upload ───────────────────────────────────────────────────────────

    def _get_content_type(self, filename: str) -> str:
        ext = filename.lower().split(".")[-1]
        if ext == "pdf":
            return "application/pdf"
        elif ext == "docx":
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif ext == "doc":
            return "application/msword"
        elif ext == "txt":
            return "text/plain; charset=utf-8"
        return "application/octet-stream"

    def _simple_upload(
        self, drive_id: str, folder: str, filename: str, file_path: str
    ) -> dict:
        url = (
            f"{self.base}/drives/{drive_id}"
            f"/root:/{quote(f'{folder}/{filename}')}:/content"
        )
        headers = {**self.headers, "Content-Type": self._get_content_type(filename)}
        with open(file_path, "rb") as f:
            resp = requests.put(url, headers=headers, data=f, timeout=120)
        resp.raise_for_status()
        return resp.json()

    def _resumable_upload(
        self, drive_id: str, folder: str, filename: str, file_path: str, file_size: int
    ) -> dict:
        url = (
            f"{self.base}/drives/{drive_id}"
            f"/root:/{quote(f'{folder}/{filename}')}:/createUploadSession"
        )
        body = {
            "item": {"@microsoft.graph.conflictBehavior": "replace", "name": filename}
        }
        upload_url = requests.post(
            url, headers=self.headers, json=body, timeout=30
        ).json()["uploadUrl"]

        with open(file_path, "rb") as f:
            offset = 0
            while offset < file_size:
                chunk = f.read(4 * 1024 * 1024)
                headers = {
                    "Content-Length": str(len(chunk)),
                    "Content-Range": f"bytes {offset}-{offset + len(chunk) - 1}/{file_size}",
                }
                requests.put(
                    upload_url, headers=headers, data=chunk, timeout=120
                ).raise_for_status()
                offset += len(chunk)
        return {"id": "resumable_upload_complete"}

    def upload_resume(
        self, file_path: str, target_filename: str, subfolder: str, metadata: dict
    ) -> dict:
        """Upload a resume file to Resumes/<subfolder>/<target_filename> with metadata."""
        drive_id = self._get_drive_id()
        full_folder = f"{Config.SHAREPOINT_BASE_FOLDER.strip('/')}/{subfolder}"
        self._ensure_folder(drive_id, full_folder)

        file_size = os.path.getsize(file_path)
        if file_size < 4 * 1024 * 1024:
            item = self._simple_upload(
                drive_id, full_folder, target_filename, file_path
            )
        else:
            item = self._resumable_upload(
                drive_id, full_folder, target_filename, file_path, file_size
            )

        self._set_metadata(drive_id, item["id"], metadata)
        return item

    def upload_text_file(
        self, local_path: Path, remote_path: str, skip_existing: bool = True
    ) -> bool:
        """
        Upload a .txt file to a specific remote path.
        Returns True if uploaded, False if skipped or failed.
        """
        if skip_existing and self.file_exists(remote_path):
            logger.info("  ⏭️  Skipping (already exists): %s", remote_path)
            return False

        drive_id = self._get_drive_id()

        # Ensure parent folder exists
        parent_folder = "/".join(remote_path.strip("/").split("/")[:-1])
        if parent_folder:
            self._ensure_folder(drive_id, parent_folder)

        filename = remote_path.strip("/").split("/")[-1]
        url = (
            f"{self.base}/drives/{drive_id}"
            f"/root:/{quote(remote_path.strip('/'))}:/content"
        )
        headers = {**self.headers, "Content-Type": "text/plain; charset=utf-8"}

        with open(local_path, "rb") as f:
            resp = requests.put(url, headers=headers, data=f, timeout=60)

        if resp.status_code in (200, 201):
            logger.info("  ✅ Uploaded: %s", remote_path)
            return True
        else:
            logger.error(
                "  ❌ Upload failed (%d): %s", resp.status_code, resp.text[:200]
            )
            return False


# ═══════════════════════════════════════════════════════════════════════════════
#  TEXT EXTRACTION — Inlined (PyPDF2, python-docx, OCR fallback)
# ═══════════════════════════════════════════════════════════════════════════════


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF using PyPDF2."""
    if not HAS_PYPDF2:
        logger.warning("PyPDF2 not installed. Cannot extract PDF text.")
        return ""
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        logger.error("PyPDF2 extraction failed for %s: %s", file_path, e)
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX using python-docx."""
    if not HAS_DOCX:
        logger.warning("python-docx not installed. Cannot extract DOCX text.")
        return ""
    try:
        doc = DocxDocument(file_path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    except Exception as e:
        logger.error("DOCX extraction failed for %s: %s", file_path, e)
        return ""


def extract_text_with_ocr(pdf_path: str) -> str:
    """
    OCR fallback: converts PDF pages to images with PyMuPDF, then runs
    Tesseract OCR on each page. Returns combined text.
    """
    if not HAS_PYMUPDF or not HAS_TESSERACT:
        logger.warning(
            "OCR dependencies missing (PyMuPDF=%s, Tesseract=%s). Skipping OCR.",
            HAS_PYMUPDF,
            HAS_TESSERACT,
        )
        return ""

    # Configure tesseract path if set
    if Config.TESSERACT_CMD and os.path.exists(Config.TESSERACT_CMD):
        pytesseract.pytesseract.tesseract_cmd = Config.TESSERACT_CMD

    logger.info("  🔍 Running OCR fallback on: %s", Path(pdf_path).name)
    try:
        doc = pymupdf.open(pdf_path)
        extracted_pages = []

        for page in doc:
            pix = page.get_pixmap(dpi=Config.OCR_DPI)
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            text = pytesseract.image_to_string(img)
            extracted_pages.append(text)

        doc.close()
        return "\n\n".join(extracted_pages)
    except Exception as e:
        logger.error("  ❌ OCR failed: %s", e)
        return ""


def extract_raw_text(local_path: Path) -> str:
    """
    Extracts text from a local file. Tries standard extraction first,
    falls back to OCR for PDFs if the result is too short.
    """
    suffix = local_path.suffix.lower()

    if suffix == ".pdf":
        text = extract_text_from_pdf(str(local_path))
    elif suffix in (".docx", ".doc"):
        text = extract_text_from_docx(str(local_path))
    else:
        logger.warning("  Unsupported file type: %s", suffix)
        return ""

    # OCR fallback for PDFs with empty/short text
    if suffix == ".pdf" and (not text or len(text.strip()) < 10):
        ocr_text = extract_text_with_ocr(str(local_path))
        if ocr_text and len(ocr_text.strip()) >= 10:
            logger.info("  ✅ OCR extraction successful.")
            return ocr_text

    return text


# ═══════════════════════════════════════════════════════════════════════════════
#  HELPERS — Resume Download, Dedup, Notifications
# ═══════════════════════════════════════════════════════════════════════════════


def download_resume_from_url(url: str, base_dest_path: str) -> tuple[bool, str, str]:
    """Download a resume from a URL, auto-detecting PDF vs DOCX."""
    try:
        resp = requests.get(url, timeout=60, stream=True)
        resp.raise_for_status()

        content_type = resp.headers.get("Content-Type", "").lower()
        ext = ".pdf"
        if "wordprocessingml" in content_type or "msword" in content_type:
            ext = ".docx"
        elif "pdf" in content_type:
            ext = ".pdf"
        elif ".docx" in url.lower():
            ext = ".docx"
        elif ".doc" in url.lower():
            ext = ".doc"

        temp_path = base_dest_path + ext
        with open(temp_path, "wb") as f:
            for chunk in resp.iter_content(chunk_size=8192):
                f.write(chunk)

        # Magic bytes check
        with open(temp_path, "rb") as f:
            header = f.read(4)
        if header == b"%PDF":
            ext = ".pdf"
        elif header == b"PK\x03\x04":
            ext = ".docx"

        final_path = temp_path
        if not temp_path.endswith(ext):
            final_path = base_dest_path + ext
            os.rename(temp_path, final_path)

        return True, final_path, ext
    except Exception as e:
        logger.error("Failed to download file from %s: %s", url, e)
        return False, "", ""


def get_unique_base_path(directory: str, candidate: CandidateInfo) -> str:
    base_name = f"{candidate.safe_name}_{candidate.safe_job_id}"
    h = hashlib.md5(f"{base_name}{datetime.now().isoformat()}".encode()).hexdigest()[:6]
    return os.path.join(directory, f"{base_name}_{h}")


def send_teams_notification(results: dict, candidates: list[dict]) -> None:
    """Send a summary Adaptive Card to Microsoft Teams."""
    if not Config.TEAMS_WEBHOOK_URL:
        return
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    candidate_rows = []
    for c in candidates[:20]:
        emoji = {
            "uploaded": "✅",
            "failed": "❌",
            "no_resume": "⚠️",
            "skipped": "⏭️",
        }.get(c["status"], "❓")
        candidate_rows.append(
            {
                "type": "TableRow",
                "cells": [
                    {
                        "type": "TableCell",
                        "items": [
                            {
                                "type": "TextBlock",
                                "text": c.get("name", "—"),
                                "wrap": True,
                            }
                        ],
                    },
                    {
                        "type": "TableCell",
                        "items": [
                            {
                                "type": "TextBlock",
                                "text": c.get("job_role", "—"),
                                "wrap": True,
                            }
                        ],
                    },
                    {
                        "type": "TableCell",
                        "items": [
                            {"type": "TextBlock", "text": f"{emoji} {c['status']}"}
                        ],
                    },
                ],
            }
        )

    card = {
        "type": "message",
        "attachments": [
            {
                "contentType": "application/vnd.microsoft.card.adaptive",
                "content": {
                    "$schema": "http://adaptivecards.io/schemas/adaptive-card.json",
                    "type": "AdaptiveCard",
                    "version": "1.4",
                    "body": [
                        {
                            "type": "TextBlock",
                            "size": "Large",
                            "weight": "Bolder",
                            "text": f"📄 Resume Pipeline — {now}",
                        },
                        {
                            "type": "Table",
                            "columns": [{"width": 2}, {"width": 2}, {"width": 1}],
                            "firstRowAsHeader": True,
                            "rows": [
                                {
                                    "type": "TableRow",
                                    "style": "accent",
                                    "cells": [
                                        {
                                            "type": "TableCell",
                                            "items": [
                                                {
                                                    "type": "TextBlock",
                                                    "text": "Name",
                                                    "weight": "Bolder",
                                                }
                                            ],
                                        },
                                        {
                                            "type": "TableCell",
                                            "items": [
                                                {
                                                    "type": "TextBlock",
                                                    "text": "Role",
                                                    "weight": "Bolder",
                                                }
                                            ],
                                        },
                                        {
                                            "type": "TableCell",
                                            "items": [
                                                {
                                                    "type": "TextBlock",
                                                    "text": "Status",
                                                    "weight": "Bolder",
                                                }
                                            ],
                                        },
                                    ],
                                }
                            ]
                            + candidate_rows,
                        },
                    ],
                },
            }
        ],
    }
    try:
        requests.post(
            Config.TEAMS_WEBHOOK_URL,
            json=card,
            headers={"Content-Type": "application/json"},
            timeout=15,
        )
    except Exception:
        pass


# ═══════════════════════════════════════════════════════════════════════════════
#  PIPELINE 1 — Email Fetch & Resume Upload
# ═══════════════════════════════════════════════════════════════════════════════


def run_email_fetch_pipeline(auth: GraphAuthProvider) -> None:
    """
    Monitors an Outlook mailbox for new application emails, downloads
    resumes (URL or attachment), and uploads them to SharePoint.
    """
    logger.info("=" * 70)
    logger.info("PIPELINE 1 — EMAIL FETCH & RESUME UPLOAD")
    logger.info("=" * 70)

    headers = auth.get_headers()
    fetcher = EmailFetcher(auth_headers=headers)
    candidates = fetcher.fetch_recent_emails()

    if not candidates:
        logger.info("No new application emails found. Pipeline 1 complete.")
        return

    # ── Deduplication: keep newest email per (email, job_id) ──
    unique_candidates = []
    seen: set[tuple[str, str]] = set()
    for c in candidates:
        key = (c.email.lower(), c.job_id)
        if key not in seen:
            seen.add(key)
            unique_candidates.append(c)
        else:
            logger.debug("Ignored older duplicate for %s", c.email)

    logger.info(
        "Processing %d unique candidates (from %d total emails).",
        len(unique_candidates),
        len(candidates),
    )

    # ── Process ──
    os.makedirs(Config.TEMP_DIR, exist_ok=True)
    sp = SharePointManager(auth_headers=headers)
    sp.ensure_base_folder()

    results = {
        "success": 0,
        "failed": 0,
        "skipped_no_resume": 0,
        "skipped_already_processed": 0,
    }
    notification_rows: list[dict] = []

    for candidate in unique_candidates:
        logger.info("─" * 60)
        logger.info(
            "Processing: %s | %s | Job: %s",
            candidate.name,
            candidate.email,
            candidate.job_role,
        )

        subfolder = Config.SUBFOLDER_TEMPLATE.format(
            job_id=candidate.safe_job_id, job_role=candidate.safe_job_role
        )
        local_base_path = get_unique_base_path(Config.TEMP_DIR, candidate)

        downloaded = False
        final_local_path = ""
        final_ext = ".pdf"

        # 1. Download resume
        if candidate.resume_url:
            downloaded, final_local_path, final_ext = download_resume_from_url(
                candidate.resume_url, local_base_path
            )

        if not downloaded and candidate.attachments:
            att = candidate.attachments[0]
            _, ext = os.path.splitext(att["name"])
            final_ext = ext.lower() if ext else ".pdf"
            final_local_path = local_base_path + final_ext

            try:
                content = fetcher.get_attachment_content(
                    candidate.source_email_id, att["id"]
                )
                with open(final_local_path, "wb") as f:
                    f.write(content)
                downloaded = True
            except Exception as e:
                logger.error("Attachment download failed: %s", e)

        if not downloaded:
            logger.warning("No downloadable resume for %s. Skipping.", candidate.name)
            results["skipped_no_resume"] += 1
            notification_rows.append(
                {
                    "name": candidate.name,
                    "job_role": candidate.job_role,
                    "status": "no_resume",
                }
            )
            continue

        # 2. Build filename & check SharePoint for duplicates
        target_filename = Config.FILE_NAME_TEMPLATE.format(
            name=candidate.safe_name, job_id=candidate.safe_job_id, ext=final_ext
        )
        full_sp_folder = f"{Config.SHAREPOINT_BASE_FOLDER.strip('/')}/{subfolder}"

        existing_meta = sp.get_file_metadata(full_sp_folder, target_filename)

        if existing_meta:
            if existing_meta.get("SourceEmailID") == candidate.source_email_id:
                logger.info(
                    "⏭️ Already processed this exact email on SharePoint. Skipping."
                )
                results["skipped_already_processed"] += 1
                if final_local_path and os.path.exists(final_local_path):
                    os.remove(final_local_path)
                continue
            else:
                logger.info(
                    "🔄 Found older application on SharePoint. Overwriting with latest."
                )

        # 3. Upload to SharePoint
        metadata = {
            "CandidateName": candidate.name,
            "CandidateEmail": candidate.email,
            "CandidatePhone": candidate.phone,
            "JobID": candidate.job_id,
            "JobRole": candidate.job_role,
            "SourceEmailID": candidate.source_email_id,
        }

        try:
            sp.upload_resume(final_local_path, target_filename, subfolder, metadata)
            results["success"] += 1
            status = "uploaded"
        except Exception as e:
            results["failed"] += 1
            status = "failed"
            logger.error("✗ Upload failed for %s: %s", candidate.name, e)

        notification_rows.append(
            {"name": candidate.name, "job_role": candidate.job_role, "status": status}
        )

        # 4. Cleanup
        try:
            if final_local_path and os.path.exists(final_local_path):
                os.remove(final_local_path)
        except OSError:
            pass

    # ── Notify & Summarize ──
    if notification_rows:
        send_teams_notification(results, notification_rows)

    logger.info("=" * 70)
    logger.info(
        "PIPELINE 1 COMPLETE — Uploaded: %d | Already done: %d | No resume: %d | Failed: %d",
        results["success"],
        results["skipped_already_processed"],
        results["skipped_no_resume"],
        results["failed"],
    )
    logger.info("=" * 70)


# ═══════════════════════════════════════════════════════════════════════════════
#  PIPELINE 2 — Text Extraction & Upload
# ═══════════════════════════════════════════════════════════════════════════════


def run_text_extraction_pipeline(auth: GraphAuthProvider) -> None:
    """
    Iterates through all resume files on SharePoint, extracts text
    (with OCR fallback for PDFs), and uploads .txt versions.
    """
    logger.info("=" * 70)
    logger.info("PIPELINE 2 — TEXT EXTRACTION & UPLOAD")
    logger.info("=" * 70)

    headers = auth.get_headers()
    sp = SharePointManager(auth_headers=headers)

    tmp_dir = Path(Config.TEMP_DIR) / "text_extraction"
    tmp_dir.mkdir(parents=True, exist_ok=True)

    # 1. List all job role folders under Resumes/
    logger.info("📂 Listing folders in '%s'…", Config.SHAREPOINT_BASE_FOLDER)
    folders = sp.list_subfolders(Config.SHAREPOINT_BASE_FOLDER)
    logger.info("   Found %d folders.", len(folders))

    total_uploaded = 0
    total_skipped = 0
    total_failed = 0
    failed_resumes: list[dict] = []

    for folder in sorted(folders, key=lambda f: f["name"]):
        role_name = folder["name"]
        logger.info("─" * 60)
        logger.info("📁 Processing Role: %s", role_name)

        # List resume files (PDF, DOCX, DOC) in this folder
        folder_path = f"{Config.SHAREPOINT_BASE_FOLDER}/{role_name}"
        resumes = sp.list_files(folder_path, extensions=(".pdf", ".docx", ".doc"))

        if not resumes:
            logger.info("   (No resume files found)")
            continue

        logger.info("   Found %d resume(s).", len(resumes))

        for res in resumes:
            fname = res["name"]
            txt_filename = Path(fname).stem + ".txt"
            remote_txt_path = f"{Config.TEXT_RESUMES_FOLDER}/{role_name}/{txt_filename}"

            # Check if text file already exists on SharePoint
            if sp.file_exists(remote_txt_path):
                logger.info("   ⏭️  Skipping (text already exists): %s", fname)
                total_skipped += 1
                continue

            local_resume_path = tmp_dir / fname
            local_txt_path = tmp_dir / txt_filename

            try:
                # A. Download the resume
                logger.info("   📄 Processing: %s", fname)

                download_url = res.get("download_url", "")
                if download_url:
                    ok = sp.download_file_by_url(download_url, local_resume_path)
                else:
                    ok = sp.download_file(res["id"], local_resume_path)

                if not ok:
                    logger.error("   ❌ Download failed for %s", fname)
                    total_failed += 1
                    failed_resumes.append(
                        {
                            "role": role_name,
                            "file": fname,
                            "reason": "Download failed",
                        }
                    )
                    continue

                # B. Extract text (with OCR fallback)
                text = extract_raw_text(local_resume_path)

                if not text or len(text.strip()) < 10:
                    logger.warning(
                        "   ⚠️ Extraction empty/short for %s. Skipping.", fname
                    )
                    total_failed += 1
                    failed_resumes.append(
                        {
                            "role": role_name,
                            "file": fname,
                            "reason": "Text extraction empty or too short (OCR attempted if PDF)",
                        }
                    )
                    continue

                # C. Save text locally
                local_txt_path.write_text(text, encoding="utf-8")

                # D. Upload to SharePoint
                uploaded = sp.upload_text_file(local_txt_path, remote_txt_path)
                if uploaded:
                    total_uploaded += 1
                else:
                    total_skipped += 1

                # E. Cleanup
                local_resume_path.unlink(missing_ok=True)
                local_txt_path.unlink(missing_ok=True)

            except Exception as e:
                logger.error("   ❌ Failed for %s: %s", fname, e)
                total_failed += 1
                failed_resumes.append(
                    {
                        "role": role_name,
                        "file": fname,
                        "reason": f"Exception: {str(e)}",
                    }
                )
                # Cleanup on error
                local_resume_path.unlink(missing_ok=True)
                local_txt_path.unlink(missing_ok=True)

    # ── Summary ──
    logger.info("=" * 70)
    logger.info(
        "PIPELINE 2 COMPLETE — Uploaded: %d | Skipped: %d | Failed: %d",
        total_uploaded,
        total_skipped,
        total_failed,
    )

    if failed_resumes:
        logger.info("")
        logger.info("🚨 FAILED / UNPROCESSED RESUMES:")
        logger.info("-" * 60)
        for idx, item in enumerate(failed_resumes, 1):
            logger.info(
                "  %02d. [%s] %s — %s", idx, item["role"], item["file"], item["reason"]
            )
        logger.info("-" * 60)

    logger.info("=" * 70)

    # Cleanup tmp dir
    if tmp_dir.exists():
        shutil.rmtree(tmp_dir, ignore_errors=True)


# ═══════════════════════════════════════════════════════════════════════════════
#  CLI ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════


def main():
    parser = argparse.ArgumentParser(
        description="Unified Resume Pipeline — fetch emails, upload resumes, extract text.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python resume_pipeline.py                  Run both pipelines sequentially
  python resume_pipeline.py --all            Same as above
  python resume_pipeline.py --fetch-emails   Pipeline 1 only (email → SharePoint)
  python resume_pipeline.py --extract-text   Pipeline 2 only (PDF/DOCX → .txt)
        """,
    )
    parser.add_argument(
        "--fetch-emails",
        action="store_true",
        help="Run Pipeline 1 only: fetch application emails and upload resumes to SharePoint.",
    )
    parser.add_argument(
        "--extract-text",
        action="store_true",
        help="Run Pipeline 2 only: extract text from resumes on SharePoint and upload .txt files.",
    )
    parser.add_argument(
        "--all",
        action="store_true",
        help="Run both pipelines sequentially (default if no flag is given).",
    )
    args = parser.parse_args()

    # Default: run both if no specific flag is given
    run_fetch = args.fetch_emails or args.all or (not args.fetch_emails and not args.extract_text)
    run_extract = args.extract_text or args.all or (not args.fetch_emails and not args.extract_text)

    setup_logging()

    logger.info("╔══════════════════════════════════════════════════════════════╗")
    logger.info("║          UNIFIED RESUME PIPELINE — RUN STARTED              ║")
    logger.info("╚══════════════════════════════════════════════════════════════╝")
    logger.info(
        "Mode: %s",
        "Both Pipelines"
        if (run_fetch and run_extract)
        else ("Email Fetch" if run_fetch else "Text Extraction"),
    )

    # Authenticate once, reuse across pipelines
    try:
        auth = GraphAuthProvider()
        _ = auth.get_access_token()  # Verify credentials
        logger.info("✅ Microsoft Graph authentication successful.")
    except Exception as e:
        logger.critical("❌ Authentication failed: %s", e)
        sys.exit(1)

    # Run selected pipelines
    if run_fetch:
        try:
            run_email_fetch_pipeline(auth)
        except Exception as e:
            logger.error("Pipeline 1 failed with error: %s", e, exc_info=True)

    if run_extract:
        try:
            run_text_extraction_pipeline(auth)
        except Exception as e:
            logger.error("Pipeline 2 failed with error: %s", e, exc_info=True)

    # Final cleanup
    try:
        if os.path.isdir(Config.TEMP_DIR):
            shutil.rmtree(Config.TEMP_DIR)
    except OSError:
        pass

    logger.info("╔══════════════════════════════════════════════════════════════╗")
    logger.info("║          UNIFIED RESUME PIPELINE — ALL DONE                 ║")
    logger.info("╚══════════════════════════════════════════════════════════════╝")


if __name__ == "__main__":
    main()