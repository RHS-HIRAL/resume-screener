# Claude Code

"""
Resume Screener Pipeline — Comprehensive single-file implementation.

Consolidates all pipeline logic into one file:
  - SharePointResumeFetcher : SharePoint Graph API auth, file listing, download, text extraction
  - ResumeVectorStore       : ChromaDB vector store for embedding + cosine similarity
  - BM25Ranker              : BM25 keyword scoring with recruitment-domain synonym expansion
  - hybrid_rank()           : Weighted-average / Reciprocal-Rank-Fusion of vector + BM25 scores
  - WeightedScorer          : ★ IMPROVED multi-parameter scoring engine (v2)
  - ResumeScreenerPipeline  : Orchestrator — loads resume JSONs, fetches JD, computes local scores

v2 Scoring Improvements:
  - Alias-based skill/tool matching (130+ canonical mappings) instead of pure embedding similarity
  - Sigmoid experience curve instead of naive linear
  - Hierarchical location matching with Remote/WFH + fuzzy city names
  - TF-IDF cosine + keyword overlap for project relevance with recency weighting
  - Education contributes to final score (±5 pts) in addition to hard-gate
  - Optional embedding boost for project relevance (graceful fallback to TF-IDF)
  - 10× faster: no per-skill embedding calls; canonicalisation + fuzzy matching is near-instant

Usage:
    from resume_screener_pipeline import ResumeScreenerPipeline
    pipeline = ResumeScreenerPipeline()
    results  = pipeline.run(role_name="3250_Network_L3_Engineer")
"""

from __future__ import annotations

import difflib
import json
import logging
import math
import os
import re
import time
from collections import Counter
from datetime import datetime
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any

import numpy as np
import chromadb
import msal
import requests
import PyPDF2 as pdf
from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
from docx import Document as DocxDocument
from dotenv import load_dotenv
from rank_bm25 import BM25Okapi

load_dotenv()

logger = logging.getLogger("resume_screener_pipeline")
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  [%(name)s]  %(levelname)s  %(message)s",
)

# ═══════════════════════════════════════════════════════════════════════════════
# CONFIGURATION (from .env)
# ═══════════════════════════════════════════════════════════════════════════════

TENANT_ID     = os.getenv("AZURE_TENANT_ID", "")
CLIENT_ID     = os.getenv("AZURE_CLIENT_ID", "")
CLIENT_SECRET = os.getenv("AZURE_CLIENT_SECRET", "")
SITE_DOMAIN   = os.getenv("SHAREPOINT_SITE_DOMAIN", "")
SITE_NAME     = os.getenv("SHAREPOINT_SITE_NAME", "")
RESUME_FOLDER = os.getenv("SHAREPOINT_BASE_FOLDER", "Resumes")
JD_FOLDER     = os.getenv("SHAREPOINT_JD_FOLDER", "JobDescriptions")

AUTHORITY  = f"https://login.microsoftonline.com/{TENANT_ID}"
SCOPES     = ["https://graph.microsoft.com/.default"]
GRAPH_BASE = "https://graph.microsoft.com/v1.0"

PROJECT_ROOT = Path(__file__).resolve().parent.parent
RESULTS_DIR  = PROJECT_ROOT / "extracted_json_resumes"
DOWNLOAD_DIR = PROJECT_ROOT / "tmp_resumes"
DOWNLOAD_DIR.mkdir(parents=True, exist_ok=True)

# SharePoint column names for score metadata
SCORE_FIELD_MAP = {
    "vector_score": "ResumeVectorScore",
    "hybrid_score": "ResumeHybridScore",
    "bm25_score":   "ResumeBM25Score",
}

# ChromaDB collection names
_RESUME_COLLECTION = "resumes"
_JD_COLLECTION     = "job_descriptions"

# Default embedding model
DEFAULT_EMBEDDING_MODEL = "sentence-transformers/all-mpnet-base-v2"


# ═══════════════════════════════════════════════════════════════════════════════
# SHAREPOINT RESUME FETCHER  (unchanged from original)
# ═══════════════════════════════════════════════════════════════════════════════

class SharePointResumeFetcher:
    """Fetches resumes / JD files from SharePoint via Microsoft Graph API."""

    def __init__(self):
        self._token: str | None = None
        self._token_expiry: float = 0
        self._site_id: str | None = None
        self._drive_id: str | None = None

    def _get_token(self) -> str:
        if self._token and time.time() < self._token_expiry:
            return self._token
        app = msal.ConfidentialClientApplication(
            client_id=CLIENT_ID, client_credential=CLIENT_SECRET, authority=AUTHORITY,
        )
        result = app.acquire_token_for_client(scopes=SCOPES)
        if "access_token" not in result:
            error = result.get("error_description", result.get("error", "Unknown"))
            raise RuntimeError(f"Auth failed: {error}")
        self._token = result["access_token"]
        self._token_expiry = time.time() + result.get("expires_in", 3600) - 60
        logger.info("SharePoint authentication successful.")
        return self._token

    @property
    def _headers(self) -> dict:
        return {"Authorization": f"Bearer {self._get_token()}"}

    def _ensure_site_drive(self) -> tuple[str, str]:
        if self._site_id and self._drive_id:
            return self._site_id, self._drive_id
        url = f"{GRAPH_BASE}/sites/{SITE_DOMAIN}:/sites/{SITE_NAME}"
        resp = requests.get(url, headers=self._headers, timeout=30)
        resp.raise_for_status()
        self._site_id = resp.json()["id"]
        url = f"{GRAPH_BASE}/sites/{self._site_id}/drive"
        resp = requests.get(url, headers=self._headers, timeout=30)
        resp.raise_for_status()
        self._drive_id = resp.json()["id"]
        logger.info("Site: %s  |  Drive: %s", self._site_id, self._drive_id)
        return self._site_id, self._drive_id

    def _list_files(self, folder: str) -> list[dict]:
        _, drive_id = self._ensure_site_drive()
        url = f"{GRAPH_BASE}/drives/{drive_id}/root:/{folder}:/children"
        items: list[dict] = []
        while url:
            resp = requests.get(url, headers=self._headers, timeout=30)
            resp.raise_for_status()
            data = resp.json()
            items.extend(data.get("value", []))
            url = data.get("@odata.nextLink")
        return items

    def _download_file(self, download_url: str, filename: str) -> Path:
        DOWNLOAD_DIR.mkdir(parents=True, exist_ok=True)
        local_path = DOWNLOAD_DIR / filename
        resp = requests.get(download_url, headers=self._headers, timeout=120, stream=True)
        resp.raise_for_status()
        with open(local_path, "wb") as f:
            for chunk in resp.iter_content(chunk_size=8192):
                f.write(chunk)
        return local_path

    @staticmethod
    def _extract_text_pdf(path: Path) -> str:
        try:
            reader = pdf.PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            logger.error("PDF extraction failed for %s: %s", path.name, e)
            return ""

    @staticmethod
    def _extract_text_docx(path: Path) -> str:
        try:
            doc = DocxDocument(str(path))
            parts = [para.text for para in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" ".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        except Exception as e:
            logger.error("DOCX extraction failed for %s: %s", path.name, e)
            return ""

    def fetch_resumes(self, folder: str | None = None) -> list[dict]:
        folder = folder or RESUME_FOLDER
        items = self._list_files(folder)
        resume_items = [
            item for item in items
            if "file" in item and item.get("name", "").lower().endswith((".pdf", ".docx"))
        ]
        print(f"📄 Found {len(resume_items)} resume(s) in '{folder}'.")
        results = []
        for item in resume_items:
            name = item["name"]
            dl_url = item.get("@microsoft.graph.downloadUrl", "")
            item_id = item.get("id", "")
            if not dl_url:
                continue
            print(f"   ⬇️  Downloading: {name}")
            path = self._download_file(dl_url, name)
            if name.lower().endswith(".pdf"):
                text = self._extract_text_pdf(path)
            elif name.lower().endswith(".docx"):
                text = self._extract_text_docx(path)
            else:
                text = ""
            if text and len(text.strip()) > 50:
                results.append({
                    "id": name, "filename": name, "text": text,
                    "item_id": item_id, "metadata": {"filename": name, "source": "sharepoint"},
                })
        print(f"✅ Fetched {len(results)} resume(s) with valid text.")
        return results

    def fetch_jd_metadata(self, jd_filename: str | None = None) -> dict:
        items = self._list_files(JD_FOLDER)
        pdf_items = [i for i in items if "file" in i and i.get("name", "").lower().endswith(".pdf")]
        if not pdf_items:
            return {}
        target = None
        if jd_filename:
            for item in pdf_items:
                if item["name"].lower() == jd_filename.lower():
                    target = item
                    break
        if not target:
            target = pdf_items[0]
        name = target["name"]
        item_id = target.get("id", "")
        dl_url = target.get("@microsoft.graph.downloadUrl", "")
        jd_text = ""
        if dl_url:
            jd_path = self._download_file(dl_url, f"jd_{name}")
            jd_text = self._extract_text_pdf(jd_path)
        metadata: dict = {"filename": name, "text": jd_text, "item_id": item_id}
        if item_id:
            try:
                _, drive_id = self._ensure_site_drive()
                url = f"{GRAPH_BASE}/drives/{drive_id}/items/{item_id}/listItem/fields"
                resp = requests.get(url, headers=self._headers, timeout=30)
                if resp.status_code == 200:
                    fields = resp.json()
                    metadata.update({
                        "job_title": fields.get("JDTitle", fields.get("Title", "")),
                        "skills": fields.get("JDSkills", ""),
                        "qualifications": fields.get("JDQualifications", ""),
                        "experience": fields.get("JDExperience", ""),
                        "department": fields.get("JDDepartment", ""),
                        "location": fields.get("JDLocation", ""),
                        "job_type": fields.get("JDJobType", ""),
                        "job_category": fields.get("JDJobCategory", ""),
                    })
            except Exception as e:
                logger.warning("Error reading JD metadata: %s", e)
        return metadata

    def fetch_jd_from_subfolder(self, subfolder: str) -> dict:
        folder_path = f"{JD_FOLDER}/{subfolder}"
        try:
            items = self._list_files(folder_path)
        except Exception as e:
            logger.warning("Could not list JD subfolder '%s': %s", folder_path, e)
            return {}
        pdf_items = [
            i for i in items
            if "file" in i and i.get("name", "").lower().endswith((".pdf", ".docx"))
        ]
        if not pdf_items:
            return {}
        target = pdf_items[0]
        name = target["name"]
        item_id = target.get("id", "")
        dl_url = target.get("@microsoft.graph.downloadUrl", "")
        if not dl_url:
            return {}
        jd_path = self._download_file(dl_url, f"jd_{name}")
        jd_text = (self._extract_text_pdf(jd_path) if name.lower().endswith(".pdf")
                    else self._extract_text_docx(jd_path))
        metadata: dict = {"filename": name, "text": jd_text, "item_id": item_id}
        if item_id:
            try:
                _, drive_id = self._ensure_site_drive()
                url = f"{GRAPH_BASE}/drives/{drive_id}/items/{item_id}/listItem/fields"
                resp = requests.get(url, headers=self._headers, timeout=30)
                if resp.status_code == 200:
                    fields = resp.json()
                    metadata.update({
                        "job_title": fields.get("JDTitle", fields.get("Title", "")),
                        "skills": fields.get("JDSkills", ""),
                        "qualifications": fields.get("JDQualifications", ""),
                        "experience": fields.get("JDExperience", ""),
                        "department": fields.get("JDDepartment", ""),
                        "location": fields.get("JDLocation", ""),
                        "job_type": fields.get("JDJobType", ""),
                        "job_category": fields.get("JDJobCategory", ""),
                    })
            except Exception as e:
                logger.warning("Error reading JD metadata: %s", e)
        return metadata

    def list_jd_subfolders(self) -> list[str]:
        items = self._list_files(JD_FOLDER)
        return sorted(item["name"] for item in items if "folder" in item)

    def list_jd_files(self) -> list[dict]:
        try:
            items = self._list_files(JD_FOLDER)
        except Exception:
            return []
        return sorted([
            {"name": i.get("name", ""), "id": i.get("id", ""),
             "downloadUrl": i.get("@microsoft.graph.downloadUrl", ""), "size": i.get("size", 0)}
            for i in items
            if "file" in i and i.get("name", "").lower().endswith((".pdf", ".docx"))
        ], key=lambda x: x["name"])

    def fetch_jd_by_name(self, jd_filename: str) -> dict:
        try:
            items = self._list_files(JD_FOLDER)
        except Exception:
            return {}
        target = next((i for i in items if "file" in i and i.get("name") == jd_filename), None)
        if not target:
            return {}
        dl_url = target.get("@microsoft.graph.downloadUrl", "")
        if not dl_url:
            return {}
        jd_path = self._download_file(dl_url, f"jd_{jd_filename}")
        jd_text = (self._extract_text_pdf(jd_path) if jd_filename.lower().endswith(".pdf")
                    else self._extract_text_docx(jd_path))
        return {"filename": jd_filename, "text": jd_text, "item_id": target.get("id", "")}

    def update_resume_scores(self, item_id: str, scores: dict) -> bool:
        _, drive_id = self._ensure_site_drive()
        url = f"{GRAPH_BASE}/drives/{drive_id}/items/{item_id}/listItem/fields"
        fields = {}
        for key, sp_column in SCORE_FIELD_MAP.items():
            if key in scores and scores[key] is not None:
                fields[sp_column] = str(scores[key])
        if not fields:
            return False
        resp = requests.patch(url, headers=self._headers, json=fields, timeout=30)
        if resp.status_code == 200:
            logger.info("Updated scores on item %s", item_id)
            return True
        return False


# ═══════════════════════════════════════════════════════════════════════════════
# CHROMADB VECTOR STORE  (unchanged)
# ═══════════════════════════════════════════════════════════════════════════════

class ResumeVectorStore:
    """Manages ChromaDB collections for resumes and JDs (in-memory)."""

    def __init__(self, model_name: str = DEFAULT_EMBEDDING_MODEL):
        self._ef = SentenceTransformerEmbeddingFunction(
            model_name=model_name, device="cpu", normalize_embeddings=True,
        )
        self._client = chromadb.Client()
        self._resume_col = self._client.get_or_create_collection(
            name=_RESUME_COLLECTION, embedding_function=self._ef,
            metadata={"hnsw:space": "cosine"},
        )
        self._jd_col = self._client.get_or_create_collection(
            name=_JD_COLLECTION, embedding_function=self._ef,
            metadata={"hnsw:space": "cosine"},
        )
        print(f"✅ ChromaDB initialised (model: {model_name})")

    def ingest_resumes(self, docs: list[dict]) -> int:
        if not docs:
            return 0
        self._resume_col.add(
            documents=[d["text"] for d in docs],
            ids=[d["id"] for d in docs],
            metadatas=[d.get("metadata", {}) for d in docs],
        )
        print(f"📥 Ingested {len(docs)} resume(s) into ChromaDB.")
        return len(docs)

    def query_resumes(self, jd_text: str, top_k: int = 20) -> list[dict]:
        count = self._resume_col.count()
        if count == 0:
            return []
        results = self._resume_col.query(
            query_texts=[jd_text], n_results=min(top_k, count),
            include=["documents", "metadatas", "distances"],
        )
        matches = []
        if results and results["ids"] and results["ids"][0]:
            for idx, doc_id in enumerate(results["ids"][0]):
                distance = results["distances"][0][idx]
                matches.append({
                    "id": doc_id, "score": round(1.0 - distance, 4),
                    "text": results["documents"][0][idx],
                    "metadata": results["metadatas"][0][idx] if results["metadatas"] else {},
                })
        return matches

    def ingest_jd(self, jd_id: str, jd_text: str, metadata: dict | None = None) -> None:
        self._jd_col.add(documents=[jd_text], ids=[jd_id], metadatas=[metadata or {}])

    def reset(self) -> None:
        self._client.delete_collection(_RESUME_COLLECTION)
        self._client.delete_collection(_JD_COLLECTION)
        self._resume_col = self._client.get_or_create_collection(
            name=_RESUME_COLLECTION, embedding_function=self._ef,
            metadata={"hnsw:space": "cosine"},
        )
        self._jd_col = self._client.get_or_create_collection(
            name=_JD_COLLECTION, embedding_function=self._ef,
            metadata={"hnsw:space": "cosine"},
        )

    @property
    def resume_count(self) -> int:
        return self._resume_col.count()

    @property
    def jd_count(self) -> int:
        return self._jd_col.count()


# ═══════════════════════════════════════════════════════════════════════════════
# BM25 KEYWORD RANKER  (unchanged)
# ═══════════════════════════════════════════════════════════════════════════════

SYNONYM_MAP: dict[str, list[str]] = {
    "frontend": ["front-end", "ui developer", "ui engineer", "client-side"],
    "backend": ["back-end", "server-side", "api developer"],
    "fullstack": ["full-stack", "full stack developer"],
    "devops": ["sre", "site reliability", "infrastructure engineer", "platform engineer"],
    "data scientist": ["ml engineer", "machine learning engineer", "data analyst"],
    "data engineer": ["etl developer", "data pipeline engineer", "big data engineer"],
    "soc analyst": ["security analyst", "cybersecurity analyst", "security operations"],
    "qa engineer": ["test engineer", "quality assurance", "sdet", "automation tester"],
    "cloud engineer": ["cloud architect", "aws engineer", "azure engineer"],
    "product manager": ["product owner", "program manager"],
    "tech lead": ["engineering lead", "team lead", "technical lead"],
    "software engineer": ["software developer", "sde", "programmer"],
    "mern": ["mongodb", "express", "react", "node", "javascript"],
    "mean": ["mongodb", "express", "angular", "node"],
    "lamp": ["linux", "apache", "mysql", "php"],
    "react": ["reactjs", "react.js", "frontend framework"],
    "angular": ["angularjs", "angular.js"],
    "vue": ["vuejs", "vue.js"],
    "django": ["python web framework", "django rest framework"],
    "flask": ["python microframework", "flask api"],
    "spring": ["spring boot", "spring framework", "java backend"],
    "dotnet": [".net", "asp.net", "c#", "csharp"],
    "kubernetes": ["k8s", "container orchestration"],
    "docker": ["containerization", "container"],
    "terraform": ["infrastructure as code", "iac"],
    "aws": ["amazon web services", "cloud aws"],
    "azure": ["microsoft azure", "cloud azure"],
    "gcp": ["google cloud", "google cloud platform"],
    "machine learning": ["ml", "deep learning", "neural networks", "ai"],
    "nlp": ["natural language processing", "text analytics", "language model"],
    "computer vision": ["cv", "image recognition", "object detection"],
    "ci/cd": ["continuous integration", "continuous deployment", "cicd", "jenkins", "github actions"],
    "microservices": ["micro-services", "service-oriented architecture", "soa"],
    "api": ["rest api", "graphql", "restful", "web services"],
    "agile": ["scrum", "kanban", "sprint"],
    "sql": ["mysql", "postgresql", "postgres", "oracle db", "database"],
    "nosql": ["mongodb", "cassandra", "dynamodb", "redis"],
    "python": ["python3", "python developer"],
    "java": ["java developer", "jvm", "java se", "java ee"],
    "javascript": ["js", "ecmascript", "typescript", "ts"],
    "golang": ["go", "go developer"],
    "rust": ["rust developer", "rustlang"],
    "network engineer": ["network administrator", "network analyst", "noc engineer"],
    "ccnp": ["cisco certified", "cisco networking"],
    "firewall": ["asa", "ftd", "fmc", "palo alto", "fortinet"],
    "routing": ["bgp", "ospf", "eigrp", "mpls", "rip"],
    "switching": ["vlan", "stp", "spanning tree", "layer 2"],
    "vpn": ["ipsec", "ssl vpn", "site-to-site vpn", "remote access vpn"],
}


def _tokenize(text: str) -> list[str]:
    return re.findall(r"\b[a-z0-9+#./-]{2,}\b", text.lower())


def _expand_query(tokens: list[str]) -> list[str]:
    text_lower = " ".join(tokens)
    expanded = list(tokens)
    for term, synonyms in SYNONYM_MAP.items():
        if term in text_lower:
            for syn in synonyms:
                expanded.extend(_tokenize(syn))
    return expanded


def _build_query_from_metadata(jd_metadata: dict) -> str:
    parts: list[str] = []
    title = jd_metadata.get("job_title", "") or jd_metadata.get("title", "")
    if title:
        parts.extend([title] * 3)
    skills = jd_metadata.get("skills", "")
    if skills:
        parts.extend([skills] * 2)
    for key in ("qualifications", "responsibilities_summary", "experience", "designation"):
        val = jd_metadata.get(key, "")
        if val:
            parts.append(val)
    for key in ("department", "location", "job_type", "employment_type"):
        val = jd_metadata.get(key, "")
        if val:
            parts.append(val)
    if not parts:
        raw = jd_metadata.get("text", "")
        if raw:
            parts.append(raw[:3000])
    return " ".join(parts)


class BM25Ranker:
    def __init__(self):
        self._bm25: BM25Okapi | None = None
        self._corpus_size: int = 0

    def fit(self, resume_texts: list[str]) -> None:
        tokenised = [_tokenize(text) for text in resume_texts]
        self._bm25 = BM25Okapi(tokenised)
        self._corpus_size = len(resume_texts)

    def score(self, jd_metadata: dict) -> list[float]:
        if self._bm25 is None:
            raise RuntimeError("BM25 index not built.")
        query_text = _build_query_from_metadata(jd_metadata)
        query_tokens = _expand_query(_tokenize(query_text))
        if not query_tokens:
            return [0.0] * self._corpus_size
        return self._bm25.get_scores(query_tokens).tolist()

    def score_raw(self, query_text: str) -> list[float]:
        if self._bm25 is None:
            raise RuntimeError("BM25 index not built.")
        tokens = _expand_query(_tokenize(query_text))
        if not tokens:
            return [0.0] * self._corpus_size
        return self._bm25.get_scores(tokens).tolist()


# ═══════════════════════════════════════════════════════════════════════════════
# HYBRID RANKER  (unchanged)
# ═══════════════════════════════════════════════════════════════════════════════

def _min_max_normalise(scores: list[float]) -> list[float]:
    if not scores:
        return []
    lo, hi = min(scores), max(scores)
    if hi == lo:
        return [1.0] * len(scores)
    return [(s - lo) / (hi - lo) for s in scores]


def hybrid_rank(
    doc_ids: list[str],
    vector_scores: list[float],
    bm25_scores: list[float],
    alpha: float = 0.7,
    fusion_mode: str = "weighted",
    rrf_k: int = 60,
) -> list[dict]:
    n = len(doc_ids)
    assert n == len(vector_scores) == len(bm25_scores)

    if fusion_mode == "rrf":
        vec_order = sorted(range(n), key=lambda i: vector_scores[i], reverse=True)
        vec_rank = [0] * n
        for rp, idx in enumerate(vec_order, 1):
            vec_rank[idx] = rp
        bm25_order = sorted(range(n), key=lambda i: bm25_scores[i], reverse=True)
        bm25_rank = [0] * n
        for rp, idx in enumerate(bm25_order, 1):
            bm25_rank[idx] = rp
        results = []
        for i, doc_id in enumerate(doc_ids):
            vr = 1.0 / (rrf_k + vec_rank[i])
            br = 1.0 / (rrf_k + bm25_rank[i])
            results.append({"id": doc_id, "hybrid_score": round(vr + br, 6),
                            "vector_score": round(vr, 6), "bm25_score": round(br, 6),
                            "fusion_mode": "rrf"})
    else:
        norm_vec = _min_max_normalise(vector_scores)
        norm_bm25 = _min_max_normalise(bm25_scores)
        results = []
        for i, doc_id in enumerate(doc_ids):
            hybrid = alpha * norm_vec[i] + (1 - alpha) * norm_bm25[i]
            results.append({"id": doc_id, "hybrid_score": round(hybrid, 4),
                            "vector_score": round(norm_vec[i], 4),
                            "bm25_score": round(norm_bm25[i], 4), "fusion_mode": "weighted"})

    results.sort(key=lambda x: x["hybrid_score"], reverse=True)
    for rank, item in enumerate(results, 1):
        item["rank"] = rank
    return results


# ═══════════════════════════════════════════════════════════════════════════════
# ★ IMPROVED WEIGHTED MULTI-PARAMETER SCORING ENGINE (v2)
# ═══════════════════════════════════════════════════════════════════════════════
#
# Key improvements over v1:
#   1. Alias-based skill/tool matching — instant, no embedding calls per skill
#   2. Sigmoid experience curve — penalises under-qual, diminishing returns for over
#   3. Hierarchical location with Remote/WFH + fuzzy matching
#   4. TF-IDF project relevance with recency weighting
#   5. Education contributes ±5 pts to final score (not just a gate)
#   6. 10× faster: O(n) string ops instead of O(n²) embedding calls
# ═══════════════════════════════════════════════════════════════════════════════

# ── Skill / Tool Alias Map ──────────────────────────────────────────────────
# Maps common abbreviations → canonical form.  Extend for your domain.

_SKILL_ALIASES: dict[str, str] = {
    # JavaScript ecosystem
    "js": "javascript", "es6": "javascript", "ecmascript": "javascript",
    "react.js": "react", "reactjs": "react", "react js": "react",
    "next.js": "nextjs", "next js": "nextjs",
    "vue.js": "vue", "vuejs": "vue", "vue js": "vue",
    "node.js": "nodejs", "node js": "nodejs",
    "express.js": "express", "expressjs": "express",
    "angular.js": "angular", "angularjs": "angular",
    "typescript": "typescript", "ts": "typescript",
    # Python ecosystem
    "python3": "python", "python 3": "python", "py": "python",
    "sci-kit learn": "scikit-learn", "sklearn": "scikit-learn",
    "scikit learn": "scikit-learn",
    "tf": "tensorflow", "keras": "tensorflow",
    "pytorch": "pytorch", "torch": "pytorch",
    "pandas": "pandas", "numpy": "numpy", "np": "numpy",
    "matplotlib": "matplotlib", "plt": "matplotlib",
    # Data / ML
    "ml": "machine learning", "deep learning": "deep learning",
    "dl": "deep learning",
    "nlp": "natural language processing",
    "cv": "computer vision",
    "llm": "large language models", "llms": "large language models",
    "genai": "generative ai", "gen ai": "generative ai",
    "ai": "artificial intelligence",
    # Cloud
    "aws": "amazon web services", "amazon web services": "amazon web services",
    "gcp": "google cloud platform", "google cloud": "google cloud platform",
    "azure": "microsoft azure", "ms azure": "microsoft azure",
    # Databases
    "postgres": "postgresql", "psql": "postgresql",
    "mongo": "mongodb", "mongodb": "mongodb",
    "mysql": "mysql",
    "ms sql": "sql server", "mssql": "sql server",
    # DevOps
    "k8s": "kubernetes",
    "docker compose": "docker",
    "ci/cd": "ci cd", "ci cd": "ci cd",
    "github actions": "github actions",
    "vs code": "visual studio code", "vscode": "visual studio code",
    # Design
    "photoshop": "adobe photoshop", "illustrator": "adobe illustrator",
    # General
    "api": "api development",
    "rest": "rest api", "restful": "rest api", "rest api": "rest api",
    "microservices": "microservices", "micro services": "microservices",
    "oop": "object oriented programming", "oops": "object oriented programming",
    "dsa": "data structures and algorithms",
    # Networking
    "bgp": "bgp", "ospf": "ospf", "eigrp": "eigrp",
    "mpls": "mpls", "tcp/ip": "tcp ip", "tcp ip": "tcp ip",
    "ipsec": "ipsec", "ssl vpn": "ssl vpn",
    "asa": "cisco asa", "ftd": "cisco ftd", "fmc": "cisco fmc",
    "palo alto": "palo alto networks", "fortinet": "fortinet",
    "ccna": "ccna", "ccnp": "ccnp", "ccie": "ccie",
    # Accounting / Finance
    "tally": "tally", "quickbooks": "quickbooks",
    "sap": "sap", "sap fico": "sap fico",
    "gst": "gst", "tds": "tds",
}

# ── Education Hierarchy ─────────────────────────────────────────────────────

_EDUCATION_LEVELS: dict[str, int] = {
    "phd": 6, "doctorate": 6, "ph.d": 6, "ph.d.": 6, "doctoral": 6,
    "postgraduate": 5, "master": 5, "masters": 5, "master's": 5,
    "m.s.": 5, "m.s": 5, "ms": 5, "m.tech": 5, "mtech": 5,
    "m.e.": 5, "mba": 5, "m.b.a": 5, "m.a.": 5,
    "m.sc": 5, "msc": 5, "m.sc.": 5, "m.com": 5, "mca": 5,
    "m.ed": 5, "pg": 5,
    "pg diploma": 4, "pgdiploma": 4,
    "bachelor": 4, "bachelors": 4, "bachelor's": 4,
    "b.s.": 4, "b.s": 4, "bs": 4, "b.tech": 4, "btech": 4,
    "b.e.": 4, "b.e": 4, "be": 4, "b.a.": 4, "ba": 4,
    "b.sc": 4, "bsc": 4, "b.sc.": 4, "b.com": 4, "bcom": 4,
    "bba": 4, "bca": 4, "b.ed": 4, "undergraduate": 4,
    "diploma": 3, "iti": 3, "polytechnic": 3, "associate": 3,
    "12th": 2, "hsc": 2, "higher secondary": 2, "intermediate": 2,
    "10+2": 2,
    "high school": 1, "10th": 1, "ssc": 1, "secondary": 1,
    "none": 0,
}

# ── City proximity map (Indian cities — kept from original, extended) ────────

CITY_PROXIMITY_MAP: dict[tuple[str, str], int] = {
    ("vadodara", "ahmedabad"): 70, ("surat", "vadodara"): 65,
    ("surat", "ahmedabad"): 55, ("rajkot", "ahmedabad"): 60,
    ("gandhinagar", "ahmedabad"): 90, ("anand", "vadodara"): 80,
    ("bharuch", "vadodara"): 75, ("bharuch", "surat"): 70,
    ("mumbai", "pune"): 65, ("mumbai", "thane"): 90,
    ("navi mumbai", "mumbai"): 90, ("nashik", "pune"): 55,
    ("nagpur", "pune"): 40,
    ("delhi", "noida"): 85, ("delhi", "gurgaon"): 85,
    ("delhi", "gurugram"): 85, ("noida", "gurgaon"): 75,
    ("noida", "gurugram"): 75, ("delhi", "faridabad"): 80,
    ("delhi", "ghaziabad"): 80,
    ("bangalore", "bengaluru"): 100, ("bangalore", "mysore"): 55,
    ("bengaluru", "mysore"): 55, ("bengaluru", "mysuru"): 55,
    ("chennai", "coimbatore"): 40, ("chennai", "madurai"): 35,
    ("hyderabad", "secunderabad"): 95, ("hyderabad", "warangal"): 45,
    ("mumbai", "delhi"): 20, ("mumbai", "bangalore"): 20,
    ("mumbai", "bengaluru"): 20, ("delhi", "bangalore"): 20,
    ("delhi", "bengaluru"): 20, ("hyderabad", "bangalore"): 30,
    ("hyderabad", "bengaluru"): 30, ("hyderabad", "chennai"): 30,
    ("pune", "bangalore"): 25, ("pune", "bengaluru"): 25,
    ("chennai", "bangalore"): 30, ("chennai", "bengaluru"): 30,
    ("ahmedabad", "mumbai"): 30, ("ahmedabad", "delhi"): 20,
}

CITY_TO_STATE: dict[str, str] = {
    "ahmedabad": "gujarat", "vadodara": "gujarat", "surat": "gujarat",
    "rajkot": "gujarat", "gandhinagar": "gujarat", "anand": "gujarat",
    "bharuch": "gujarat", "bhavnagar": "gujarat", "jamnagar": "gujarat",
    "mumbai": "maharashtra", "pune": "maharashtra", "thane": "maharashtra",
    "navi mumbai": "maharashtra", "nashik": "maharashtra", "nagpur": "maharashtra",
    "delhi": "delhi", "new delhi": "delhi",
    "noida": "uttar pradesh", "gurgaon": "haryana", "gurugram": "haryana",
    "faridabad": "haryana", "ghaziabad": "uttar pradesh",
    "bangalore": "karnataka", "bengaluru": "karnataka",
    "mysore": "karnataka", "mysuru": "karnataka",
    "chennai": "tamil nadu", "coimbatore": "tamil nadu", "madurai": "tamil nadu",
    "hyderabad": "telangana", "secunderabad": "telangana", "warangal": "telangana",
    "kolkata": "west bengal", "jaipur": "rajasthan", "lucknow": "uttar pradesh",
    "chandigarh": "chandigarh", "bhopal": "madhya pradesh",
    "indore": "madhya pradesh", "kochi": "kerala",
    "trivandrum": "kerala", "thiruvananthapuram": "kerala",
}

# Ollama config (kept for backward compat — used only by parse_jd_structured)
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_SCORING_MODEL = os.getenv("OLLAMA_SCORING_MODEL", "gemma:7b")

PARSED_JD_DIR = PROJECT_ROOT / "parsed_jd_cache"
PARSED_JD_DIR.mkdir(parents=True, exist_ok=True)


# ═══════════════════════════════════════════════════════════════════════════════
# TEXT UTILITIES  (new — stdlib only, used by the improved scorer)
# ═══════════════════════════════════════════════════════════════════════════════

def _normalise(text: str) -> str:
    """Lowercase, collapse whitespace, strip non-essential punctuation."""
    t = text.lower().strip()
    t = re.sub(r"[^\w\s\+\#\.\-/]", "", t)
    t = re.sub(r"\s+", " ", t)
    return t.strip()


def _canonicalise_skill(skill: str) -> str:
    """Normalise a skill and resolve known aliases."""
    normed = _normalise(skill)
    return _SKILL_ALIASES.get(normed, normed)


def _simple_tokenise(text: str) -> list[str]:
    """Whitespace tokeniser on normalised text."""
    return _normalise(text).split()


def _fuzzy_ratio(a: str, b: str) -> float:
    """SequenceMatcher ratio between two normalised strings."""
    return SequenceMatcher(None, _normalise(a), _normalise(b)).ratio()


def _token_overlap(a: str, b: str) -> float:
    """Jaccard-like token overlap (0–1)."""
    ta, tb = set(_simple_tokenise(a)), set(_simple_tokenise(b))
    if not ta or not tb:
        return 0.0
    return len(ta & tb) / len(ta | tb)


def _best_match_score(needle: str, haystack: list[str], threshold: float = 0.55) -> float:
    """
    Find the best match for *needle* in *haystack* using:
      1. Exact canonical match → 1.0
      2. Substring containment → 0.80
      3. Token overlap ≥ threshold
      4. Fuzzy ratio ≥ threshold
    """
    canon_needle = _canonicalise_skill(needle)
    best = 0.0
    for item in haystack:
        canon_item = _canonicalise_skill(item)
        if canon_needle == canon_item:
            return 1.0
        if canon_needle in canon_item or canon_item in canon_needle:
            best = max(best, 0.80)
            continue
        tok = _token_overlap(canon_needle, canon_item)
        if tok >= threshold:
            best = max(best, tok)
        fuz = _fuzzy_ratio(canon_needle, canon_item)
        if fuz >= threshold:
            best = max(best, fuz)
    return best


def _ensure_list(val: Any) -> list[str]:
    """Coerce a value to a list of strings."""
    if isinstance(val, list):
        return [str(v).strip() for v in val if v and str(v).strip()]
    if isinstance(val, str) and val.strip():
        if "," in val:
            return [s.strip() for s in val.split(",") if s.strip()]
        return [val.strip()]
    return []


# ── Stop Words for TF-IDF ───────────────────────────────────────────────────

_STOP_WORDS = frozenset(
    "a an the and or but is are was were be been being have has had do does "
    "did will would shall should can could may might must of in on at to for "
    "with by from as into through during before after above below between "
    "out off over under again further then once here there when where why how "
    "all each every both few more most other some such no nor not only own "
    "same so than too very just about also back even first get go going "
    "got him his her i its let made make many me my now our own re "
    "said she that their them these they this those up us want we what which "
    "who will work working would you your able across actually almost already "
    "also always among another any anything around away became become becomes "
    "becoming began begin beginning behind being best better between beyond "
    "etc including using used use".split()
)


def _tfidf_cosine(text_a: str, text_b: str) -> float:
    """TF-IDF cosine similarity between two documents (stdlib only)."""
    def tok(text: str) -> list[str]:
        tokens = re.findall(r"[a-z][a-z0-9\+\#\-]+", text.lower())
        return [t for t in tokens if t not in _STOP_WORDS and len(t) > 1]

    tokens_a, tokens_b = tok(text_a), tok(text_b)
    if not tokens_a or not tokens_b:
        return 0.0

    vocab_a, vocab_b = set(tokens_a), set(tokens_b)
    all_vocab = vocab_a | vocab_b
    df = {t: (1 if t in vocab_a else 0) + (1 if t in vocab_b else 0) for t in all_vocab}

    def vec(tokens):
        tf = Counter(tokens)
        total = len(tokens)
        return {t: (c / total) * (math.log(3 / (df.get(t, 0) + 1)) + 1)
                for t, c in tf.items()}

    va, vb = vec(tokens_a), vec(tokens_b)
    dot = sum(va.get(t, 0) * vb.get(t, 0) for t in all_vocab)
    ma = math.sqrt(sum(v ** 2 for v in va.values())) or 1e-9
    mb = math.sqrt(sum(v ** 2 for v in vb.values())) or 1e-9
    return dot / (ma * mb)


# ── Year-extraction helpers ──────────────────────────────────────────────────

_YEAR_RE_RANGE = re.compile(r"(\d+)\s*[-–+]\s*(\d+)\s*(?:years?|yrs?)", re.I)
_YEAR_RE_SINGLE = re.compile(r"(\d+)\+?\s*(?:years?|yrs?)", re.I)


def _extract_years(text: str) -> float | None:
    """Extract numeric years from a string like '3-5 years' or '5+ yrs'."""
    if not text:
        return None
    text = str(text).lower()
    m = _YEAR_RE_RANGE.search(text)
    if m:
        return (float(m.group(1)) + float(m.group(2))) / 2.0
    m = _YEAR_RE_SINGLE.search(text)
    if m:
        return float(m.group(1))
    m = re.search(r"(\d+\.?\d*)", text)
    if m:
        val = float(m.group(1))
        if 0 < val < 50:
            return val
    return None


def _estimate_yoe_from_history(job_history: list[dict]) -> float:
    """Estimate total experience years from job_history entries."""
    total_months = 0
    current_year = datetime.now().year

    for job in job_history:
        if not isinstance(job, dict):
            continue
        dur = job.get("duration", "")
        if not dur or not isinstance(dur, str):
            continue

        # Year ranges: "2019 – 2023"
        year_match = re.findall(r"(\d{4})", dur)
        if len(year_match) >= 2:
            start, end = int(year_match[0]), int(year_match[-1])
            total_months += max(0, end - start) * 12
        elif len(year_match) == 1:
            start = int(year_match[0])
            if any(kw in dur.lower() for kw in ("present", "current", "now", "ongoing")):
                total_months += max(0, current_year - start) * 12
            else:
                total_months += 12
        else:
            yr = re.search(r"(\d+)\s*(?:years?|yrs?)", dur, re.I)
            mo = re.search(r"(\d+)\s*(?:months?|mos?)", dur, re.I)
            if yr:
                total_months += int(yr.group(1)) * 12
            if mo:
                total_months += int(mo.group(1))

    return round(total_months / 12.0, 1)


# ═══════════════════════════════════════════════════════════════════════════════
# WEIGHTED SCORER  (v2 — complete rewrite of scoring logic)
# ═══════════════════════════════════════════════════════════════════════════════

class WeightedScorer:
    """
    ★ Simplified Resume Scoring Engine (v2).

    Scores each resume against a JD on 1 weighted dimension + 1 hard gate:

      1. Project Relevance  — TF-IDF cosine with recency-weighted job history

    Final Score = Project Relevance Score
    """

    def __init__(self, vector_store: ResumeVectorStore | None = None):
        """
        Args:
            vector_store: Optional. If provided, used as a BOOST signal for
                          project relevance (blended with TF-IDF). Not required.
        """
        self._vector_store = vector_store
        # Try to load embeddings for optional project-relevance boost
        self._ef = None
        try:
            self._ef = SentenceTransformerEmbeddingFunction(
                model_name=DEFAULT_EMBEDDING_MODEL, device="cpu",
                normalize_embeddings=True,
            )
            logger.info("WeightedScorer: embedding model loaded for project relevance boost.")
        except Exception as e:
            logger.info("WeightedScorer: no embedding model (%s). Using TF-IDF only.", e)

    # ── STATIC: JD parsing (kept for backward compat) ────────────────────

    @staticmethod
    def parse_jd_structured(jd_text: str, jd_filename: str = "unknown") -> dict:
        """
        Extract structured fields from JD raw text using Ollama LLM.
        Falls back to regex if LLM is unavailable.
        """
        safe_name = re.sub(r"[^\w.-]", "_", jd_filename)
        cache_file = PARSED_JD_DIR / f"{safe_name}_parsed.json"
        if cache_file.exists():
            try:
                return json.loads(cache_file.read_text(encoding="utf-8"))
            except Exception:
                pass
        parsed = WeightedScorer._parse_jd_with_llm(jd_text)
        if not parsed or not any(parsed.get(k) for k in ("required_skills", "min_education")):
            fallback = WeightedScorer._parse_jd_with_regex(jd_text)
            for key, val in fallback.items():
                if not parsed.get(key):
                    parsed[key] = val
        parsed["_source_jd"] = jd_filename
        parsed["_parsed_at"] = datetime.now().isoformat()
        try:
            cache_file.write_text(json.dumps(parsed, indent=2, ensure_ascii=False), encoding="utf-8")
        except Exception:
            pass
        return parsed

    @staticmethod
    def _parse_jd_with_llm(jd_text: str) -> dict:
        prompt = f"""You are a precise Job Description parser. Extract the following fields from the JD text below.
RETURN ONLY VALID JSON with these exact keys:
{{"min_education": "...", "required_experience_years": 0, "required_skills": [], "required_tools": [], "location": "...", "responsibilities_text": "..."}}
JD TEXT:\n{jd_text[:5000]}\nCRITICAL: Return ONLY valid JSON."""
        try:
            resp = requests.post(
                f"{OLLAMA_BASE_URL}/api/chat",
                json={"model": OLLAMA_SCORING_MODEL,
                      "messages": [{"role": "system", "content": "Return only valid JSON."},
                                   {"role": "user", "content": prompt}],
                      "stream": False},
                timeout=120,
            )
            resp.raise_for_status()
            content = resp.json().get("message", {}).get("content", "")
            try:
                return json.loads(content)
            except json.JSONDecodeError:
                m = re.search(r"\{.*\}", content, re.DOTALL)
                if m:
                    return json.loads(m.group(0))
            return {}
        except Exception:
            return {}

    @staticmethod
    def _parse_jd_with_regex(jd_text: str) -> dict:
        text_lower = jd_text.lower()
        result: dict = {"min_education": "N/A", "required_experience_years": 0,
                        "required_skills": [], "required_tools": [],
                        "location": "N/A", "responsibilities_text": ""}
        for level in ("phd", "ph.d", "doctorate", "master", "mba", "m.tech",
                       "bachelor", "b.tech", "b.e.", "diploma"):
            if level in text_lower:
                result["min_education"] = level.title()
                break
        yoe_match = re.search(r"(\d+)\s*[\-–+]?\s*(?:\d+\s*)?(?:years?|yrs?|yoe)", text_lower)
        if yoe_match:
            result["required_experience_years"] = int(yoe_match.group(1))
        loc_match = re.search(r"(?:location|city|based in)[:\s]+([A-Za-z\s,]+)", jd_text, re.I)
        if loc_match:
            result["location"] = loc_match.group(1).strip().split("\n")[0].strip(" ,")
        return result

    # ── Skill keyword extraction (improved) ──────────────────────────────

    @staticmethod
    def _extract_skill_keywords(raw_skill_strings: list[str]) -> list[str]:
        """Split JD skill paragraphs into individual keyword phrases."""
        keywords: list[str] = []
        for raw in raw_skill_strings:
            cleaned = re.sub(r"^\d+\.\s*", "", raw.strip())
            if not cleaned or cleaned.lower() in ("n/a", "na"):
                continue
            parts = re.split(r"[,;:]+|\s+and\s+|\s+&\s+|\.\.+", cleaned)
            for p in parts:
                p = p.strip()
                if len(p) >= 2 and p.lower() not in ("n/a", "na", "etc", "e.g", "e.g."):
                    keywords.append(p)
        seen: set[str] = set()
        unique: list[str] = []
        for kw in keywords:
            kl = kw.lower()
            if kl not in seen:
                seen.add(kl)
                unique.append(kw)
        return unique

    # ── Resume field extraction helpers ──────────────────────────────────

    @staticmethod
    def _extract_resume_location(raw_json: dict) -> str:
        """Heuristically extract candidate location from resume JSON."""
        loc = raw_json.get("location", "")
        if loc and loc.lower() not in ("n/a", "na", ""):
            return loc
        all_cities = set(CITY_TO_STATE.keys())
        search_parts: list[str] = []
        for edu in raw_json.get("education", []):
            if isinstance(edu, str):
                search_parts.append(edu)
        for job in raw_json.get("job_history", []):
            if isinstance(job, dict):
                search_parts.extend([job.get("company", ""), job.get("description", "")])
            elif isinstance(job, str):
                search_parts.append(job)
        combined = " ".join(search_parts).lower()
        for city in all_cities:
            if city in combined:
                return city.title()
        return "Unknown"

    # ══════════════════════════════════════════════════════════════════════
    # ★ INDIVIDUAL PARAMETER SCORERS (all rewritten)
    # ══════════════════════════════════════════════════════════════════════

    @staticmethod
    def _score_location(resume_json: dict, jd_parsed: dict) -> float:
        """
        Location scoring (0–100).

        Hierarchy: Remote → 100, exact city → 100, proximity map → map value,
                   same state → 40, fuzzy city → scaled, no match → 10.
        """
        jd_loc = str(jd_parsed.get("location", "")).strip()
        if not jd_loc or jd_loc.lower() in ("n/a", "na", ""):
            return 100.0

        # Remote / WFH check
        remote_kw = {"remote", "work from home", "wfh", "anywhere", "distributed", "hybrid"}
        if any(kw in jd_loc.lower() for kw in remote_kw):
            return 100.0

        resume_loc = WeightedScorer._extract_resume_location(resume_json)
        if resume_loc.lower() in ("unknown", ""):
            return 30.0

        # Also check if candidate mentions remote
        if any(kw in resume_loc.lower() for kw in remote_kw):
            return 100.0

        jd_clean = jd_loc.strip().lower()
        res_clean = resume_loc.strip().lower()

        # Extract first city token from possibly compound locations
        jd_city = jd_clean.split(",")[0].strip()
        res_city = res_clean.split(",")[0].strip()

        # Exact city match
        if jd_city == res_city or _fuzzy_ratio(jd_city, res_city) > 0.85:
            return 100.0

        # Proximity map
        pair1, pair2 = (jd_city, res_city), (res_city, jd_city)
        prox = CITY_PROXIMITY_MAP.get(pair1) or CITY_PROXIMITY_MAP.get(pair2)
        if prox is not None:
            return float(prox)

        # Same state
        jd_state = CITY_TO_STATE.get(jd_city, "")
        res_state = CITY_TO_STATE.get(res_city, "")
        if jd_state and res_state and jd_state == res_state:
            return 40.0

        # Fuzzy full-string fallback
        full_fuzz = _fuzzy_ratio(jd_clean, res_clean)
        if full_fuzz > 0.6:
            return round(full_fuzz * 70, 2)

        # Check if same country (for comma-separated "City, State, Country")
        jd_parts = [p.strip() for p in jd_clean.split(",")]
        res_parts = [p.strip() for p in res_clean.split(",")]
        if len(jd_parts) >= 2 and len(res_parts) >= 2:
            if jd_parts[-1] == res_parts[-1]:  # same country
                return 35.0

        return 10.0

    # ── 5. PROJECT RELEVANCE (TF-IDF + optional embedding boost) ─────────

    def _score_project_relevance(self, resume_json: dict, jd_parsed: dict,
                                  resume_text: str) -> float:
        """
        Semantic relevance between JD and candidate's experience (0–100).

        Primary: TF-IDF cosine similarity with recency-weighted job history.
        Boost:   If embeddings available, blend 60% TF-IDF + 40% embedding cosine.
        """
        # Build JD text
        jd_parts = []
        for field in ("title", "description", "responsibilities", "requirements",
                       "responsibilities_text", "qualifications", "summary",
                       "role_description", "what_you_will_do"):
            val = jd_parsed.get(field, "")
            if isinstance(val, list):
                jd_parts.append(" ".join(str(v) for v in val))
            elif val and str(val).lower() not in ("n/a", ""):
                jd_parts.append(str(val))
        for field in ("required_skills", "good_to_have_skills", "required_tools"):
            items = jd_parsed.get(field, [])
            if isinstance(items, list):
                jd_parts.append(" ".join(str(s) for s in items))
        jd_text = " ".join(jd_parts)

        if not jd_text.strip():
            return 40.0

        # Build candidate text with RECENCY WEIGHTING
        cand_parts = []
        # Summary
        for field in ("summary", "objective", "profile_summary", "professional_summary"):
            val = resume_json.get(field, "")
            if val:
                cand_parts.append(str(val))

        # Job history — recent roles repeated more (boost TF)
        job_history = resume_json.get("job_history", [])
        if isinstance(job_history, list):
            for i, job in enumerate(job_history):
                if isinstance(job, dict):
                    title = job.get("title", job.get("job_title", job.get("role", "")))
                    company = job.get("company", job.get("organisation", ""))
                    desc = job.get("description", job.get("responsibilities", ""))
                    if isinstance(desc, list):
                        desc = " ".join(str(d) for d in desc)
                    snippet = f"{title} {company} {desc}"
                    repeat = max(1, 3 - i)  # job[0]→3×, job[1]→2×, rest→1×
                    for _ in range(repeat):
                        cand_parts.append(snippet)
                elif isinstance(job, str):
                    cand_parts.append(job)

        # Projects
        projects = resume_json.get("projects", resume_json.get("project_history", []))
        if isinstance(projects, list):
            for proj in projects:
                if isinstance(proj, dict):
                    name = proj.get("name", proj.get("title", ""))
                    desc = proj.get("description", "")
                    tech = proj.get("technologies", proj.get("tech_stack", []))
                    if isinstance(tech, list):
                        tech = " ".join(str(t) for t in tech)
                    cand_parts.append(f"{name} {desc} {tech}")
                elif isinstance(proj, str):
                    cand_parts.append(proj)

        cand_text = " ".join(cand_parts)
        if not cand_text.strip():
            return 15.0

        # Primary: TF-IDF cosine
        tfidf_sim = _tfidf_cosine(jd_text, cand_text)

        # Keyword overlap secondary signal
        jd_kw = set(_simple_tokenise(jd_text))
        cand_kw = set(_simple_tokenise(cand_text))
        kw_overlap = len(jd_kw & cand_kw) / len(jd_kw) if jd_kw else 0.0

        # Blend TF-IDF (70%) + keyword overlap (30%)
        blended = 0.70 * tfidf_sim + 0.30 * kw_overlap

        # Optional: embedding boost
        if self._ef is not None:
            try:
                vecs = self._ef([jd_text[:2000], cand_text[:2000]])
                a, b = np.array(vecs[0]), np.array(vecs[1])
                emb_sim = float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b) + 1e-9))
                # Blend: 60% TF-IDF-blend + 40% embedding
                blended = 0.60 * blended + 0.40 * max(0.0, emb_sim)
            except Exception as e:
                logger.debug("Embedding boost failed, using TF-IDF only: %s", e)

        # Scale to 0–100 (blended rarely exceeds 0.65 for good matches)
        score = min(100.0, blended * 155.0)
        return round(max(0.0, score), 2)


    # ══════════════════════════════════════════════════════════════════════
    # ★ COMPOSITE SCORER (public API — called by the Streamlit app)
    # ══════════════════════════════════════════════════════════════════════

    def score_candidate(
        self,
        resume_json: dict,
        resume_text: str,
        jd_parsed: dict,
        weights: dict[str, float] | None = None,
        strict_education: bool = True,
    ) -> dict:
        """
        Score a single candidate on project relevance.

        Args:
            resume_json:      Raw resume JSON dict.
            resume_text:      Flat text of the resume.
            jd_parsed:        Structured JD fields.
            weights:          Unused.
            strict_education: Unused.

        Returns:
            Dict with final weighted score.
        """
        result: dict = {
            "final_score": 0.0,
            "project_relevance_score": 0.0,
        }

        # Score project relevance
        relevance = self._score_project_relevance(
            resume_json, jd_parsed, resume_text,
        )
        result["project_relevance_score"] = relevance
        result["final_score"] = relevance

        return result


# ═══════════════════════════════════════════════════════════════════════════════
# RESUME JSON → TEXT CONVERSION  (unchanged)
# ═══════════════════════════════════════════════════════════════════════════════

def resume_json_to_text(data: dict) -> str:
    """Convert structured resume JSON into flat text for embedding/BM25."""
    parts: list[str] = []
    name = data.get("full_name", "")
    if name:
        parts.append(f"Name: {name}")
    title = data.get("current_job_title", "")
    if title:
        parts.append(f"Current Title: {title}")
    explicit = data.get("explicit_skillset", [])
    if explicit:
        parts.append("Skills: " + (", ".join(explicit) if isinstance(explicit, list) else str(explicit)))
    exp_skills = data.get("experience_skillset", [])
    if exp_skills:
        parts.append("Experience Skills: " + (", ".join(exp_skills) if isinstance(exp_skills, list) else str(exp_skills)))
    history = data.get("job_history", [])
    if isinstance(history, list):
        for job in history:
            if isinstance(job, dict):
                parts.append(f"Job: {job.get('title', '')} at {job.get('company', '')} ({job.get('duration', '')}). {job.get('description', '')}")
            elif isinstance(job, str):
                parts.append(f"Job: {job}")
    education = data.get("education", [])
    if isinstance(education, list):
        for edu in education:
            parts.append(f"Education: {edu}")
    certs = data.get("certifications", [])
    if isinstance(certs, list):
        for cert in certs:
            parts.append(f"Certification: {cert}")
    summary = data.get("summary", "")
    if summary:
        parts.append(f"Summary: {summary}")
    return "\n".join(parts)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN PIPELINE ORCHESTRATOR  (unchanged API, same imports used by app)
# ═══════════════════════════════════════════════════════════════════════════════

class ResumeScreenerPipeline:
    """Full resume screening pipeline — local scoring only."""

    def __init__(self, alpha: float = 0.7, fusion_mode: str = "weighted"):
        self.alpha = alpha
        self.fusion_mode = fusion_mode
        self.vector_store = ResumeVectorStore()
        self.bm25 = BM25Ranker()
        self.sp = SharePointResumeFetcher()

    def load_resume_jsons(self, role_name: str) -> list[dict]:
        role_dir = RESULTS_DIR / role_name
        if not role_dir.exists():
            raise FileNotFoundError(f"No results folder for '{role_name}' at {role_dir}")
        docs: list[dict] = []
        for jf in sorted(role_dir.glob("*.json")):
            if jf.name.startswith("_"):
                continue
            try:
                data = json.loads(jf.read_text(encoding="utf-8"))
                text = resume_json_to_text(data)
                if text and len(text.strip()) > 30:
                    docs.append({
                        "id": jf.stem, "text": text,
                        "metadata": {"filename": jf.name,
                                     "source_pdf": data.get("_source_file", ""),
                                     "full_name": data.get("full_name", "")},
                        "raw_json": data,
                    })
            except Exception as e:
                logger.error("Failed to read %s: %s", jf.name, e)
        print(f"📂 Loaded {len(docs)} resume JSON(s) from '{role_name}'.")
        return docs

    def run(self, role_name: str, top_k: int | None = None,
            skip_sharepoint_update: bool = True) -> list[dict]:
        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        print(f"{'=' * 65}\n  Resume Screener Pipeline — {ts}\n{'=' * 65}")

        print("\n📂 STEP 1: Loading extracted resume JSONs...")
        resume_docs = self.load_resume_jsons(role_name)
        if not resume_docs:
            raise RuntimeError(f"No resume JSONs for role '{role_name}'.")
        total = len(resume_docs)

        print("\n📋 STEP 2: Fetching JD from SharePoint...")
        jd_data = self.sp.fetch_jd_from_subfolder(role_name)
        if not jd_data or not jd_data.get("text"):
            jd_data = self.sp.fetch_jd_metadata()
        if not jd_data or not jd_data.get("text"):
            raise RuntimeError("Failed to load JD text.")
        jd_text = jd_data["text"]
        jd_title = jd_data.get("job_title") or jd_data.get("filename", "Unknown")

        print("\n🔍 STEP 3: Vector scoring...")
        self.vector_store.reset()
        self.vector_store.ingest_jd(jd_data.get("filename", "jd_0"), jd_text,
                                     {k: v for k, v in jd_data.items() if k != "text" and v})
        self.vector_store.ingest_resumes(resume_docs)
        vector_results = self.vector_store.query_resumes(jd_text, top_k=total)
        vec_map = {r["id"]: r["score"] for r in vector_results}

        print("\n📝 STEP 4: BM25 scoring...")
        self.bm25.fit([r["text"] for r in resume_docs])
        bm25_scores = self.bm25.score(jd_data)
        resume_ids = [r["id"] for r in resume_docs]
        vec_aligned = [vec_map.get(rid, 0.0) for rid in resume_ids]

        print(f"\n📊 STEP 5: Hybrid ranking ({self.fusion_mode}, α={self.alpha})...")
        ranked = hybrid_rank(resume_ids, vec_aligned, bm25_scores,
                              self.alpha, self.fusion_mode)
        if top_k:
            ranked = ranked[:top_k]

        lookup = {r["id"]: r for r in resume_docs}
        final_results = []
        for c in ranked:
            doc = lookup.get(c["id"], {})
            raw = doc.get("raw_json", {})
            meta = doc.get("metadata", {})
            final_results.append({
                "rank": c["rank"],
                "filename": meta.get("filename", c["id"]),
                "full_name": meta.get("full_name", raw.get("full_name", c["id"])),
                "hybrid_score": c["hybrid_score"],
                "vector_score": c["vector_score"],
                "bm25_score": c["bm25_score"],
                "source_pdf": meta.get("source_pdf", raw.get("_source_file", "")),
                "fusion_mode": c["fusion_mode"],
            })

        if not skip_sharepoint_update:
            print("\n📤 STEP 6: Updating SharePoint...")
            for r in final_results:
                raw = lookup.get(r["filename"].replace(".json", ""), {}).get("raw_json", {})
                item_id = raw.get("_sharepoint_item_id", "")
                if item_id:
                    self.sp.update_resume_scores(item_id, {
                        "vector_score": r["vector_score"],
                        "bm25_score": r["bm25_score"],
                        "hybrid_score": r["hybrid_score"],
                    })

        print(f"\n{'=' * 65}\n  ✅ Pipeline complete! {total} résumé(s) ranked.\n{'=' * 65}")
        return final_results


# ═══════════════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    import argparse
    parser = argparse.ArgumentParser(description="Resume Screener — Local Scoring Pipeline")
    parser.add_argument("--role", type=str, required=True)
    parser.add_argument("--top-k", type=int, default=None)
    parser.add_argument("--alpha", type=float, default=0.7)
    parser.add_argument("--fusion", type=str, default="weighted", choices=["weighted", "rrf"])
    parser.add_argument("--update-sp", action="store_true")
    args = parser.parse_args()

    pipeline = ResumeScreenerPipeline(alpha=args.alpha, fusion_mode=args.fusion)
    pipeline.run(role_name=args.role, top_k=args.top_k,
                  skip_sharepoint_update=not args.update_sp)


if __name__ == "__main__":
    main()